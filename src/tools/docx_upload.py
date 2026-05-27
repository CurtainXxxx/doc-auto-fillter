"""S3 上传与文档校验模块。

将原本散落在三个 generate_* 函数中的重复代码统一抽取：
  - S3SyncStorage 初始化（环境变量配置）
  - 文件上传 + 签名 URL 生成
  - 本地副本保存
  - 生成后文档校验（validate_doc）

使用方式:
    from tools.docx_upload import upload_and_validate
    result = upload_and_validate(docx_bytes, "edu_report", template_path)
"""

import logging
import os
import time
from typing import Any

from docx import Document
from tools.docx_validator import validate_docx, diff_docx, fix_docx

logger = logging.getLogger(__name__)

# ── 环境变量键名常量 ──────────────────────────────────────

_ENV_ENDPOINT_URL = "COZE_BUCKET_ENDPOINT_URL"
_ENV_ACCESS_KEY = "COZE_BUCKET_ACCESS_KEY"
_ENV_SECRET_KEY = "COZE_BUCKET_SECRET_KEY"
_ENV_BUCKET_NAME = "COZE_BUCKET_NAME"
_ENV_REGION = "COZE_BUCKET_REGION"

# ── 默认值 ────────────────────────────────────────────────

_DEFAULT_REGION = "cn-beijing"
_DEFAULT_EXPIRE_SECONDS = 86400  # 24h
_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)


def _create_storage() -> Any:
    """根据环境变量创建 S3SyncStorage 实例。

    所有配置项均从环境变量读取，不再硬编码。
    缺少必要变量时抛出 EnvironmentError。

    Returns:
        S3SyncStorage 实例

    Raises:
        EnvironmentError: 缺少必要的 bucket 配置
    """
    # 延迟导入——此包仅沙箱环境可用，本地开发走 mock
    from coze_coding_dev_sdk.s3 import S3SyncStorage

    endpoint_url = os.getenv(_ENV_ENDPOINT_URL)
    bucket_name = os.getenv(_ENV_BUCKET_NAME)

    if not bucket_name:
        raise EnvironmentError(
            f"缺少必要的环境变量 {_ENV_BUCKET_NAME}，无法初始化 S3 存储"
        )

    return S3SyncStorage(
        endpoint_url=endpoint_url,
        access_key=os.getenv(_ENV_ACCESS_KEY, ""),
        secret_key=os.getenv(_ENV_SECRET_KEY, ""),
        bucket_name=bucket_name,
        region=os.getenv(_ENV_REGION, _DEFAULT_REGION),
    )


def _generate_file_name(prefix: str, display_name: str) -> str:
    """生成唯一的 S3 文件名。

    Args:
        prefix: 路径前缀，如 "edu_report"、"custom_report"、"edu_form"
        display_name: 展示用名称（含空格会被替换为下划线）

    Returns:
        格式: {prefix}/{safe_name}_{timestamp}.docx
    """
    safe_name = display_name.replace(" ", "_")
    timestamp = time.strftime("%Y%m%d%H%M%S")
    return f"{prefix}/{safe_name}_{timestamp}.docx"


def _save_local_copy(docx_bytes: bytes, file_name: str) -> str:
    """将 docx 字节流保存到 /tmp 本地副本。

    Args:
        docx_bytes: 文档字节流
        file_name: S3 文件名（仅取 basename 作为本地文件名）

    Returns:
        本地文件绝对路径
    """
    local_name = os.path.basename(file_name)
    local_path = os.path.join("/tmp", local_name)

    with open(local_path, "wb") as f:
        f.write(docx_bytes)

    logger.debug("本地副本已保存: %s", local_path)
    return local_path


def upload_and_validate(
    docx_bytes: bytes,
    name_prefix: str,
    display_name: str,
    template_path: str,
    *,
    validate: bool = True,
) -> dict:
    """上传 docx 到 S3 + 保存本地副本 + 可选校验，返回统一结果。

    Args:
        docx_bytes: 文档字节流
        name_prefix: S3 路径前缀，如 "edu_report"、"custom_report"、"edu_form"
        display_name: 展示用名称（用于生成文件名）
        template_path: 模板文件路径（校验用）
        validate: 是否执行文档校验，默认 True

    Returns:
        {
            "file_name": str,       # S3 文件名
            "download_url": str,    # 签名下载 URL（24h 有效）
            "local_path": str,      # 本地副本路径
            "validation": dict|None # 校验结果（validate=True 时）
        }

    Raises:
        EnvironmentError: S3 配置缺失
        Exception: 上传失败
    """
    storage = _create_storage()
    file_name = _generate_file_name(name_prefix, display_name)

    # 上传
    file_key = storage.upload_file(
        file_content=docx_bytes,
        file_name=file_name,
        content_type=_DOCX_CONTENT_TYPE,
    )
    download_url = storage.generate_presigned_url(
        key=file_key,
        expire_time=_DEFAULT_EXPIRE_SECONDS,
    )
    logger.info("文档已上传: %s → %s", file_name, file_key)

    # 本地副本
    local_path = _save_local_copy(docx_bytes, file_name)

    # 校验
    validation = None
    if validate:
        validation = validate_doc(template_path, local_path)

    return {
        "file_name": file_name,
        "download_url": download_url,
        "local_path": local_path,
        "validation": validation,
    }


# ── 文档校验 ──────────────────────────────────────────────


def validate_doc(template_path: str, output_path: str) -> dict:
    """对比模板和生成文档的结构完整性（增强版校验管线）。

    整合了 docx_validator 模块的多层校验：
    1. 生成文件能否被 python-docx 正常打开
    2. OpenXML元素顺序校验（pPr→runs, rPr→t, tcPr→p, sectPr位置）
    3. 单元格最少段落校验
    4. 表格维度对比（数量、行数、列数）
    5. 合并单元格连续性校验
    6. 格式污染检测
    7. 模板与填写文档对比（diff）

    Args:
        template_path: 原始模板文件路径
        output_path: 生成的文档文件路径

    Returns:
        {"valid": bool, "errors": [...], "warnings": [...], "diff": dict, "checks": list}
    """
    errors: list[str] = []
    warnings: list[str] = []

    # 1. 检查生成文件能否打开
    try:
        output_doc = Document(output_path)
    except Exception as e:
        return {"valid": False, "errors": [f"生成的文档无法打开: {e}"], "warnings": [], "diff": None, "checks": []}

    try:
        template_doc = Document(template_path)
    except Exception as e:
        warnings.append(f"模板文件无法打开，跳过结构对比: {e}")
        return {"valid": True, "errors": [], "warnings": warnings, "diff": None, "checks": []}

    # 2. 运行校验管线（validate_docx）
    validation = validate_docx(output_doc)

    # 合并管线校验结果
    errors.extend(validation["errors"])
    warnings.extend(validation["warnings"])

    # 3. 表格数量和维度对比
    t_tables = len(template_doc.tables)
    o_tables = len(output_doc.tables)
    if t_tables != o_tables:
        errors.append(f"表格数量变化: 模板{t_tables}个, 生成{o_tables}个")

    for i in range(min(t_tables, o_tables)):
        t_rows = len(template_doc.tables[i].rows)
        o_rows = len(output_doc.tables[i].rows)
        if t_rows != o_rows:
            if o_rows < t_rows:
                errors.append(f"表格{i+1}行数减少: 模板{t_rows}行, 生成{o_rows}行")
            else:
                warnings.append(f"表格{i+1}行数增加: 模板{t_rows}行, 生成{o_rows}行（可能是行组填充）")

        t_cols = len(template_doc.tables[i].columns)
        o_cols = len(output_doc.tables[i].columns)
        if t_cols != o_cols:
            errors.append(f"表格{i+1}列数变化: 模板{t_cols}列, 生成{o_cols}列")

    # 4. 生成diff对比
    diff_result = None
    try:
        diff_result = diff_docx(template_path, output_path)
    except Exception as e:
        warnings.append(f"文档对比失败: {e}")

    # 5. 如果有可自动修复的错误，尝试修复
    if not validation["passed"]:
        try:
            fix_result = fix_docx(output_path)
            if fix_result["fixed"]:
                warnings.append(f"已自动修复 {len(fix_result['fixes'])} 个结构问题")
        except Exception:
            pass  # 修复失败不影响主流程

    return {
        "valid": len(errors) == 0,
        "errors": errors,
        "warnings": warnings,
        "diff": diff_result,
        "checks": validation["checks"],
    }
