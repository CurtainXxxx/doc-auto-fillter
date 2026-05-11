"""将 docx 文档转为 HTML 预览，标记可填写区域

输出：
- html: 文档的 HTML 渲染（表格、段落），可填写单元格带 contenteditable
- field_map: field_id → { label, table_idx, row_idx, col_idx } 映射
"""

import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 占位符文字
_PLACEHOLDER_TEXTS = {'%', '…', '……', '...', '—', '--', '___', '□', '○'}


def _resolve_template_path(template_path: str) -> str:
    """将模板名或路径解析为实际文件路径"""
    # 如果已经是存在的文件路径，直接返回
    if os.path.isfile(template_path):
        return template_path

    workspace = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")

    # 尝试作为内置模板名解析
    _BUILTIN_TEMPLATES = {
        "评价报告": "2023-2024-2《xxx》 岭南师范学院专业课程目标达成度评价报告模板.docx",
        "试卷分析": "2023-2024-2《xxx》 试卷分析模板.docx",
        "关联矩阵": "2023-2024-2《xxx》岭南师范学院考题与课程目标及毕业要求关联矩阵表模板.docx",
    }
    if template_path in _BUILTIN_TEMPLATES:
        fname = _BUILTIN_TEMPLATES[template_path]
        full = os.path.join(workspace, "assets", fname)
        if os.path.isfile(full):
            return full

    # 尝试作为 assets 下的文件名
    full = os.path.join(workspace, "assets", template_path)
    if os.path.isfile(full):
        return full

    # 尝试作为 workspace 下的相对路径
    full = os.path.join(workspace, template_path)
    if os.path.isfile(full):
        return full

    # 最后尝试原始路径（可能不存在，让 Document() 报错）
    return template_path


def docx_to_html(template_path: str) -> dict:
    """将 docx 转为 HTML，返回 { html, field_map }"""
    resolved = _resolve_template_path(template_path)
    doc = Document(resolved)

    # 1. 用 analyze_template 获取字段信息
    from tools.template_analyzer import analyze_template
    analysis = analyze_template(resolved)

    # 建立 field_id → field 信息映射
    field_map = {}
    for f in analysis["label_fields"]:
        fid = f.get("field_id", "")
        if fid:
            field_map[fid] = {
                "label": f["label"],
                "table_idx": f["table_idx"],
                "row_idx": f["row_idx"],
                "col_idx": f["col_idx"],
                "line_idx": f.get("line_idx", 0),
                "fill_mode": f["fill_mode"],
                "existing_value": f.get("existing_value", ""),
                "pattern": f.get("pattern", ""),
                "repeat_count": f.get("repeat_count", 1),
                "row_indices": f.get("row_indices", [f["row_idx"]]),
            }

    # 2. 构建 (table_idx, row_idx, col_idx) → [field_info] 映射
    # 冒号模式可能一个格有多个字段（多行）
    cell_fields = {}  # (t, r, c) → [field_info]
    for fid, info in field_map.items():
        for ri in info.get("row_indices", [info["row_idx"]]):
            key = (info["table_idx"], ri, info["col_idx"])
            if key not in cell_fields:
                cell_fields[key] = []
            cell_fields[key].append({**info, "field_id": fid})

    # 3. 构建 label → field_ids 映射
    label_to_fields = {}
    for fid, info in field_map.items():
        label = info["label"]
        if label not in label_to_fields:
            label_to_fields[label] = []
        label_to_fields[label].append(fid)

    # 4. 遍历文档生成 HTML
    parts = []

    body_el = doc.element.body  # type: ignore[attr-defined]
    for element in body_el:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'tbl':
            t_idx = None
            for i, table in enumerate(doc.tables):
                if table._element is element:
                    t_idx = i
                    break
            if t_idx is None:
                continue

            table = doc.tables[t_idx]
            parts.append(_render_table(table, t_idx, cell_fields))

        elif tag == 'p':
            para_text = ""
            for child in element.iter(qn('w:t')):
                if child.text:
                    para_text += child.text
            if para_text.strip():
                align = ""
                pPr = element.find(qn('w:pPr'))
                if pPr is not None:
                    jc = pPr.find(qn('w:jc'))
                    if jc is not None:
                        align = jc.get(qn('w:val'), '')

                style = ""
                if align == 'center':
                    style = ' style="text-align:center"'
                elif align == 'right':
                    style = ' style="text-align:right"'

                parts.append(f'<p{style}>{_escape(para_text.strip())}</p>')

    html = '\n'.join(parts)

    return {
        "html": html,
        "field_map": field_map,
        "label_to_fields": label_to_fields,
    }


def _render_table(table, t_idx, cell_fields) -> str:
    """渲染一个表格为 HTML，标记可填写单元格"""
    rows_html = []

    for r_idx, row in enumerate(table.rows):
        cells_html = []
        unique_cells = _get_unique_cells(row)

        for c_idx, cell in enumerate(unique_cells):
            text = cell.text.strip()

            # 处理合并单元格
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            colspan = 1
            attrs = ""

            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    colspan = int(gs.get(qn('w:val'), 1))
                    if colspan > 1:
                        attrs += f' colspan="{colspan}"'
                vm = tcPr.find(qn('w:vMerge'))
                if vm is not None:
                    vm_val = vm.get(qn('w:val'), '')
                    if vm_val == 'restart':
                        rowspan_val = _calc_rowspan(table, r_idx, c_idx)
                        if rowspan_val > 1:
                            attrs += f' rowspan="{rowspan_val}"'
                    else:
                        continue  # 续行跳过

            cell_style = _get_cell_style(cell)
            key = (t_idx, r_idx, c_idx)
            fields_in_cell = cell_fields.get(key, [])

            if fields_in_cell:
                # 有字段映射的格 - 渲染为标签+可编辑值
                cells_html.append(
                    _render_field_cell(attrs, cell_style, text, fields_in_cell)
                )
            elif _is_fillable_cell(text):
                # 空白格但没有字段映射 - 仍标记为可编辑
                fid = f"T{t_idx}_R{r_idx}_C{c_idx}"
                cells_html.append(
                    f'<td{attrs} style="{cell_style}" class="doc-cell editable empty" '
                    f'data-field-id="{_escape(fid)}" contenteditable="true"></td>'
                )
            else:
                # 普通标签格
                display_text = _escape(text) if text else '&nbsp;'
                cells_html.append(
                    f'<td{attrs} style="{cell_style}" class="doc-cell label">{display_text}</td>'
                )

        if cells_html:
            rows_html.append('<tr>' + ''.join(cells_html) + '</tr>')

    return '<table class="doc-table">' + '\n'.join(rows_html) + '</table>'


def _render_field_cell(attrs, cell_style, text, fields) -> str:
    """渲染包含字段的单元格

    冒号模式: "课程名称:" → "课程名称: [可编辑空白]"
    label_blank模式: 整个格可编辑
    """
    # 判断是否是冒号模式
    has_colon = any(f["pattern"] == "colon" for f in fields)

    if has_colon:
        # 冒号模式 - 需要拆分标签和值
        # 按行拆分文本
        lines = text.split('\n')
        if len(fields) == 1 and len(lines) <= 1:
            # 单字段冒号格: "课程名称:" → 标签 + 可编辑
            f = fields[0]
            fid = f["field_id"]
            label = f["label"]
            return (
                f'<td{attrs} style="{cell_style}" class="doc-cell colon-field">'
                f'<span class="cell-label">{_escape(label)}：</span>'
                f'<span class="cell-value editable empty" '
                f'data-field-id="{_escape(fid)}" '
                f'data-label="{_escape(label)}" '
                f'contenteditable="true"></span>'
                f'</td>'
            )
        else:
            # 多行冒号格（如 "考试类别:\n平时:\n期末:"）
            parts = []
            for i, f in enumerate(fields):
                fid = f["field_id"]
                label = f["label"]
                parts.append(
                    f'<div class="colon-line">'
                    f'<span class="cell-label">{_escape(label)}：</span>'
                    f'<span class="cell-value editable empty" '
                    f'data-field-id="{_escape(fid)}" '
                    f'data-label="{_escape(label)}" '
                    f'contenteditable="true"></span>'
                    f'</div>'
                )
            return (
                f'<td{attrs} style="{cell_style}" class="doc-cell colon-field multi-line">'
                + ''.join(parts)
                + '</td>'
            )
    else:
        # 非冒号模式 - 整个格可编辑
        f = fields[0]
        fid = f["field_id"]
        label = f["label"]
        existing = f.get("existing_value", "")

        if existing and existing not in _PLACEHOLDER_TEXTS:
            # 已有内容 - 显示并可编辑
            return (
                f'<td{attrs} style="{cell_style}" class="doc-cell editable filled" '
                f'data-field-id="{_escape(fid)}" '
                f'data-label="{_escape(label)}" '
                f'contenteditable="true">{_escape(existing)}</td>'
            )
        else:
            # 空白/占位符格
            return (
                f'<td{attrs} style="{cell_style}" class="doc-cell editable empty" '
                f'data-field-id="{_escape(fid)}" '
                f'data-label="{_escape(label)}" '
                f'contenteditable="true"></td>'
            )


def _is_fillable_cell(text: str) -> bool:
    """判断单元格是否是可填写的（空白或仅含占位符）"""
    if not text:
        return True
    if text in _PLACEHOLDER_TEXTS:
        return True
    return False


def _get_unique_cells(row) -> list:
    """获取行中的独立单元格 (跳过被合并重复的)"""
    seen = set()
    cells = []
    for cell in row.cells:
        cid = id(cell._element)
        if cid not in seen:
            seen.add(cid)
            cells.append(cell)
    return cells


def _calc_rowspan(table, r_idx, c_idx):
    """计算从 r_idx 开始的垂直合并行数"""
    span = 1
    for r in range(r_idx + 1, len(table.rows)):
        row = table.rows[r]
        unique = _get_unique_cells(row)
        if c_idx < len(unique):
            tc = unique[c_idx]._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                vm = tcPr.find(qn('w:vMerge'))
                if vm is not None:
                    vm_val = vm.get(qn('w:val'), '')
                    if vm_val != 'restart':
                        span += 1
                    else:
                        break
                else:
                    break
            else:
                break
        else:
            break
    return span


def _get_cell_style(cell) -> str:
    """提取单元格的基本样式"""
    styles = []

    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'), '')
            if fill and fill.upper() not in ('FFFFFF', 'AUTO') and fill != 'auto':
                styles.append(f'background-color:#{fill}')

    for para in cell.paragraphs:
        for run in para.runs:
            if run.font.size:
                pt = run.font.size.pt
                if pt:
                    styles.append(f'font-size:{pt}pt')
                    break
        if styles:
            break

    for para in cell.paragraphs:
        if para.alignment:
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                styles.append('text-align:center')
            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                styles.append('text-align:right')
            break

    return ';'.join(styles)


def _escape(text: str) -> str:
    """HTML 转义"""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#39;'))
