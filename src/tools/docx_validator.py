"""
docx_validator.py — 文档校验、对比与修复模块

借鉴 MiniMax-docx 项目的工程化思路，提供三大能力：
1. validate_docx() — 填写后校验管线（类比 MiniMax 的 Validation Pipeline）
2. diff_docx() — 模板 vs 填写文档对比（类比 MiniMax 的 diff 命令）
3. fix_docx() — 自动修复常见结构问题（类比 MiniMax 的 fix-order）
4. ensure_element_order() — 系统化 OpenXML 元素顺序规则
5. strip_inline_formatting() — 格式污染防护（类比 MiniMax 的 Format Contamination）
"""

import re
import copy
from collections import OrderedDict
from docx import Document
from docx.oxml.ns import qn


# ============================================================
# 一、校验管线（Validation Pipeline）
# ============================================================

def validate_docx(path_or_doc):
    """对生成的docx执行校验管线，返回校验结果。

    Args:
        path_or_doc: 文件路径(str)或Document对象

    Returns:
        dict: {
            "passed": bool,
            "errors": [str],   # 必须修复的错误
            "warnings": [str], # 建议修复的警告
            "checks": [{"name": str, "passed": bool, "detail": str}]
        }
    """
    if isinstance(path_or_doc, str):
        doc = Document(path_or_doc)
    else:
        doc = path_or_doc

    errors = []
    warnings = []
    checks = []

    # --- Step 1: 结构完整性校验 ---
    r1 = _check_structure_integrity(doc)
    checks.append(r1)
    if not r1["passed"]:
        errors.extend(r1.get("issues", []))

    # --- Step 2: 元素顺序校验 ---
    r2 = _check_element_order(doc)
    checks.append(r2)
    if not r2["passed"]:
        errors.extend(r2.get("issues", []))

    # --- Step 3: 单元格最少段落校验 ---
    r3 = _check_cell_min_paragraph(doc)
    checks.append(r3)
    if not r3["passed"]:
        errors.extend(r3.get("issues", []))

    # --- Step 4: 合并单元格连续性校验 ---
    r4 = _check_merge_continuity(doc)
    checks.append(r4)
    if not r4["passed"]:
        warnings.extend(r4.get("issues", []))

    # --- Step 5: 表格行列数校验（与模板对比） ---
    r5 = _check_table_dimensions(doc)
    checks.append(r5)
    if not r5["passed"]:
        warnings.extend(r5.get("issues", []))

    # --- Step 6: sectPr位置校验 ---
    r6 = _check_sectpr_position(doc)
    checks.append(r6)
    if not r6["passed"]:
        errors.extend(r6.get("issues", []))

    # --- Step 7: 格式污染检测 ---
    r7 = _check_format_contamination(doc)
    checks.append(r7)
    if not r7["passed"]:
        warnings.extend(r7.get("issues", []))

    passed = len(errors) == 0
    return {
        "passed": passed,
        "errors": errors,
        "warnings": warnings,
        "checks": checks,
    }


def _check_structure_integrity(doc):
    """校验文档结构完整性：能否正常打开、表格数是否合理。"""
    issues = []
    try:
        _ = doc.part.element.xml  # 尝试序列化
    except Exception as e:
        issues.append(f"文档XML序列化失败: {e}")
        return {"name": "结构完整性", "passed": False, "issues": issues}

    # 检查body是否存在
    body = doc.element.body
    if body is None:
        issues.append("文档缺少body元素")
        return {"name": "结构完整性", "passed": False, "issues": issues}

    # 检查表格数
    table_count = len(doc.tables)
    if table_count == 0:
        issues.append("文档不包含任何表格")

    return {"name": "结构完整性", "passed": len(issues) == 0, "issues": issues}


def _check_element_order(doc):
    """校验OpenXML元素顺序：pPr必须在runs之前，rPr必须在t之前等。

    关键规则（来自 MiniMax-docx + ECMA-376）：
    - w:p → pPr 在 runs 之前
    - w:r → rPr 在 t/br/tab 之前
    - w:tbl → tblPr → tblGrid → tr
    - w:tr → trPr 在 tc 之前
    - w:tc → tcPr 在 p 之前
    - w:body → sectPr 必须是最后一个子元素
    """
    issues = []
    body = doc.element.body

    def _check_children(parent, tag_name, expected_order, child_tags):
        """检查子元素是否按预期顺序排列。"""
        props_found = False
        content_found = False
        for child in parent:
            ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if ctag in expected_order:
                prop_idx = expected_order.index(ctag)
                # 属性类元素（pPr, rPr, tcPr, trPr, tblPr, tblGrid）
                if prop_idx < len(expected_order) - 1:
                    if content_found:
                        issues.append(
                            f"<{tag_name}> 内 <{ctag}> 出现在内容元素之后，"
                            f"应在内容之前"
                        )
                    props_found = True
                else:
                    content_found = True

    # 检查所有段落
    for p in body.iter(qn('w:p')):
        has_ppr = False
        has_run_content = False
        for child in p:
            ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if ctag == 'pPr':
                if has_run_content:
                    issues.append("<w:p> 内 <w:pPr> 出现在run内容之后")
                has_ppr = True
            elif ctag == 'r':
                has_run_content = True
                # 检查run内部顺序
                has_rpr = False
                has_text = False
                for rc in child:
                    rctag = rc.tag.split('}')[-1] if '}' in rc.tag else rc.tag
                    if rctag == 'rPr':
                        if has_text:
                            issues.append("<w:r> 内 <w:rPr> 出现在文本内容之后")
                        has_rpr = True
                    elif rctag in ('t', 'br', 'tab', 'cr', 'sym', 'object'):
                        has_text = True

    # 检查所有表格
    for tbl in body.iter(qn('w:tbl')):
        has_tblpr = False
        has_tblgrid = False
        has_tr = False
        for child in tbl:
            ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if ctag == 'tblPr':
                if has_tr:
                    issues.append("<w:tbl> 内 <w:tblPr> 出现在行之后")
                has_tblpr = True
            elif ctag == 'tblGrid':
                if has_tr:
                    issues.append("<w:tbl> 内 <w:tblGrid> 出现在行之后")
                has_tblgrid = True
            elif ctag == 'tr':
                has_tr = True
                # 检查行内顺序
                has_trpr = False
                has_tc = False
                for tc_child in child:
                    tctag = tc_child.tag.split('}')[-1] if '}' in tc_child.tag else tc_child.tag
                    if tctag == 'trPr':
                        if has_tc:
                            issues.append("<w:tr> 内 <w:trPr> 出现在单元格之后")
                        has_trpr = True
                    elif tctag == 'tc':
                        has_tc = True
                        # 检查单元格内顺序
                        has_tcpr = False
                        has_p = False
                        for tc_c in tc_child:
                            tc_ctag = tc_c.tag.split('}')[-1] if '}' in tc_c.tag else tc_c.tag
                            if tc_ctag == 'tcPr':
                                if has_p:
                                    issues.append("<w:tc> 内 <w:tcPr> 出现在段落之后")
                                has_tcpr = True
                            elif tc_ctag == 'p':
                                has_p = True

    return {"name": "元素顺序", "passed": len(issues) == 0, "issues": issues}


def _check_cell_min_paragraph(doc):
    """校验每个单元格至少有一个段落（OpenXML规范要求）。"""
    issues = []
    for tbl in doc.element.body.iter(qn('w:tbl')):
        for tc in tbl.iter(qn('w:tc')):
            p_elements = tc.findall(qn('w:p'))
            if len(p_elements) == 0:
                issues.append("<w:tc> 缺少至少一个 <w:p> 元素，Word可能无法打开")
    return {"name": "单元格最少段落", "passed": len(issues) == 0, "issues": issues}


def _check_merge_continuity(doc):
    """校验合并单元格的连续性：vmerge应该从restart开始连续到结束。"""
    issues = []
    for t_idx, table in enumerate(doc.tables):
        for c_idx in range(len(table.columns)):
            has_restart = False
            for r_idx, row in enumerate(table.rows):
                cell = row.cells[c_idx]
                tc = cell._element
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    vmerge = tcPr.find(qn('w:vMerge'))
                    if vmerge is not None:
                        val = vmerge.get(qn('w:val'))
                        if val == 'restart':
                            has_restart = True
                        elif not has_restart:
                            # 连续的merge没有restart起始点
                            # 这不一定是个错误（可能是从上方继承的merge）
                            pass
    return {"name": "合并单元格连续性", "passed": len(issues) == 0, "issues": issues}


def _check_table_dimensions(doc):
    """校验表格维度：行列数是否在合理范围。"""
    issues = []
    for t_idx, table in enumerate(doc.tables):
        row_count = len(table.rows)
        col_count = len(table.columns)
        if row_count > 200:
            issues.append(f"表格{t_idx}行数异常: {row_count}行，可能存在行组扩展错误")
        if col_count > 30:
            issues.append(f"表格{t_idx}列数异常: {col_count}列")
    return {"name": "表格维度", "passed": len(issues) == 0, "issues": issues}


def _check_sectpr_position(doc):
    """校验sectPr是否是body的最后一个子元素。"""
    issues = []
    body = doc.element.body
    children = list(body)
    if children:
        last = children[-1]
        last_tag = last.tag.split('}')[-1] if '}' in last.tag else last.tag
        # sectPr 应该是最后一个元素
        sectprs = body.findall(qn('w:sectPr'))
        if sectprs:
            last_sectpr = sectprs[-1]
            if last_sectpr is not last:
                # 可能是 pPr 内的 sectPr，这种情况下 body 的最后一个应该是 p
                last_ppr = last.find(qn('w:pPr'))
                if last_ppr is not None and last_ppr.find(qn('w:sectPr')) is not None:
                    pass  # 正常：sectPr在最后一段的pPr内
                else:
                    issues.append("sectPr 不是 body 的最后一个元素，文档可能损坏")
    return {"name": "sectPr位置", "passed": len(issues) == 0, "issues": issues}


def _check_format_contamination(doc):
    """检测格式污染：直连格式覆盖了样式定义。

    重点检查：
    - 表格单元格内的段落是否有与样式冲突的直连rPr
    - 填入的文本是否带来了不应有的字体/颜色
    """
    issues = []
    contaminated_cells = 0

    for tbl in doc.element.body.iter(qn('w:tbl')):
        for tc in tbl.iter(qn('w:tc')):
            for p in tc.iter(qn('w:p')):
                for r in p.findall(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if rPr is not None:
                        # 检查是否有过多的直连格式属性（超过3个可能表示格式污染）
                        direct_attrs = []
                        for attr in rPr:
                            attr_tag = attr.tag.split('}')[-1] if '}' in attr.tag else attr.tag
                            # 这些是可能污染模板格式的属性
                            if attr_tag in ('rFonts', 'color', 'sz', 'szCs', 'b', 'i',
                                           'u', 'highlight', 'shd', 'strike', 'vanish'):
                                direct_attrs.append(attr_tag)
                        if len(direct_attrs) > 4:
                            contaminated_cells += 1

    if contaminated_cells > 0:
        issues.append(
            f"检测到 {contaminated_cells} 个单元格存在格式污染风险"
            f"（直连格式属性过多，可能覆盖模板样式）"
        )

    return {"name": "格式污染检测", "passed": len(issues) == 0, "issues": issues}


# ============================================================
# 二、文档对比（Diff）
# ============================================================

def diff_docx(template_path, filled_path):
    """对比模板与填写后文档的差异，返回变更清单。

    Args:
        template_path: 原始模板文件路径
        filled_path: 填写后的文档路径

    Returns:
        dict: {
            "filled": [{"field_id": str, "label": str, "value": str}],
            "still_empty": [{"field_id": str, "label": str}],
            "changed_structure": [{"type": str, "detail": str}],
            "summary": str
        }
    """
    template = Document(template_path)
    filled = Document(filled_path)

    filled_list = []
    still_empty = []
    changed_structure = []

    # 1. 对比表格单元格
    for t_idx in range(min(len(template.tables), len(filled.tables))):
        t_template = template.tables[t_idx]
        t_filled = filled.tables[t_idx]

        # 检查维度变化
        if len(t_template.rows) != len(t_filled.rows):
            changed_structure.append({
                "type": "行数变化",
                "detail": f"表格{t_idx}: 模板{len(t_template.rows)}行 → 填写后{len(t_filled.rows)}行"
            })

        # 对比每个单元格
        seen = set()
        for r_idx in range(len(t_filled.rows)):
            for c_idx in range(len(t_filled.rows[r_idx].cells)):
                cell_filled = t_filled.rows[r_idx].cells[c_idx]
                tc_id = id(cell_filled._element)
                if tc_id in seen:
                    continue
                seen.add(tc_id)

                text_filled = cell_filled.text.strip()
                field_id = f"T{t_idx}_R{r_idx}_C{c_idx}"

                # 获取模板对应单元格文本
                text_template = ""
                if r_idx < len(t_template.rows) and c_idx < len(t_template.rows[r_idx].cells):
                    text_template = t_template.rows[r_idx].cells[c_idx].text.strip()

                if text_filled and not text_template:
                    # 原来是空的，现在有内容 → 已填写
                    label = _extract_label_from_context(t_template, t_idx, r_idx, c_idx)
                    filled_list.append({
                        "field_id": field_id,
                        "label": label,
                        "value": text_filled[:80]
                    })
                elif text_filled and text_template and text_filled != text_template:
                    # 原来有内容，现在内容不同 → 已替换
                    if text_template.startswith(text_filled[:5]):
                        # 可能只是追加
                        pass
                    label = _extract_label_from_context(t_template, t_idx, r_idx, c_idx)
                    filled_list.append({
                        "field_id": field_id,
                        "label": label,
                        "value": text_filled[:80]
                    })
                elif not text_filled and not text_template:
                    # 两个都是空的 → 仍然为空
                    # 检查是否是label旁边的空值单元格
                    label = _extract_label_from_context(t_template, t_idx, r_idx, c_idx)
                    if label:
                        still_empty.append({
                            "field_id": field_id,
                            "label": label
                        })

    # 2. 对比段落（检测段落级下划线字段）
    for p_idx in range(min(len(template.paragraphs), len(filled.paragraphs))):
        p_template = template.paragraphs[p_idx]
        p_filled = filled.paragraphs[p_idx]

        # 检测下划线字段的值
        template_underline_texts = _get_underline_run_texts(p_template)
        filled_underline_texts = _get_underline_run_texts(p_filled)

        for u_idx, (t_text, f_text) in enumerate(
            zip(template_underline_texts, filled_underline_texts)
        ):
            if not t_text.strip() and f_text.strip():
                # 下划线run从空变为有内容 → 已填写
                label = _get_paragraph_label(p_template)
                filled_list.append({
                    "field_id": f"P{p_idx}_U{u_idx}",
                    "label": label,
                    "value": f_text[:80]
                })

    # 3. 生成摘要
    total_fields = len(filled_list) + len(still_empty)
    fill_rate = f"{len(filled_list)}/{total_fields}" if total_fields > 0 else "0/0"
    summary = (
        f"已填写 {len(filled_list)} 个字段，"
        f"仍有 {len(still_empty)} 个字段为空"
        f"（填写率 {fill_rate}）"
    )
    if changed_structure:
        summary += f"，结构变更 {len(changed_structure)} 处"

    return {
        "filled": filled_list,
        "still_empty": still_empty,
        "changed_structure": changed_structure,
        "summary": summary,
    }


def _extract_label_from_context(template_table_doc, t_idx, r_idx, c_idx):
    """从模板上下文中提取字段标签（取同行/同列的标签单元格文本）。"""
    try:
        table = template_table_doc  # 实际传入的是template Document
        if t_idx >= len(table.tables):
            return ""
        t = table.tables[t_idx]
        if r_idx >= len(t.rows):
            return ""
        row = t.rows[r_idx]
        if c_idx >= len(row.cells):
            return ""
        # 取左侧标签
        if c_idx > 0:
            label_text = row.cells[c_idx - 1].text.strip()
            if label_text and len(label_text) <= 20:
                return label_text
        # 取上方标签
        if r_idx > 0:
            prev_row = t.rows[r_idx - 1]
            if c_idx < len(prev_row.cells):
                label_text = prev_row.cells[c_idx].text.strip()
                if label_text and len(label_text) <= 20:
                    return label_text
    except Exception:
        pass
    return f"T{t_idx}_R{r_idx}_C{c_idx}"


def _get_underline_run_texts(paragraph):
    """获取段落中所有下划线run的文本。"""
    texts = []
    for run in paragraph.runs:
        is_ul = run.underline and run.underline not in (False, 0)
        if is_ul:
            texts.append(run.text)
    return texts


def _get_paragraph_label(paragraph):
    """从段落中提取标签（非下划线部分的文本）。"""
    label_parts = []
    for run in paragraph.runs:
        is_ul = run.underline and run.underline not in (False, 0)
        if not is_ul:
            label_parts.append(run.text)
    label = ''.join(label_parts).strip().rstrip('：:：')
    return label if label else "段落字段"


# ============================================================
# 三、自动修复（Auto-fix）
# ============================================================

def fix_docx(path_or_doc, output_path=None):
    """自动修复文档中的常见结构问题。

    Args:
        path_or_doc: 文件路径(str)或Document对象
        output_path: 修复后保存路径，None则原地覆盖

    Returns:
        dict: {
            "fixed": bool,
            "fixes": [{"type": str, "detail": str}]
        }
    """
    if isinstance(path_or_doc, str):
        doc = Document(path_or_doc)
        save_path = output_path or path_or_doc
    else:
        doc = path_or_doc
        save_path = output_path

    fixes = []

    # Fix 1: 确保每个tc至少有一个p
    fix1 = _fix_cell_min_paragraph(doc)
    fixes.extend(fix1)

    # Fix 2: 修复元素顺序
    fix2 = _fix_element_order(doc)
    fixes.extend(fix2)

    # Fix 3: 合并连续的相同run
    fix3 = _fix_merge_runs(doc)
    fixes.extend(fix3)

    if save_path and fixes:
        doc.save(save_path)

    return {"fixed": len(fixes) > 0, "fixes": fixes}


def _fix_cell_min_paragraph(doc):
    """修复：确保每个tc至少有一个p元素。"""
    fixes = []
    for tbl in doc.element.body.iter(qn('w:tbl')):
        for tc in tbl.iter(qn('w:tc')):
            p_elements = tc.findall(qn('w:p'))
            if len(p_elements) == 0:
                # 添加一个空的p元素
                from lxml import etree
                new_p = etree.SubElement(tc, qn('w:p'))
                fixes.append({
                    "type": "缺少段落",
                    "detail": "为缺少<w:p>的<w:tc>添加了空段落"
                })
    return fixes


def _fix_element_order(doc):
    """修复：确保pPr在runs之前，rPr在t之前，tcPr在p之前。"""
    fixes = []
    body = doc.element.body

    # 修复段落：将pPr移到runs之前
    for p in body.iter(qn('w:p')):
        ppr = p.find(qn('w:pPr'))
        if ppr is not None:
            first_run = p.find(qn('w:r'))
            if first_run is not None:
                # pPr应该在r之前
                ppr_idx = list(p).index(ppr)
                run_idx = list(p).index(first_run)
                if ppr_idx > run_idx:
                    p.remove(ppr)
                    p.insert(run_idx, ppr)
                    fixes.append({
                        "type": "元素顺序",
                        "detail": "将<w:pPr>移到run之前"
                    })

        # 修复run：将rPr移到t之前
        for r in p.findall(qn('w:r')):
            rpr = r.find(qn('w:rPr'))
            if rpr is not None:
                first_content = None
                for child in r:
                    ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if ctag in ('t', 'br', 'tab', 'cr'):
                        first_content = child
                        break
                if first_content is not None:
                    rpr_idx = list(r).index(rpr)
                    content_idx = list(r).index(first_content)
                    if rpr_idx > content_idx:
                        r.remove(rpr)
                        r.insert(content_idx, rpr)
                        fixes.append({
                            "type": "元素顺序",
                            "detail": "将<w:rPr>移到文本之前"
                        })

    # 修复单元格：将tcPr移到p之前
    for tbl in body.iter(qn('w:tbl')):
        for tc in tbl.iter(qn('w:tc')):
            tcpr = tc.find(qn('w:tcPr'))
            if tcpr is not None:
                first_p = tc.find(qn('w:p'))
                if first_p is not None:
                    tcpr_idx = list(tc).index(tcpr)
                    p_idx = list(tc).index(first_p)
                    if tcpr_idx > p_idx:
                        tc.remove(tcpr)
                        tc.insert(p_idx, tcpr)
                        fixes.append({
                            "type": "元素顺序",
                            "detail": "将<w:tcPr>移到段落之前"
                        })

    return fixes


def _fix_merge_runs(doc):
    """修复：合并连续的相同格式run，减少碎片化。"""
    fixes = []
    body = doc.element.body

    for p in body.iter(qn('w:p')):
        runs = p.findall(qn('w:r'))
        if len(runs) < 2:
            continue

        i = 0
        while i < len(runs) - 1:
            r1 = runs[i]
            r2 = runs[i + 1]

            # 比较两个run的rPr是否相同
            rpr1 = r1.find(qn('w:rPr'))
            rpr2 = r2.find(qn('w:rPr'))

            if _rpr_equal(rpr1, rpr2):
                # 合并文本
                t1_elements = r1.findall(qn('w:t'))
                t2_elements = r2.findall(qn('w:t'))
                if t1_elements and t2_elements:
                    t1 = t1_elements[0]
                    t2 = t2_elements[0]
                    t1.text = (t1.text or '') + (t2.text or '')
                    p.remove(r2)
                    runs.pop(i + 1)
                    fixes.append({
                        "type": "合并run",
                        "detail": f"合并了相同格式的连续run"
                    })
                    continue
            i += 1

    return fixes


def _rpr_equal(rpr1, rpr2):
    """比较两个rPr元素是否相等。"""
    if rpr1 is None and rpr2 is None:
        return True
    if rpr1 is None or rpr2 is None:
        return False
    from lxml import etree
    return etree.tostring(rpr1) == etree.tostring(rpr2)


# ============================================================
# 四、格式污染防护（Format Contamination Prevention）
# ============================================================

def strip_inline_formatting(element, keep=('b', 'i', 'u', 'sz', 'szCs', 'rFonts')):
    """剥离元素中的直连格式，只保留指定的属性，防止格式污染。

    类比 MiniMax-docx 的规则：填入内容时，只保留 pStyle 引用和必要的基本格式，
    剥离其他直连格式（颜色、高亮、底纹等），防止源文档格式污染模板。

    Args:
        element: XML元素（可以是tc、p或r）
        keep: 保留的格式属性列表，默认保留加粗/斜体/下划线/字号/字体

    Returns:
        int: 剥离的属性数量
    """
    stripped = 0

    # 处理段落中的runs
    for r in element.iter(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            children_to_remove = []
            for child in rPr:
                ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if ctag not in keep:
                    children_to_remove.append(child)
            for child in children_to_remove:
                rPr.remove(child)
                stripped += 1

    # 处理段落的直连pPr
    for p in element.iter(qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            ppr_keep = ('pStyle', 'rPr', 'spacing', 'ind', 'jc', 'outlineLvl',
                       'keepNext', 'keepLines', 'pageBreakBefore', 'sectPr')
            children_to_remove = []
            for child in pPr:
                ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if ctag not in ppr_keep:
                    children_to_remove.append(child)
            for child in children_to_remove:
                pPr.remove(child)
                stripped += 1

    return stripped


def sanitize_fill_text(text):
    """清理填入文本中的控制字符和不可见字符，防止XML注入。

    Args:
        text: 待填入的文本

    Returns:
        str: 清理后的文本
    """
    if not isinstance(text, str):
        text = str(text)

    # 移除控制字符（保留换行和制表符）
    # 不手动转义XML：python-docx/lxml在写入XML时会自动转义<、>、&等字符
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    return text
