"""知识文件解析工具 - 通过LLM智能提取文档所需信息"""

import os
import re
import json
from langchain.tools import tool
from langchain_core.messages import SystemMessage, HumanMessage
from coze_coding_dev_sdk import LLMClient
from coze_coding_utils.log.write_log import request_context
from coze_coding_utils.runtime_ctx.context import new_context


def _extract_text_from_file(file_path: str) -> str:
    """根据文件扩展名提取文本内容"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    if ext == '.csv':
        import csv
        lines = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                lines.append(', '.join(row))
        return '\n'.join(lines)

    if ext in ('.docx', '.doc'):
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    # 用unique cells避免合并单元格重复
                    seen = set()
                    cells = []
                    for cell in row.cells:
                        cid = id(cell._element)
                        if cid not in seen:
                            seen.add(cid)
                            cells.append(cell.text.strip())
                    non_empty = [c for c in cells if c]
                    if non_empty:
                        text_parts.append(' | '.join(non_empty))
            return '\n'.join(text_parts)
        except Exception as e:
            return f"[docx解析失败: {str(e)}]"

    if ext in ('.xlsx', '.xls'):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        lines.append(', '.join(cells))
            wb.close()
            return '\n'.join(lines)
        except Exception as e:
            return f"[xlsx解析失败: {str(e)}]"

    if ext == '.pdf':
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
            return '\n'.join(text_parts)
        except Exception as e:
            return f"[PDF解析失败: {str(e)}]"

    return f"[不支持的文件格式: {ext}]"


def extract_text_from_upload(file_path: str) -> dict:
    """
    解析上传文件并提取文本内容（供API调用，非@tool）

    Returns:
        dict: {success: bool, extracted_text: str, error: str}
    """
    try:
        text = _extract_text_from_file(file_path)
        if text.startswith('['):
            return {"success": False, "extracted_text": "", "error": text}
        return {"success": True, "extracted_text": text, "error": ""}
    except Exception as e:
        return {"success": False, "extracted_text": "", "error": str(e)}


def _llm_extract_fields(field_list: list, file_content: str, ctx=None) -> dict:
    """使用LLM从文件内容中智能提取字段值。

    Args:
        field_list: 需要提取的字段名列表，如 ["课程名称", "学时数", "教师姓名"]
        file_content: 知识文件的文本内容
        ctx: 请求上下文

    Returns:
        dict: {字段名: 提取到的值}，未提取到的字段不在结果中
    """
    # 限制文件内容长度（避免token过多）
    max_chars = 8000
    if len(file_content) > max_chars:
        file_content = file_content[:max_chars] + "\n...(内容过长已截断)"

    fields_str = "、".join(field_list)

    system_prompt = """你是一个精确的信息提取助手。你的任务是从给定的文件内容中，提取指定字段的值。

规则：
1. 只提取明确出现在文件内容中的信息，不要编造
2. 对于每个字段，找到最匹配的值。字段名可能和文件中的表述不完全一致，你需要理解语义匹配
   - 例如：字段"课程名称"可能对应文件中的"课程名"、"课程"、"科目名称"等
   - 例如：字段"学时数"可能对应文件中的"学时"、"总学时"、"课时"等
   - 例如：字段"教师姓名"可能对应文件中的"任课教师"、"授课教师"、"主讲教师"等
3. 提取的值应该简洁，去掉多余的前缀和后缀
4. 如果文件中找不到某个字段的值，就不要在结果中包含该字段
5. 必须返回合法的JSON格式"""

    user_prompt = f"""请从以下文件内容中提取这些字段的值：{fields_str}

文件内容：
{file_content}

请返回JSON格式，例如：
{{"课程名称": "高等数学", "学时数": "64", "教师姓名": "张明"}}

只返回JSON，不要其他文字。"""

    try:
        # 支持外部模型API（如 DeepSeek）
        ext_api_key = os.getenv("EXTERNAL_LLM_API_KEY")
        ext_base_url = os.getenv("EXTERNAL_LLM_BASE_URL")

        if ext_api_key and ext_base_url:
            # 使用外部API
            from langchain_openai import ChatOpenAI
            ext_model = os.getenv("EXTERNAL_LLM_MODEL", "deepseek-chat")
            ext_llm = ChatOpenAI(
                model=ext_model,
                api_key=ext_api_key,
                base_url=ext_base_url,
                temperature=0.1,
                max_tokens=4096,
            )
            from langchain_core.messages import SystemMessage as SM, HumanMessage as HM
            response = ext_llm.invoke([SM(content=system_prompt), HM(content=user_prompt)])
            content = response.content
        else:
            # 使用平台内置LLM
            client = LLMClient(ctx=ctx)
            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=user_prompt),
            ]
            response = client.invoke(
                messages=messages,
                model="doubao-seed-1-6-lite-251015",
                temperature=0.1,
                max_completion_tokens=4096,
            )
            content = response.content
            if isinstance(content, list):
                content = " ".join(
                    item.get("text", "") for item in content
                    if isinstance(item, dict) and item.get("type") == "text"
                )

        # 提取JSON部分
        content = content.strip()
        # 尝试直接解析
        try:
            result = json.loads(content)
        except json.JSONDecodeError:
            # 尝试提取JSON代码块
            json_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', content)
            if json_match:
                result = json.loads(json_match.group(1).strip())
            else:
                # 尝试提取花括号内容
                brace_match = re.search(r'\{[\s\S]*\}', content)
                if brace_match:
                    result = json.loads(brace_match.group(0))
                else:
                    return {}

        # 过滤：只保留field_list中的字段，值必须是字符串且非空
        extracted = {}
        for field in field_list:
            if field in result and result[field]:
                val = str(result[field]).strip()
                if val and val not in ("null", "None", "无", "未知", "-"):
                    extracted[field] = val

        return extracted

    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"LLM提取失败: {e}")
        return {}


def _rule_extract_fields(field_list: list, file_content: str) -> dict:
    """使用规则匹配从文件内容中提取字段值（兜底方案）。

    支持多种格式：
    - 冒号模式: 课程性质：专业必修课
    - 等号模式: 课程性质＝专业必修课
    - 管道符/表格模式: 课程性质 | 专业必修课
    - "是"模式: 课程性质是专业必修课
    - 【】模式: 【课程性质】专业必修课
    """
    extracted = {}

    # 字段名别名映射（同义词扩展）
    ALIAS_MAP = {
        "课程名称": ["课程名", "课程", "科目名称", "科目", "课程代码名称"],
        "学时数": ["学时", "总学时", "课时", "总课时", "教学时数"],
        "教师姓名": ["任课教师", "授课教师", "主讲教师", "教师", "任课老师"],
        "开课单位": ["开课院系", "教学单位", "院系", "所在院系", "开课学院"],
        "课程性质": ["课程类别", "课程类型", "修读性质"],
        "学生班级": ["教学班", "班级", "授课班级", "教学班级"],
        "开课时间": ["学期", "开课学期", "授课学期"],
        "考试类别": ["考核方式", "考试方式"],
        "考试形式": ["考试类型"],
        "平均分": ["平均成绩", "均分"],
        "最高分": ["最高成绩"],
        "最低分": ["最低成绩"],
        "标准差": ["标准偏差", "成绩标准差"],
    }

    for field in field_list:
        # 构建搜索词列表：原名 + 别名
        search_terms = [field]
        if field in ALIAS_MAP:
            search_terms.extend(ALIAS_MAP[field])

        found_values = []

        for term in search_terms:
            for line in file_content.split('\n'):
                line = line.strip()
                if not line:
                    continue

                # ── 冒号/等号分隔 ──
                for sep in ['：', ':', '＝', '=']:
                    pat = re.compile(
                        r'(?:^|\|)\s*' + re.escape(term) + r'\s*' + re.escape(sep) + r'\s*(.+?)(?:\s*\||$)'
                    )
                    m = pat.search(line)
                    if m:
                        val = m.group(1).strip().rstrip('，。,.;；')
                        if val and len(val) < 100 and val != term:
                            found_values.append(val)
                            break
                else:
                    continue
                break

                if found_values:
                    break

            if found_values:
                break

        if found_values:
            extracted[field] = found_values[0]
            continue

        # ── 管道符/表格格式 ──
        for term in search_terms:
            done = False
            for line in file_content.split('\n'):
                line = line.strip()
                if '|' not in line:
                    continue
                parts = [p.strip() for p in line.split('|') if p.strip()]
                for i, part in enumerate(parts):
                    if part == term and i + 1 < len(parts):
                        val = parts[i + 1].rstrip('，。,.;；')
                        if val and len(val) < 100 and val != term:
                            extracted[field] = val
                            done = True
                            break
                if done:
                    break
            if done:
                break

        if field in extracted:
            continue

        # ── "是"连接 ──
        for term in search_terms:
            for line in file_content.split('\n'):
                m = re.search(re.escape(term) + r'是(.+?)(?:[,，。.;；\|]|$)', line)
                if m:
                    val = m.group(1).strip()
                    if val and len(val) < 100:
                        extracted[field] = val
                        break
            if field in extracted:
                break

        if field in extracted:
            continue

        # ── 【字段】值 ──
        for term in search_terms:
            for line in file_content.split('\n'):
                m = re.search(r'【' + re.escape(term) + r'】\s*(.+?)(?:[,，。.;；\|]|$)', line)
                if m:
                    val = m.group(1).strip()
                    if val and len(val) < 100:
                        extracted[field] = val
                        break
            if field in extracted:
                break

    return extracted


@tool
def parse_knowledge_file(file_description: str, file_content: str, missing_fields: str) -> str:
    """解析用户上传的知识文件内容，智能提取文档所需的缺失信息。

    优先使用LLM进行语义理解提取，规则匹配作为兜底。
    当Agent发现用户缺少某些字段信息，且用户上传了文件时，调用此工具从文件内容中提取缺失字段。

    Args:
        file_description: 文件描述（如文件名、类型等）
        file_content: 文件的文本内容
        missing_fields: 需要提取的缺失字段列表，用逗号分隔

    Returns:
        提取结果JSON字符串，包含每个字段的提取值
    """
    ctx = request_context.get() or new_context(method="parse_knowledge_file")

    fields = [f.strip() for f in missing_fields.split(',') if f.strip()]
    content = file_content[:10000]  # 扩大到10000字符

    result = {
        "file": file_description,
        "missing_fields": fields,
        "extracted": {},
        "summary": f"已从文件中智能提取以下字段：{', '.join(fields)}"
    }

    # 优先：LLM智能提取
    llm_extracted = _llm_extract_fields(fields, content, ctx=ctx)

    # 兜底：规则匹配提取
    rule_extracted = _rule_extract_fields(fields, content)

    # 合并结果：LLM优先，规则补充
    for field in fields:
        if field in llm_extracted:
            result["extracted"][field] = llm_extracted[field]
        elif field in rule_extracted:
            result["extracted"][field] = rule_extracted[field]

    # 生成提示
    lines = []
    found_fields = []
    for field in fields:
        if field in result["extracted"]:
            lines.append(f"✅ {field}：{result['extracted'][field]}")
            found_fields.append(field)

    unfound = [f for f in fields if f not in result["extracted"]]
    if unfound:
        lines.append(f"⚠ 未找到：{', '.join(unfound)}")

    result["choice_prompt"] = '\n'.join(lines)
    result["found_count"] = len(found_fields)
    result["total_count"] = len(fields)

    return json.dumps(result, ensure_ascii=False)
