"""
教务报告生成工具 - 核心逻辑
"""
import os
import json
import copy
import re
import tempfile
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from langchain.tools import tool

from coze_coding_utils.log.write_log import request_context
from coze_coding_utils.runtime_ctx.context import new_context
from storage.memory.memory_saver import get_memory_saver

from tools.template_analyzer import analyze_template


# ── 模板注册表 ──
TEMPLATE_REGISTRY = {
    "评价报告": "assets/2023-2024-2《xxx》 岭南师范学院专业课程目标达成度评价报告模板.docx",
    "试卷分析": "assets/2023-2024-2《xxx》 试卷分析模板.docx",
    "关联矩阵": "assets/2023-2024-2《xxx》岭南师范学院考题与课程目标及毕业要求关联矩阵表模板.docx",
}

# 复杂行组最大列数阈值
_MAX_SIMPLE_GROUP_COLS = 15


def _get_template_path(name: str) -> str:
    workspace = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")
    file = TEMPLATE_REGISTRY.get(name)
    if not file:
        avail = ", ".join(TEMPLATE_REGISTRY.keys())
        raise ValueError(f"未找到模板'{name}'，可用模板: [{avail}]")
    return os.path.join(workspace, file)


# ── 单元格操作 ──

def _get_unique_cells(row):
    """获取行中的独立单元格（通过XML元素去重）。"""
    seen = set()
    unique = []
    for cell in row.cells:
        cid = id(cell._element)
        if cid not in seen:
            seen.add(cid)
            unique.append(cell)
    return unique


def _is_vmerge_continue(cell) -> bool:
    """判断单元格是否为 vMerge=continue（延续格）。"""
    tc = cell._element
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        vm = tcPr.find(qn("w:vMerge"))
        if vm is not None:
            val = vm.get(qn("w:val"))
            if val is None or val == "":
                return True
    return False


def _is_vmerge_restart(cell) -> bool:
    """判断单元格是否为 vMerge=restart。"""
    tc = cell._element
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        vm = tcPr.find(qn("w:vMerge"))
        if vm is not None:
            val = vm.get(qn("w:val"))
            if val == "restart":
                return True
    return False


def _clone_rpr(rPr_source):
    """复制 run 格式。"""
    if rPr_source is None:
        return None
    if getattr(rPr_source, "tag", None) == qn("w:rPr"):
        return copy.deepcopy(rPr_source)
    src_rPr = rPr_source.find(qn("w:rPr"))
    if src_rPr is not None:
        return copy.deepcopy(src_rPr)
    return None


def _get_first_run_rpr(paragraph):
    """获取段落中第一个 run 的 rPr。"""
    for run in paragraph.findall(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is not None:
            return rPr
    return None


def _ensure_run_text(run, text: str, rPr_source=None):
    """尽量复用已有 run，只改文本，不破坏段落结构。"""
    if run.find(qn("w:rPr")) is None:
        cloned = _clone_rpr(rPr_source)
        if cloned is not None:
            run.insert(0, cloned)

    text_nodes = run.findall(qn("w:t"))
    if not text_nodes:
        t_elem = run.makeelement(qn("w:t"), {})
        run.append(t_elem)
        text_nodes = [t_elem]

    text_nodes[0].text = text
    text_nodes[0].set(qn("xml:space"), "preserve")
    for extra_t in text_nodes[1:]:
        extra_t.text = ""


def _set_paragraph_text_preserve_runs(paragraph, text: str, rPr_source=None):
    """尽量保留现有 paragraph/run 结构，只替换文本。"""
    runs = paragraph.findall(qn("w:r"))
    target_run = None
    for run in runs:
        if run.findall(qn("w:t")):
            target_run = run
            break

    if target_run is None:
        if runs:
            target_run = runs[0]
        else:
            target_run = paragraph.makeelement(qn("w:r"), {})
            paragraph.append(target_run)

    effective_rpr = _get_first_run_rpr(paragraph) or rPr_source
    _ensure_run_text(target_run, text, effective_rpr)

    for run in runs:
        if run is target_run:
            continue
        for t_elem in run.findall(qn("w:t")):
            t_elem.text = ""


def _append_paragraph_with_text(tc, text: str, rPr_source=None, pPr_source=None):
    """在单元格末尾追加一个带文本的段落，并继承段落/字体格式。"""
    paragraph = tc.makeelement(qn("w:p"), {})
    if pPr_source is not None:
        paragraph.append(copy.deepcopy(pPr_source))
    run = paragraph.makeelement(qn("w:r"), {})
    cloned = _clone_rpr(rPr_source)
    if cloned is not None:
        run.append(cloned)
    t_elem = run.makeelement(qn("w:t"), {})
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    run.append(t_elem)
    paragraph.append(run)
    tc.append(paragraph)
    return paragraph


def _set_tc_text(tc, text: str, rPr_source=None):
    """替换 tc 元素中的文本，尽量保留原模板的段落/run结构。"""
    paragraphs = tc.findall(qn("w:p"))
    if not paragraphs:
        _append_paragraph_with_text(tc, str(text), rPr_source=rPr_source)
        return

    lines = str(text).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if not lines:
        lines = [""]

    first_ppr = paragraphs[0].find(qn("w:pPr"))
    for idx, line in enumerate(lines):
        if idx < len(paragraphs):
            effective_rpr = _get_first_run_rpr(paragraphs[idx]) or rPr_source
            _set_paragraph_text_preserve_runs(paragraphs[idx], line, effective_rpr)
        else:
            _append_paragraph_with_text(tc, line, rPr_source=rPr_source, pPr_source=first_ppr)

    for paragraph in paragraphs[len(lines):]:
        effective_rpr = _get_first_run_rpr(paragraph) or rPr_source
        _set_paragraph_text_preserve_runs(paragraph, "", effective_rpr)


def _set_cell_text(cell, text: str, rPr_source=None):
    """设置单元格文本（清空后写入），保留格式。"""
    _set_tc_text(cell._element, text, rPr_source=rPr_source)


def _replace_signature_and_date_in_tc(tc, label: str, existing_value: str, value: str):
    """在tc元素中，处理"负责人签名：年月日"格式——在标签后追加签名，可选替换日期占位符。
    
    场景1：值只有签名（如"李明"）→ 在签名后追加名字，日期占位符保持不变
    场景2：值包含签名+日期（如"李明 2024年12月15日"）→ 同时替换日期占位符
    """
    import re as _re
    _EMBEDDED_DATE_PAT = _re.compile(r'年\s*月\s*日(?:\s*(?:上午|下午|上|下))?')
    _DATE_VALUE_PAT = _re.compile(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日')
    
    # 判断value是否包含日期
    date_match = _DATE_VALUE_PAT.search(value)
    if date_match:
        # 值包含日期，同时替换签名和日期
        _replace_embedded_date_in_tc(tc, label, existing_value, value)
    else:
        # 值只是签名，在标签后追加签名，保留日期占位符
        _append_value_to_tc_after_label(tc, label, value)


def _replace_embedded_date_in_tc(tc, label: str, existing_value: str, value: str):
    """在tc元素中，替换嵌入式日期占位符为实际值。
    
    例如：单元格文本 "考试时间  年  月  日  上午  下午"
    label="考试时间", existing_value="年  月  日  上午  下午", value="2024年12月15日 上午"
    结果: "考试时间  2024年12月15日 上午"
    """
    import re as _re
    _EMBEDDED_DATE_PAT = _re.compile(r'年\s*月\s*日(?:\s*(?:上午|下午|上|下))?')
    
    for paragraph in tc.findall(qn("w:p")):
        text_nodes = paragraph.findall(".//" + qn("w:t"))
        para_text = "".join((t.text or "") for t in text_nodes)
        
        # 在段落文本中找到标签位置，然后替换标签后面的日期占位符
        label_pos = para_text.find(label)
        if label_pos >= 0:
            after_label = para_text[label_pos + len(label):]
            replaced = _EMBEDDED_DATE_PAT.sub(value, after_label, count=1)
            new_para_text = para_text[:label_pos + len(label)] + replaced
            
            effective_rpr = _get_first_run_rpr(paragraph)
            _set_paragraph_text_preserve_runs(paragraph, new_para_text, effective_rpr)
            return True
    
    # 回退：直接设置整个单元格文本
    rPr_source = _find_label_rPr_in_row(tc.getparent())
    _set_tc_text(tc, f"{label} {value}", rPr_source=rPr_source)
    return True


def _append_value_to_tc_after_label(tc, label: str, value: str):
    """在tc元素中，在标签run后追加一个新的value run，保留标签run不变。

    改进点：不再把标签和值合并到同一个run，而是在标签run之后
    插入一个新的run来放置值，这样标签和值可以保持各自的格式。
    """
    for paragraph in tc.findall(qn("w:p")):
        text_nodes = paragraph.findall(".//" + qn("w:t"))
        para_text = "".join((t.text or "") for t in text_nodes)
        if label in para_text:
            # 找到包含标签文字的run
            runs = paragraph.findall(qn("w:r"))
            for run in runs:
                run_text = "".join((t.text or "") for t in run.findall(qn("w:t")))
                if label in run_text:
                    # 在此run后面插入新run
                    new_run = copy.deepcopy(run)
                    # 新run只包含value
                    for t_elem in new_run.findall(qn("w:t")):
                        t_elem.text = value
                        t_elem.set(qn("xml:space"), "preserve")
                    # 如果label后面还有文字（如冒号后的空格），保留
                    # 只清空新run中不需要的部分
                    run.addnext(new_run)
                    return True

            # 如果没有找到包含完整label的run（跨run标签），回退到简单方式
            effective_rpr = _get_first_run_rpr(paragraph)
            replaced = para_text.replace(label, f"{label}{value}", 1)
            _set_paragraph_text_preserve_runs(paragraph, replaced, effective_rpr)
            return True

    paragraphs = tc.findall(qn("w:p"))
    if paragraphs:
        paragraph = paragraphs[0]
        para_text = "".join((t.text or "") for t in paragraph.findall(".//" + qn("w:t")))
        effective_rpr = _get_first_run_rpr(paragraph)
        _set_paragraph_text_preserve_runs(paragraph, para_text + str(value), effective_rpr)
        return True
    return False


def _build_field_id_value_map(label_fields, data):
    """把已填数据映射为精确的 field_id -> value，供前端精确回填。"""
    field_values = {}
    for field in label_fields:
        label = field["label"]
        if label not in data:
            continue
        value = data[label]
        if value is None:
            continue
        value_str = str(value).strip()
        if not value_str:
            continue
        field_values[field["field_id"]] = value_str
    return field_values


def validate_doc(template_path: str, output_path: str) -> dict:
    """对比模板和生成文档的结构完整性。

    检查项:
    1. 生成文件能否被 python-docx 正常打开
    2. 表格数量是否一致
    3. 每张表的行数是否一致
    4. 每张表的列数是否一致

    Returns:
        {"valid": bool, "errors": [...], "warnings": [...]}
        errors → 阻断性错误（不应返回下载链接）
        warnings → 警告（可返回下载链接但需提示）
    """
    errors = []
    warnings = []

    # 1. 检查生成文件能否打开
    try:
        output_doc = Document(output_path)
    except Exception as e:
        return {"valid": False, "errors": [f"生成的文档无法打开: {e}"], "warnings": []}

    try:
        template_doc = Document(template_path)
    except Exception as e:
        warnings.append(f"模板文件无法打开，跳过结构对比: {e}")
        return {"valid": True, "errors": [], "warnings": warnings}

    # 2. 表格数量
    t_tables = len(template_doc.tables)
    o_tables = len(output_doc.tables)
    if t_tables != o_tables:
        errors.append(f"表格数量变化: 模板{t_tables}个, 生成{o_tables}个")

    # 3 & 4. 每张表的行数和列数
    for i in range(min(t_tables, o_tables)):
        t_rows = len(template_doc.tables[i].rows)
        o_rows = len(output_doc.tables[i].rows)
        if t_rows != o_rows:
            # 行数增加可能是行组填充（合法），减少则是有问题
            if o_rows < t_rows:
                errors.append(f"表格{i+1}行数减少: 模板{t_rows}行, 生成{o_rows}行")
            else:
                warnings.append(f"表格{i+1}行数增加: 模板{t_rows}行, 生成{o_rows}行（可能是行组填充）")

        t_cols = len(template_doc.tables[i].columns)
        o_cols = len(output_doc.tables[i].columns)
        if t_cols != o_cols:
            errors.append(f"表格{i+1}列数变化: 模板{t_cols}列, 生成{o_cols}列")

    return {"valid": len(errors) == 0, "errors": errors, "warnings": warnings}


def _add_table_row_after(table, after_row_idx):
    """在指定行后复制一行（用于行组填充）。"""
    src_row = table.rows[after_row_idx]
    new_tr = copy.deepcopy(src_row._tr)
    after_row_idx + 1  # noqa
    src_row._tr.addnext(new_tr)
    return new_tr


# ── 选项词集合 ──
_OPTION_WORDS_SET = frozenset({
    "选修", "必修", "开卷", "闭卷", "半开卷", "是", "否",
    "试题库", "试卷库", "教师组题", "本人阅卷", "同行阅卷",
    "集体阅卷", "机器阅卷", "其他",
})


# ── 勾选框行检测和填充 ──

def _normalize_option_text(text: str) -> str:
    """标准化选项文字：去除多余空格，使'闭  卷'匹配'闭卷'。"""
    return text.replace(" ", "").replace("　", "")


def _find_section_title_above(table, row_idx: int) -> str:
    """从上方行查找章节标题（如'二、考试方式'），用于无行标签的勾选框组。"""
    for r in range(row_idx - 1, max(row_idx - 3, -1), -1):
        row = table.rows[r]
        seen = set()
        for cell in row.cells:
            elem_id = id(cell._element)
            if elem_id not in seen:
                seen.add(elem_id)
                text = cell.text.strip()
                # 检查是否是跨整行的章节标题（gridSpan较大）
                from docx.oxml.ns import qn as _qn
                tc = cell._tc
                tcPr = tc.find(_qn("w:tcPr"))
                if tcPr is not None:
                    gs = tcPr.find(_qn("w:gridSpan"))
                    if gs is not None:
                        span = int(gs.get(_qn("w:val"), "1"))
                        if span >= 5 and text:  # 跨5列以上视为章节标题
                            return text
    return ""


def _detect_checkbox_row(unique_cells):
    """检测一行是否是勾选框行，返回标签组列表。
    
    勾选框行模式: [行标签] [选项1] [空白1] [选项2] [空白2] ... [行标签2] [选项3] [空白3] ...
    也支持无行标签但有上方章节标题的模式。
    返回: [{"label": 行标签, "option_blanks": {选项文字: 空白格索引}}, ...]
    """
    groups = []
    current_label = None
    current_options = {}
    
    for ci in range(len(unique_cells) - 1):
        cell_text = unique_cells[ci].text.strip()
        next_text = unique_cells[ci + 1].text.strip()
        
        # 标准化选项文字以匹配（处理"闭  卷" → "闭卷"的情况）
        normalized_text = _normalize_option_text(cell_text)
        
        if (normalized_text in _OPTION_WORDS_SET or cell_text in _OPTION_WORDS_SET) and not next_text:
            # 找到选项+空白格对
            # 优先使用原始文本（保留空格版）作为key，便于后续匹配
            opt_key = normalized_text if normalized_text in _OPTION_WORDS_SET else cell_text
            current_options[opt_key] = ci + 1
        elif cell_text and normalized_text not in _OPTION_WORDS_SET and cell_text not in _OPTION_WORDS_SET and not _is_vmerge_continue(unique_cells[ci]):
            # 遇到非选项标签文字
            # 如果当前组已有选项，先保存
            if current_options and current_label:
                groups.append({"label": current_label, "option_blanks": current_options})
            elif current_options:
                # 无行标签但有选项——使用空标签或从上方标题获取
                groups.append({"label": current_label or "_checkbox_group", "option_blanks": current_options})
            current_label = cell_text
            current_options = {}
    
    # 保存最后一组
    if current_options:
        if current_label:
            groups.append({"label": current_label, "option_blanks": current_options})
        else:
            groups.append({"label": "_checkbox_group", "option_blanks": current_options})
    
    return groups


def _fill_checkbox_row(unique_cells, group, user_value):
    """根据用户值在勾选框组打√。
    
    user_value: 如 "必修" 或 "闭卷" （单选）
    """
    selected = [v.strip() for v in user_value.replace("，", ",").split(",")]
    option_blanks = group["option_blanks"]
    
    # 找到同行中第一个有格式的cell作为格式源
    rPr_source = None
    for cell in unique_cells:
        tc = cell._element
        for p in tc.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is not None:
                    rPr_source = rPr
                    break
            if rPr_source is not None:
                break
        if rPr_source is not None:
            break
    
    for opt_text, blank_idx in option_blanks.items():
        if opt_text in selected:
            _set_cell_text(unique_cells[blank_idx], "√", rPr_source=rPr_source)
        # 未选中的空白格保持空白


# ── 标签字段填充 ──

def _find_label_rPr_in_row(tr):
    """从一行中找到第一个有rPr格式的标签格，返回其rPr元素（深拷贝）。
    
    用于空白格填充时继承同行标签格的字体/字号等格式。
    """
    tcs = tr.findall(qn("w:tc"))
    for tc in tcs:
        for p in tc.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is not None:
                    return rPr
    return None


def _fill_label_fields(doc, label_fields, data):
    """填充标签字段，支持多种填充模式。保留原有格式。"""
    import re as _re
    _EMBEDDED_DATE_PAT = _re.compile(r'年\s*月\s*日(?:\s*(?:上午|下午|上|下))?')
    
    for f in label_fields:
        label = f["label"]
        if label not in data:
            continue
        # 跳过段落级字段（由_fill_paragraph_fields单独处理）
        if f.get("pattern") == "paragraph_underline":
            continue
        value = data[label]
        fill_mode = f.get("fill_mode", "append")
        pattern = f.get("pattern", "colon")
        existing_value = f.get("existing_value", "")
        
        for ri in f["row_indices"]:
            t_idx = f["table_idx"]
            table = doc.tables[t_idx]
            
            if fill_mode == "set":
                # 标签格+空白格模式：直接设置空白格内容
                col_idx = f["col_idx"]
                # 用tr级别获取tc
                tr = table.rows[ri]._tr
                tcs = tr.findall(qn("w:tc"))
                if col_idx < len(tcs):
                    # 获取同行标签格的格式作为格式源
                    rPr_source = _find_label_rPr_in_row(tr)
                    _set_tc_text(tcs[col_idx], str(value), rPr_source=rPr_source)
            
            elif fill_mode == "replace":
                # 占位符替换模式
                col_idx = f["col_idx"]
                tr = table.rows[ri]._tr
                tcs = tr.findall(qn("w:tc"))
                if col_idx < len(tcs):
                    tc = tcs[col_idx]
                    raw_label = f.get("raw_label", label)
                    existing_value = f.get("existing_value", "")
                    # colon模式 + 日期占位符：在标签后追加签名并替换日期占位符
                    if pattern == "colon" and existing_value and _EMBEDDED_DATE_PAT.search(existing_value):
                        _replace_signature_and_date_in_tc(tc, raw_label, existing_value, str(value))
                    elif pattern == "colon" and not existing_value:
                        # colon模式空值：在标签后追加
                        _append_value_to_tc_after_label(tc, raw_label, str(value))
                    else:
                        # 非colon模式：清空格内容后写入新值
                        rPr_source = _find_label_rPr_in_row(tr)
                        _set_tc_text(tc, str(value), rPr_source=rPr_source)
            
            elif fill_mode == "append":
                col_idx = f["col_idx"]
                tr = table.rows[ri]._tr
                tcs = tr.findall(qn("w:tc"))
                if col_idx < len(tcs):
                    tc = tcs[col_idx]
                    # 特殊处理：嵌入式日期占位符
                    # 如果existing_value包含"年月日"等日期占位符，替换占位符而非追加
                    if existing_value and _EMBEDDED_DATE_PAT.search(existing_value):
                        _replace_embedded_date_in_tc(tc, label, existing_value, str(value))
                    else:
                        # 标准冒号模式：在标签后追加值
                        _append_value_to_tc_after_label(tc, label, str(value))


def _fill_paragraph_fields(doc, paragraph_fields, data):
    """填充正文段落中的下划线横线字段。
    
    格式：段落中有"标签(无下划线) + 空白(有下划线)"的run结构，
    需要将值填入下划线run中，保留下划线格式。
    """
    for f in paragraph_fields:
        label = f["label"]
        if label not in data:
            continue
        value = str(data[label])
        
        p_idx = f["row_idx"]
        if p_idx >= len(doc.paragraphs):
            continue
        
        para = doc.paragraphs[p_idx]
        
        # 找到下划线run并填入值
        for run in para.runs:
            is_underline = False
            if run.underline and run.underline not in (False, 0):
                is_underline = True
            if not is_underline:
                rPr = run._element.find(qn('w:rPr'))
                if rPr is not None:
                    u_elem = rPr.find(qn('w:u'))
                    if u_elem is not None:
                        val = u_elem.get(qn('w:val'))
                        if val and val not in ('none',):
                            is_underline = True
            
            if is_underline and not run.text.strip():
                # 将下划线空白run替换为填入的值
                run.text = f" {value} "
                break


def _fill_checkbox_rows_in_table(doc, analysis, data):
    """扫描所有表格行，检测勾选框行并填充。"""
    for t_idx, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            unique = _get_unique_cells(row)
            groups = _detect_checkbox_row(unique)
            
            if not groups:
                continue
            
            # 对每个标签组，查找data中对应的值并填充
            for group in groups:
                row_label = group["label"]
                user_value = None
                
                # 1. 精确匹配
                if row_label in data:
                    user_value = data[row_label]
                else:
                    # 2. 模糊匹配
                    clean_label = row_label.replace("\n", "").replace(" ", "")
                    for key in data:
                        clean_key = key.replace("\n", "").replace(" ", "")
                        if clean_key == clean_label:
                            user_value = data[key]
                            break
                
                # 3. 对_checkbox_group标签（无行标签），尝试从上方章节标题匹配
                if user_value is None and row_label == "_checkbox_group":
                    section_title = _find_section_title_above(table, ri)
                    if section_title:
                        clean_section = section_title.replace("\n", "").replace(" ", "")
                        for key in data:
                            clean_key = key.replace("\n", "").replace(" ", "")
                            if clean_key == clean_section or clean_key in clean_section:
                                user_value = data[key]
                                break
                
                # 4. 检查选项本身是否在data中
                if user_value is None:
                    for opt in group["option_blanks"]:
                        if opt in data:
                            user_value = data[opt]
                            break
                
                # 5. 对_checkbox_group，也检查data中是否有该组任意选项的同义词
                if user_value is None and row_label == "_checkbox_group":
                    for opt in group["option_blanks"]:
                        for key in data:
                            if _normalize_option_text(key) == _normalize_option_text(opt):
                                user_value = data[key]
                                break
                        if user_value is not None:
                            break
                
                if user_value is not None:
                    _fill_checkbox_row(unique, group, user_value)


def _fill_simple_row_groups(doc, groups, data):
    """填充简单行组（重复行数据）。
    
    Returns:
        dict: field_id → value 映射，格式为 T{t}_G{g}_R{r}_C{c}
    """
    rg_field_values = {}
    for g in groups:
        gid = g["group_id"]
        if gid not in data:
            continue
        
        row_data_list = data[gid]
        if not isinstance(row_data_list, list):
            continue
        
        t_idx = g["table_idx"]
        table = doc.tables[t_idx]
        start_row = g["start_row"]
        template_row_count = g["template_row_count"]
        
        # 从 group_id 提取 group 序号 (如 "T0_G1" → 1)
        g_num = int(gid.split("_G")[1]) if "_G" in gid else 0
        
        # 填充模板行
        for row_offset, row_data in enumerate(row_data_list):
            if row_offset >= template_row_count:
                break
            actual_row = start_row + row_offset
            
            # 使用tr级别的tc元素
            tr = table.rows[actual_row]._tr
            tcs = tr.findall(qn("w:tc"))
            
            data_idx = 0
            for ti, tc in enumerate(tcs):
                if data_idx >= len(row_data):
                    break
                
                # 检查vMerge
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is not None:
                    vm = tcPr.find(qn("w:vMerge"))
                    if vm is not None:
                        val = vm.get(qn("w:val"))
                        if val is None or val == "":
                            # vMerge=continue：消耗数据但不填值
                            # 仍然记录映射
                            fid = f"T{t_idx}_G{g_num}_R{row_offset}_C{data_idx}"
                            rg_field_values[fid] = str(row_data[data_idx])
                            data_idx += 1
                            continue
                
                # 获取格内文本
                all_text = []
                for p in tc.findall(qn("w:p")):
                    for r in p.findall(qn("w:r")):
                        for t in r.findall(qn("w:t")):
                            if t.text:
                                all_text.append(t.text)
                text = "".join(all_text).strip()
                
                # 空白格或占位符格：填入数据
                if not text or text in ("%", "…", "……"):
                    rPr_source = _find_label_rPr_in_row(tr)
                    _set_tc_text(tc, str(row_data[data_idx]), rPr_source=rPr_source)
                    fid = f"T{t_idx}_G{g_num}_R{row_offset}_C{data_idx}"
                    rg_field_values[fid] = str(row_data[data_idx])
                    data_idx += 1
                else:
                    # 有文字的格：如果不是标签则跳过
                    pass
        
        # 如果数据行数 > 模板行数，需要复制行
        if len(row_data_list) > template_row_count:
            for extra_idx in range(template_row_count, len(row_data_list)):
                # 复制最后一行
                last_row = table.rows[start_row + template_row_count - 1]
                new_tr = copy.deepcopy(last_row._tr)
                last_row._tr.addnext(new_tr)
                
                # 填充新行
                tcs = new_tr.findall(qn("w:tc"))
                row_data = row_data_list[extra_idx]
                data_idx = 0
                for ti, tc in enumerate(tcs):
                    if data_idx >= len(row_data):
                        break
                    tcPr = tc.find(qn("w:tcPr"))
                    if tcPr is not None:
                        vm = tcPr.find(qn("w:vMerge"))
                        if vm is not None:
                            val = vm.get(qn("w:val"))
                            if val is None or val == "":
                                fid = f"T{t_idx}_G{g_num}_R{extra_idx}_C{data_idx}"
                                rg_field_values[fid] = str(row_data[data_idx])
                                data_idx += 1
                                continue
                    all_text = []
                    for p in tc.findall(qn("w:p")):
                        for r in p.findall(qn("w:r")):
                            for t in r.findall(qn("w:t")):
                                if t.text:
                                    all_text.append(t.text)
                    text = "".join(all_text).strip()
                    if not text or text in ("%", "…", "……"):
                        _set_tc_text(tc, str(row_data[data_idx]))
                        fid = f"T{t_idx}_G{g_num}_R{extra_idx}_C{data_idx}"
                        rg_field_values[fid] = str(row_data[data_idx])
                        data_idx += 1
    
    return rg_field_values


# ── 用户字段简化 ──

def _simplify_generic_fields(label_fields):
    """通用字段简化：自动识别前缀归组，将92+个内部字段压缩为用户友好的少量字段。
    
    归组规则：
    1. "前缀_数字" 模式 → 归为"前缀"组（如 课程目标1_1, 课程目标1_2 → 课程目标1）
    2. "前缀_第N列" 模式 → 归入已有组（如 课程目标2_第6列 → 课程目标2）
    3. 已有group类型字段 → 保留
    4. 单独字段 → 直接保留
    """
    import re
    
    # ── 第1步：识别所有字段的基础前缀 ──
    # 模式: "课程目标1_1" → base="课程目标1", sub="1"
    # 模式: "课程目标2_第6列" → base="课程目标2", sub="第6列"
    # 模式: "期末考试_课程目标" → base="期末考试", sub="课程目标"
    
    groups = {}       # base_label → {"type", "sub_labels", "sub_fields" / "field"}
    field_order = []  # 保持字段出现顺序
    
    # 列索引后缀（不单独成字段，归入父组）
    col_suffixes_re = re.compile(r'^(第\d+列|第\d+格)$')
    # 纯数字后缀
    num_suffixes_re = re.compile(r'^(\d+)$')
    
    for f in label_fields:
        label = f["label"]
        
        # 如果字段本身就是group类型（template_analyzer识别的multi_col字段），直接保留
        if f.get("fill_mode") == "group":
            if label not in groups:
                field_order.append(label)
            groups[label] = {
                "type": "group",
                "base_label": label,
                "sub_labels": f.get("sub_labels", []),
                "sub_fields": [f],
            }
            continue
        
        # 尝试按下划线拆分
        parts = label.split("_", 1)
        if len(parts) == 2:
            base, suffix = parts
            
            # 判断后缀是否应该归组
            should_group = (
                col_suffixes_re.match(suffix) or   # 第N列
                num_suffixes_re.match(suffix) or   # 纯数字
                suffix in ("课程目标", "实现途径", "评价方法",  # 语义子字段
                           "实际平均分", "目标达成评价值") or
                any(k in suffix for k in ["列", "题", "列"])
            )
            
            if should_group:
                if base in groups:
                    # 追加到已有组
                    if groups[base]["type"] == "single":
                        # 升级single → group
                        existing = groups[base]["field"]
                        groups[base] = {
                            "type": "group",
                            "base_label": base,
                            "sub_labels": [base, suffix],
                            "sub_fields": [existing, f],
                        }
                    else:
                        groups[base]["sub_labels"].append(suffix)
                        groups[base]["sub_fields"].append(f)
                else:
                    field_order.append(base)
                    groups[base] = {
                        "type": "group",
                        "base_label": base,
                        "sub_labels": [suffix],
                        "sub_fields": [f],
                    }
                continue
        
        # 未匹配归组规则 → 独立字段
        if label not in groups:
            field_order.append(label)
            groups[label] = {
                "type": "single",
                "base_label": label,
                "field": f,
            }
    
    # ── 第2步：合并连续的列索引子字段为范围描述 ──
    # 如 sub_labels=["1","2","3","4","5","第6列",...,"第26列"] 
    # → sub_labels=["数据1~5", "数据6~26"]
    
    user_fields = []
    for label in field_order:
        g = groups[label]
        if g["type"] == "single":
            f = g["field"]
            user_fields.append({
                "label": f["label"],
                "description": f.get("description", f"请填写{f['label']}"),
                "fill_mode": f.get("fill_mode", "set"),
            })
        else:
            subs = g["sub_labels"]
            base = g["base_label"]
            
            # 压缩子标签：将连续的列索引合并
            compressed = _compress_sub_labels(subs)
            
            user_fields.append({
                "label": base,
                "description": f"请填写{base}的相关数据",
                "fill_mode": "group",
                "sub_labels": compressed,
            })
    
    return user_fields


def _compress_sub_labels(subs):
    """压缩子标签列表：连续的列索引用范围表示。
    
    ["1","2","3","4","5","第6列","第7列",...,"第26列"]
    → ["数据1~5", "数据6~26"]
    """
    if len(subs) <= 5:
        return subs
    
    import re
    num_re = re.compile(r'^(\d+)$')
    col_re = re.compile(r'^第(\d+)列$')
    
    # 分类：纯数字 vs 列索引 vs 其他
    num_items = []   # (index, value, sort_key)
    col_items = []   # (index, value, sort_key)
    other_items = [] # (index, value)
    
    for i, s in enumerate(subs):
        m = num_re.match(s)
        if m:
            num_items.append((i, s, int(m.group(1))))
            continue
        m = col_re.match(s)
        if m:
            col_items.append((i, s, int(m.group(1))))
            continue
        other_items.append((i, s))
    
    result = []
    
    # 纯数字范围
    if num_items:
        sorted_nums = sorted(num_items, key=lambda x: x[2])
        min_n = sorted_nums[0][2]
        max_n = sorted_nums[-1][2]
        if max_n - min_n + 1 == len(sorted_nums):
            result.append(f"数据{min_n}~{max_n}（共{len(sorted_nums)}项）")
        else:
            result.append(f"数据{len(sorted_nums)}项")
    
    # 列索引范围
    if col_items:
        sorted_cols = sorted(col_items, key=lambda x: x[2])
        min_c = sorted_cols[0][2]
        max_c = sorted_cols[-1][2]
        if max_c - min_c + 1 == len(sorted_cols):
            result.append(f"第{min_c}列~第{max_c}列（共{len(sorted_cols)}项）")
        else:
            result.append(f"列数据{len(sorted_cols)}项")
    
    # 其他
    for _, s in other_items:
        result.append(s)
    
    return result if result else subs


def _simplify_fields(label_fields):
    """将内部字段列表简化为用户友好的字段列表。
    
    将 _第1列 等子字段归组到基础字段下。
    """
    # 子字段后缀列表
    sub_suffixes = ["_第1列", "_第2列", "_第3列", "_第4列", "_第5列",
                    "_一题", "_二题", "_三题", "_四题", "_五题",
                    "_六题", "_七题", "_八题", "_九题",
                    "_<60", "_60-69", "_70-79", "_80-89", "_90-100",
                    "_应到", "_实到", "_缺考", "_缓考", "_作弊", "_取消考试资格"]
    
    groups = {}  # base_label → {"type": "single"/"group", ...}
    field_order = []
    
    for f in label_fields:
        label = f["label"]
        
        base = None
        suffix = None
        for s in sub_suffixes:
            if label.endswith(s):
                base = label[:-len(s)]
                suffix = s[1:]
                break
        
        if base and suffix:
            if base in groups and groups[base]["type"] == "group":
                # 追加到已有组
                groups[base]["sub_labels"].append(suffix)
                groups[base]["sub_fields"].append(f)
            elif base in groups and groups[base]["type"] == "single":
                # 升级为group
                existing = groups[base]["field"]
                groups[base] = {
                    "type": "group",
                    "base_label": base,
                    "sub_fields": [existing, f],
                    "sub_labels": [base, suffix],
                }
            else:
                if base not in groups:
                    field_order.append(base)
                groups[base] = {
                    "type": "group",
                    "base_label": base,
                    "sub_fields": [f],
                    "sub_labels": [suffix],
                }
        else:
            if label not in groups:
                field_order.append(label)
                groups[label] = {
                    "type": "single",
                    "base_label": label,
                    "field": f,
                }
    
    user_fields = []
    for label in field_order:
        g = groups[label]
        if g["type"] == "single":
            f = g["field"]
            user_fields.append({
                "field_id": f["field_id"],
                "label": f["field_id"],
                "raw_label": f["label"],
                "description": f"{f['label']} - 请填写{f['label']}",
                "fill_mode": f["fill_mode"],
            })
        else:
            # 分组字段 - 检查是否需要拆分考勤数据
            subs = g["sub_labels"]
            base = g["base_label"]
            
            # 检查是否包含考勤子字段
            attendance_suffixes = {"应到", "实到", "缺考", "缓考", "作弊", "取消考试资格"}
            has_attendance = any(s in attendance_suffixes for s in subs)
            
            if has_attendance and base != "考勤":
                # 拆分：base字段 + 考勤组
                # 先添加base字段，跳过"_第N列"等内部子字段
                internal_suffixes = {"第1列", "第2列", "第3列", "第4列", "第5列"}
                base_fields = [sf for sf in g["sub_fields"] 
                              if not any(sf["label"].endswith(f"_{s}") for s in attendance_suffixes)
                              and not any(sf["label"].endswith(f"_{s}") for s in internal_suffixes)]
                for sf in base_fields:
                    user_fields.append({
                        "field_id": sf["field_id"],
                        "label": sf["field_id"],
                        "raw_label": sf["label"],
                        "description": f"{sf['label']} - 请填写{sf['label']}",
                        "fill_mode": sf["fill_mode"],
                    })
                # 添加考勤组
                att_labels = [s for s in subs if s in attendance_suffixes]
                if att_labels:
                    # 找出考勤子字段的 field_id
                    att_fids = [sf["field_id"] for sf in g["sub_fields"] 
                                if any(sf["label"].endswith(f"_{s}") for s in att_labels)]
                    att_labels_with_fid = [f"[{fid}] {lbl}" for fid, lbl in zip(att_fids, att_labels)]
                    user_fields.append({
                        "label": "考勤数据",
                        "description": "请填写考勤数据：应到、实到、缺考、缓考、作弊、取消考试资格人数",
                        "fill_mode": "group",
                        "sub_labels": att_fids,
                        "sub_field_ids": att_fids,
                    })
            else:
                # 提取子字段的 field_id
                sub_fids = [sf["field_id"] for sf in g["sub_fields"]]
                user_fields.append({
                    "label": base,
                    "description": f"请填写{base}的相关数据",
                    "fill_mode": "group",
                    "sub_labels": sub_fids,
                    "sub_field_ids": sub_fids,
                })
    
    return user_fields


# ── 数据扩展 ──

def _expand_report_data(template_path: str, user_data: dict) -> dict:
    """将用户友好的字段数据扩展为完整的内部字段数据。
    
    处理逻辑：
    1. 逗号分隔的分组字段（如"及百分比": "15,15,20,..."）自动展开为子字段
    2. 勾选框字段（如"课程性质": "必修"）由_fill_checkbox_rows_in_table处理，此处原样保留
    3. 考勤数据自动映射到子字段
    4. 分数段比例自动计算（如果用户未提供）
    """
    analysis = analyze_template(template_path)
    label_fields = analysis["label_fields"]
    
    # 构建基础名→子字段映射
    sub_suffixes = ["_第1列", "_第2列", "_第3列", "_第4列", "_第5列",
                    "_一题", "_二题", "_三题", "_四题", "_五题",
                    "_六题", "_七题", "_八题", "_九题",
                    "_<60", "_60-69", "_70-79", "_80-89", "_90-100",
                    "_应到", "_实到", "_缺考", "_缓考", "_作弊", "_取消考试资格"]
    
    base_to_subs = {}
    for f in label_fields:
        label = f["label"]
        for s in sub_suffixes:
            if label.endswith(s):
                base = label[:-len(s)]
                if base not in base_to_subs:
                    base_to_subs[base] = []
                base_to_subs[base].append({"full_label": label, "suffix": s[1:], "field": f})
                break
    
    expanded = {}
    
    # 考勤子字段名映射
    attendance_suffixes = {"应到", "实到", "缺考", "缓考", "作弊", "取消考试资格"}
    
    # 先处理独立的考勤字段（应到=45, 实到=43等），映射到full_label
    for key, value in list(user_data.items()):
        if key in attendance_suffixes:
            # 找到匹配的full_label（如 学生班级_应到）
            for base_name, subs in base_to_subs.items():
                for sub in subs:
                    if sub["suffix"] == key:
                        expanded[sub["full_label"]] = str(value)
                        if key in user_data:
                            del user_data[key]  # 已处理，避免重复
    
    for key, value in user_data.items():
        if key == "考勤数据":
            # 处理考勤数据组
            if isinstance(value, dict):
                for att_key, att_val in value.items():
                    # 查找匹配的考勤子字段
                    for base_name, subs in base_to_subs.items():
                        for sub in subs:
                            if sub["suffix"] == att_key:
                                expanded[sub["full_label"]] = str(att_val)
            elif isinstance(value, str) and "," in value:
                # 逗号分隔格式：应到,实到,缺考,缓考,作弊,取消资格
                parts = [p.strip() for p in value.split(",")]
                # 按考勤子字段的顺序映射
                for base_name, subs in base_to_subs.items():
                    att_subs = [s for s in subs if s["suffix"] in attendance_suffixes]
                    att_subs.sort(key=lambda s: ["应到", "实到", "缺考", "缓考", "作弊", "取消考试资格"].index(s["suffix"]))
                    for i, sub in enumerate(att_subs):
                        if i < len(parts):
                            expanded[sub["full_label"]] = parts[i]
            continue
        
        if key in base_to_subs:
            subs = base_to_subs[key]
            # 检查是否包含考勤子字段
            has_attendance = any(s["suffix"] in attendance_suffixes for s in subs)
            
            if has_attendance:
                # 拆分：先填非考勤子字段，考勤数据另存
                non_att_subs = [s for s in subs if s["suffix"] not in attendance_suffixes]
                if isinstance(value, str) and "," in value and non_att_subs:
                    parts = [p.strip() for p in value.split(",")]
                    for i, sub in enumerate(non_att_subs):
                        if i < len(parts):
                            expanded[sub["full_label"]] = parts[i]
                else:
                    expanded[key] = value
            elif isinstance(value, list) and len(value) == len(subs):
                # 列表值，按顺序映射
                for i, sub in enumerate(subs):
                    expanded[sub["full_label"]] = value[i]
            elif isinstance(value, str) and "," in value:
                # 逗号分隔值，按顺序映射到子字段
                parts = [p.strip() for p in value.split(",")]
                for i, sub in enumerate(subs):
                    if i < len(parts):
                        expanded[sub["full_label"]] = parts[i]
            else:
                # 单个值，保留原样（勾选框逻辑由_fill_checkbox_rows_in_table处理）
                expanded[key] = value
        else:
            # 直接字段，原样保留
            expanded[key] = value
    
    # 自动计算分数段比例（如果用户未提供但有分数段人数）
    dist_bases = [b for b in base_to_subs if b.startswith("分数分布")]
    for base_name in dist_bases:
        subs = base_to_subs[base_name]
        ratio_subs = [s for s in subs if s["suffix"].startswith("第")]
        count_subs = [s for s in subs if s["suffix"] in ["<60", "60-69", "70-79", "80-89", "90-100"]]
        
        if ratio_subs and count_subs:
            # 检查是否已有比例数据
            has_ratio = any(s["full_label"] in expanded for s in ratio_subs)
            if not has_ratio:
                # 计算比例
                total = 0
                counts = []
                for s in count_subs:
                    v = expanded.get(s["full_label"], "0")
                    try:
                        n = float(v)
                    except (ValueError, TypeError):
                        n = 0
                    counts.append(n)
                    total += n
                
                if total > 0:
                    for i, s in enumerate(ratio_subs):
                        if i < len(counts):
                            pct = counts[i] / total * 100
                            expanded[s["full_label"]] = f"{pct:.1f}%"
    
    return expanded


# ── 主生成函数 ──

def _build_report_docx(template_path: str, user_data: dict):
    """根据模板和用户数据生成填充后的docx文件。"""
    analysis = analyze_template(template_path)
    doc = Document(template_path)
    
    # 1. 扩展用户数据
    expanded_data = _expand_report_data(template_path, user_data)
    
    # 2. 识别勾选框行，排除与勾选框冲突的标签字段
    checkbox_rows = set()  # {(table_idx, row_idx)}
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            unique = _get_unique_cells(row)
            groups = _detect_checkbox_row(unique)
            if groups:
                checkbox_rows.add((t_idx, r_idx))
    
    # 构建勾选框空白格的位置集合 (table_idx, row_idx, col_idx)
    checkbox_blank_cells = set()
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            unique = _get_unique_cells(row)
            groups = _detect_checkbox_row(unique)
            for g in groups:
                for ci in g["option_blanks"].values():
                    checkbox_blank_cells.add((t_idx, r_idx, ci))

    # 过滤掉勾选框字段（pattern=checkbox由_fill_checkbox_rows_in_table单独处理）
    # 以及勾选框空白格位置的label字段（避免重复填充冲突）
    # 但保留勾选框行中非空白格位置的label字段（如"教师姓名"）
    # 同时分离出段落级下划线字段（pattern=paragraph_underline由_fill_paragraph_fields处理）
    filtered_fields = []
    paragraph_fields = []
    for f in analysis["label_fields"]:
        # 1. checkbox模式字段直接跳过
        if f.get("pattern") == "checkbox":
            continue
        # 2. paragraph_underline模式字段单独处理
        if f.get("pattern") == "paragraph_underline":
            paragraph_fields.append(f)
            continue
        # 3. 勾选框空白格位置的字段跳过
        t_idx = f["table_idx"]
        skip = False
        for r_idx in f["row_indices"]:
            if (t_idx, r_idx, f["col_idx"]) in checkbox_blank_cells:
                skip = True
                break
        if not skip:
            filtered_fields.append(f)
    
    # 3. 填充标签字段（不含勾选框行）
    _fill_label_fields(doc, filtered_fields, expanded_data)
    
    # 3.5 填充段落级下划线字段
    _fill_paragraph_fields(doc, paragraph_fields, expanded_data)
    
    # 4. 填充勾选框行（独立于标签字段，智能识别选项+空白格模式）
    _fill_checkbox_rows_in_table(doc, analysis, expanded_data)
    
    # 5. 填充行组
    rg_field_values = _fill_simple_row_groups(doc, analysis["row_groups"], expanded_data)
    
    # 保存到临时文件
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.seek(0)
    content = tmp.read()
    tmp.close()
    os.unlink(tmp.name)
    
    return content, analysis, expanded_data, rg_field_values


def _expand_generic_data(label_fields, user_data):
    """将通用模板中归组字段的逗号分隔值展开成内部子字段名。
    
    例: 用户输入 {"课程目标1": "0.8,0.7,0.6,..."} 
    → {"课程目标1_1": "0.8", "课程目标1_2": "0.7", ...}
    """
    expanded = dict(user_data)
    
    # 收集所有归组前缀及其子字段
    groups = {}  # base_label -> [(sub_field_label, sub_field)]
    for f in label_fields:
        label = f["label"]
        for sep in ("_", "·"):
            if sep in label:
                base, sub = label.split(sep, 1)
                if base not in groups:
                    groups[base] = []
                groups[base].append((label, f))
                break
    
    # 对每个归组，展开用户数据
    for base, subs in groups.items():
        if base in user_data and isinstance(user_data[base], str):
            values = [v.strip() for v in user_data[base].split(",")]
            # 按子字段顺序赋值
            for i, (sub_label, sub_f) in enumerate(subs):
                if i < len(values) and values[i]:
                    expanded[sub_label] = values[i]
            # 删除原始归组键（避免重复填充）
            if base in expanded:
                del expanded[base]
    
    return expanded


def _fill_custom_template(template_path: str, user_data: dict):
    """通用模板纯填充：只向空白格/待填格写入文本，绝不改变文档格式。
    
    与 _build_report_docx 的区别：
    - 不调用 _expand_report_data（内置模板专用扩展逻辑）
    - 使用 _expand_generic_data 展开通用归组字段
    - 完整保留原有字体、字号、加粗、对齐等格式
    """
    analysis = analyze_template(template_path)
    doc = Document(template_path)
    
    # 0. 展开归组字段数据
    expanded_data = _expand_generic_data(analysis["label_fields"], user_data)
    
    # 1. 构建勾选框行检测，排除冲突的标签字段
    checkbox_blank_cells = set()
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            unique = _get_unique_cells(row)
            groups = _detect_checkbox_row(unique)
            for g in groups:
                for ci in g["option_blanks"].values():
                    checkbox_blank_cells.add((t_idx, r_idx, ci))
    
    # 过滤掉勾选框空白格位置的label字段
    # 同时分离出段落级下划线字段
    filtered_fields = []
    paragraph_fields = []
    for f in analysis["label_fields"]:
        if f.get("pattern") == "paragraph_underline":
            paragraph_fields.append(f)
            continue
        t_idx = f["table_idx"]
        skip = False
        for r_idx in f["row_indices"]:
            if (t_idx, r_idx, f["col_idx"]) in checkbox_blank_cells:
                skip = True
                break
        if not skip:
            filtered_fields.append(f)
    
    # 2. 填充标签字段（使用展开后的数据）
    _fill_label_fields(doc, filtered_fields, expanded_data)
    
    # 2.5 填充段落级下划线字段
    _fill_paragraph_fields(doc, paragraph_fields, expanded_data)
    
    # 3. 填充勾选框行
    _fill_checkbox_rows_in_table(doc, analysis, expanded_data)
    
    # 4. 填充行组（使用展开后的数据）
    rg_field_values = _fill_simple_row_groups(doc, analysis["row_groups"], expanded_data)
    
    # 5. 填充多列字段（multi_col字段，如评价报告中的课程目标表）
    for f in analysis["label_fields"]:
        if f.get("pattern") == "multi_col" and f["label"] in expanded_data:
            _fill_multi_col_field(doc, f, expanded_data[f["label"]])
    
    # 保存
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.seek(0)
    content = tmp.read()
    tmp.close()
    os.unlink(tmp.name)
    
    return content, analysis, expanded_data, rg_field_values


def _fill_multi_col_field(doc, field, value):
    """填充multi_col模式字段（一行内多个待填列）。"""
    t_idx = field["table_idx"]
    table = doc.tables[t_idx]
    
    # value可能是逗号分隔字符串或列表
    if isinstance(value, str):
        values = [v.strip() for v in value.replace("，", ",").split(",")]
    else:
        values = list(value)
    
    # 获取可填充列索引
    fillable_cols = field.get("fillable_cols", [])
    
    for i, col_idx in enumerate(fillable_cols):
        if i >= len(values):
            break
        for ri in field["row_indices"]:
            tr = table.rows[ri]._tr
            tcs = tr.findall(qn("w:tc"))
            if col_idx < len(tcs):
                rPr_source = _find_label_rPr_in_row(tr)
                _set_tc_text(tcs[col_idx], str(values[i]), rPr_source=rPr_source)


# ── 工具函数 ──

@tool
def list_templates() -> str:
    """列出所有可用的教务报告模板。"""
    templates = []
    for name, path in TEMPLATE_REGISTRY.items():
        templates.append({"name": name, "file": os.path.basename(path)})
    return json.dumps({"success": True, "templates": templates}, ensure_ascii=False)


@tool
def analyze_report_template(template_name: str) -> str:
    """解析指定模板，返回用户友好的字段清单和信息收集指引。
    在开始收集用户信息前必须先调用此工具。

    Args:
        template_name: 模板名称，从list_templates返回的名称中选择
    """
    try:
        path = _get_template_path(template_name)
        analysis = analyze_template(path)
        
        # 简化字段列表
        user_fields = _simplify_fields(analysis["label_fields"])
        
        # 构建收集指引
        guide_parts = []
        guide_parts.append(f"模板【{template_name}】共需填写 {len(user_fields)} 项信息：\n")
        
        for i, f in enumerate(user_fields, 1):
            if f["fill_mode"] == "group":
                subs = f.get("sub_labels", [])
                if len(subs) <= 5:
                    guide_parts.append(f"  {i}. {f['label']}（{', '.join(subs)}）")
                else:
                    guide_parts.append(f"  {i}. {f['label']}（共{len(subs)}项，可用逗号分隔）")
            else:
                guide_parts.append(f"  {i}. {f['label']}")
        
        guide_parts.append('\n提示：多值字段可用逗号分隔一次性提供，如"及百分比: 15,15,20,20,10,10,5,5,0"')
        
        return json.dumps({
            "success": True,
            "template_name": template_name,
            "total_fields": len(user_fields),
            "user_fields": user_fields,
            "row_groups_count": analysis["summary"]["total_row_groups"],
            "collection_guide": "\n".join(guide_parts),
        }, ensure_ascii=False)
        
    except Exception as e:
        return json.dumps({"success": False, "message": f"模板解析失败: {e}"}, ensure_ascii=False)


@tool
def generate_edu_report(template_name: str, report_data: str) -> str:
    """根据预设模板和用户数据生成教务报告文档（仅支持预设的3种模板）。

    Args:
        template_name: 模板名称（评价报告/试卷分析/关联矩阵）
        report_data: JSON格式的报告数据，键为字段名，值为字段值。
                     多值字段用逗号分隔，如 {"及百分比": "15,15,20,20,10,10,5,5,0"}
                     行组数据用二维数组，如 {"T0_G0": [["值1","值2"],...]}
    """
    ctx = request_context.get() or new_context(method="generate_edu_report")
    
    try:
        path = _get_template_path(template_name)
        
        # 解析报告数据
        if isinstance(report_data, str):
            data = json.loads(report_data)
        else:
            data = report_data
        
        # 生成文档
        docx_bytes, analysis, expanded_data, rg_field_values = _build_report_docx(path, data)
        
        # 上传到对象存储
        from coze_coding_dev_sdk.s3 import S3SyncStorage
        import time
        
        storage = S3SyncStorage(
            endpoint_url=os.getenv("COZE_BUCKET_ENDPOINT_URL"),
            access_key="",
            secret_key="",
            bucket_name=os.getenv("COZE_BUCKET_NAME"),
            region="cn-beijing",
        )
        
        timestamp = time.strftime("%Y%m%d%H%M%S")
        safe_name = template_name.replace(" ", "_")
        file_name = f"edu_report/{safe_name}_{timestamp}.docx"
        
        file_key = storage.upload_file(
            file_content=docx_bytes,
            file_name=file_name,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        url = storage.generate_presigned_url(key=file_key, expire_time=86400)
        
        # 保存本地副本供打印转PDF使用
        local_path = os.path.join("/tmp", os.path.basename(file_name))
        with open(local_path, "wb") as lf:
            lf.write(docx_bytes)

        # 生成后文档校验
        validation = validate_doc(path, local_path)

        # 提取非空字段数据供前端更新预览
        filled_data = {k: v for k, v in data.items() if v and str(v).strip()}
        filled_field_values = _build_field_id_value_map(analysis["label_fields"], expanded_data)
        filled_field_values.update(rg_field_values)  # 合并行组field_id映射

        result = {
            "success": True,
            "message": "报告已成功生成并上传",
            "file_name": file_name,
            "download_url": url,
            "local_path": local_path,
            "filled_data": filled_data,
            "filled_field_values": filled_field_values,
            "validation": validation,
        }

        # 如果校验发现硬伤，标记但不阻断（前端可提示用户）
        if not validation["valid"]:
            result["message"] = f"报告已生成，但存在结构问题: {'; '.join(validation['errors'])}"

        return json.dumps(result, ensure_ascii=False)
        
    except Exception as e:
        return json.dumps({
            "success": False,
            "message": f"报告生成失败: {e}",
        }, ensure_ascii=False)


# ═══════════════════════════════════════════════════════════════
# 通用模板工具：支持任意docx文件的自动识别和填充
# ═══════════════════════════════════════════════════════════════

@tool
def analyze_uploaded_template(file_path: str) -> str:
    """分析用户上传的任意Word模板文件，自动识别待填字段。
    支持识别：冒号字段、标签+空白格、勾选框、占位符、多列数据行、行组等。

    Args:
        file_path: 上传的docx文件路径（通常是上传后的临时文件路径）
    """
    ctx = request_context.get() or new_context(method="analyze_uploaded_template")
    
    try:
        import os
        workspace = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")
        if not os.path.isabs(file_path):
            full_path = os.path.join(workspace, file_path)
        else:
            full_path = file_path
        
        if not os.path.exists(full_path):
            return json.dumps({"success": False, "message": f"文件不存在: {file_path}"}, ensure_ascii=False)
        
        analysis = analyze_template(full_path)
        # 通用模板使用通用简化函数（比_simplify_fields更智能地归组）
        user_fields = _simplify_generic_fields(analysis['label_fields'])
        
        # 构建用户友好的字段描述
        field_descriptions = []
        for f in user_fields:
            if f['fill_mode'] == 'group':
                subs = f['sub_labels']
                field_descriptions.append({
                    "label": f['label'],
                    "type": "group",
                    "sub_items": subs,
                    "hint": f"请提供{len(subs)}个值，用逗号分隔",
                })
            else:
                desc = {"label": f['label'], "type": "single"}
                # 添加提示
                label_lower = f['label']
                if any(k in label_lower for k in ['签字', '签名']):
                    desc['hint'] = "请填写姓名"
                elif any(k in label_lower for k in ['日期', '时间']):
                    desc['hint'] = "请填写日期"
                elif any(k in label_lower for k in ['百分比', '比例', '占比']):
                    desc['hint'] = "请填写数值"
                field_descriptions.append(desc)
        
        # 行组信息
        row_group_info = []
        for g in analysis['row_groups']:
            row_group_info.append({
                "group_id": g['group_id'],
                "num_cols": g['num_cols'],
                "template_row_count": g['template_row_count'],
                "column_labels": g.get('column_labels', []),
            })
        
        return json.dumps({
            "success": True,
            "file_name": os.path.basename(full_path),
            "total_fields": len(user_fields),
            "fields": field_descriptions,
            "row_groups": row_group_info,
            "summary": f"识别到{len(user_fields)}个待填字段和{len(row_group_info)}个数据行组",
        }, ensure_ascii=False)
        
    except Exception as e:
        return json.dumps({"success": False, "message": f"模板分析失败: {e}"}, ensure_ascii=False)


@tool
def generate_from_template(file_path: str, report_data: str) -> str:
    """根据用户上传的模板文件和填写数据，生成填充后的文档。
    支持任意docx模板文件的自动填充。

    Args:
        file_path: 模板文件路径（与analyze_uploaded_template使用的路径相同）
        report_data: JSON格式的填写数据，键为字段名，值为字段值。
                     多值字段用逗号分隔，如 {"及百分比": "15,15,20,20,10,10,5,5,0"}
                     行组数据用二维数组，如 {"T0_G0": [["值1","值2"],...]}
    """
    ctx = request_context.get() or new_context(method="generate_from_template")
    
    try:
        import os
        workspace = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")
        if not os.path.isabs(file_path):
            full_path = os.path.join(workspace, file_path)
        else:
            full_path = file_path
        
        if not os.path.exists(full_path):
            return json.dumps({"success": False, "message": f"文件不存在: {file_path}"}, ensure_ascii=False)
        
        data = json.loads(report_data)
        
        # 通用模板：直接填充，不调用expand（保留原格式）
        doc_bytes, analysis, expanded_data, rg_field_values = _fill_custom_template(full_path, data)
        
        # 上传到对象存储
        from coze_coding_dev_sdk.s3 import S3SyncStorage
        import time
        
        storage = S3SyncStorage(
            endpoint_url=os.getenv("COZE_BUCKET_ENDPOINT_URL"),
            access_key="",
            secret_key="",
            bucket_name=os.getenv("COZE_BUCKET_NAME"),
            region="cn-beijing",
        )
        
        timestamp = time.strftime("%Y%m%d%H%M%S")
        safe_name = os.path.splitext(os.path.basename(full_path))[0].replace(" ", "_")
        file_name = f"custom_report/{safe_name}_{timestamp}.docx"
        
        file_key = storage.upload_file(
            file_content=doc_bytes,
            file_name=file_name,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        url = storage.generate_presigned_url(key=file_key, expire_time=86400)
        
        # 保存本地副本供打印转PDF使用
        local_path = os.path.join("/tmp", os.path.basename(file_name))
        with open(local_path, "wb") as lf:
            lf.write(doc_bytes)

        # 提取非空字段数据供前端更新预览
        filled_data = {k: v for k, v in data.items() if v and str(v).strip()}
        filled_field_values = _build_field_id_value_map(analysis["label_fields"], expanded_data)
        filled_field_values.update(rg_field_values)  # 合并行组field_id映射

        # 生成后文档校验
        validation = validate_doc(full_path, local_path)

        result = {
            "success": True,
            "message": "文档已成功生成并上传",
            "file_name": file_key,
            "download_url": url,
            "local_path": local_path,
            "filled_data": filled_data,
            "filled_field_values": filled_field_values,
            "validation": validation,
        }

        if not validation["valid"]:
            result["message"] = f"文档已生成，但存在结构问题: {'; '.join(validation['errors'])}"

        return json.dumps(result, ensure_ascii=False)
        
    except Exception as e:
        return json.dumps({
            "success": False,
            "message": f"文档生成失败: {e}",
        }, ensure_ascii=False)
