"""
old_report_extractor — 旧报告反向提取 + 数据准备清单

核心能力：
1. 从已填写的旧报告 docx 中反向提取字段值
2. 将提取结果自动预填入 FormFillingState
3. 生成数据准备清单（分类：必填/可继承/可计算/可选）
"""

import os
import re
import json
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from langchain.tools import tool
from coze_coding_utils.log.write_log import request_context
from coze_coding_utils.runtime_ctx.context import new_context

from tools.template_analyzer import analyze_template
from tools.form_filling_state import FormFillingState


# ── 模板注册表（与 edu_report_tool.py 同步） ──

TEMPLATE_REGISTRY = {
    "评价报告": "assets/2023-2024-2《xxx》 岭南师范学院专业课程目标达成度评价报告模板.docx",
    "试卷分析": "assets/2023-2024-2《xxx》 试卷分析模板.docx",
    "关联矩阵": "assets/2023-2024-2《xxx》岭南师范学院考题与课程目标及毕业要求关联矩阵表模板.docx",
}

# ── 会话缓存（与 edu_report_tool.py 共享引用） ──
# 通过 import 注入，避免循环依赖
_active_form_states: dict = {}


def inject_form_states(states_dict: dict) -> None:
    """由 agent.py 启动时调用，注入 edu_report_tool 的 _active_form_states 引用"""
    global _active_form_states
    _active_form_states = states_dict


# ── 模板路径解析 ──

def _resolve_template_path(template_name_or_path: str) -> str:
    """将模板名称或路径解析为实际文件路径"""
    workspace = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")
    for name, rel_path in TEMPLATE_REGISTRY.items():
        if name == template_name_or_path:
            return os.path.join(workspace, rel_path)
    if os.path.exists(template_name_or_path):
        return template_name_or_path
    assets_path = os.path.join(workspace, "assets", template_name_or_path)
    if os.path.exists(assets_path):
        return assets_path
    raise ValueError(
        f"找不到模板: {template_name_or_path}，可用模板: [{', '.join(TEMPLATE_REGISTRY.keys())}]"
    )


# ── 字段 ID 解析 ──

_FIELD_ID_RE = re.compile(r"T(\d+)_R(\d+)_C(\d+)(?:_L(\d+))?")


def _parse_field_id(field_id: str) -> Optional[tuple]:
    """解析 field_id 为 (table_idx, row_idx, col_idx, label_run_idx)"""
    m = _FIELD_ID_RE.match(field_id)
    if not m:
        return None
    return (int(m.group(1)), int(m.group(2)), int(m.group(3)),
            int(m.group(4)) if m.group(4) is not None else -1)


# ── 标签剥离 ──

def _strip_label(cell_text: str, label: str) -> str:
    """从单元格文本中剥离标签前缀，提取纯值部分。

    处理格式：
      "课程名称：高等数学" → "高等数学"
      "课程名称:高等数学"   → "高等数学"
      "课程名称 高等数学"   → "高等数学"
      "课程名称高等数学"    → "高等数学"
    """
    if not cell_text or not label:
        return cell_text.strip()

    text = cell_text.strip()

    # 尝试常见分隔符
    for sep in ("：", ":", " "):
        pattern = label + sep
        if pattern in text:
            return text.split(pattern, 1)[1].strip()

    # 无分隔符：标签是前缀
    if text.startswith(label):
        remainder = text[len(label):].strip()
        if remainder:
            return remainder

    # 合并单元格导致多个标签混入，如 "课程名称:高等数学\n开课时间:2024-2025-2"
    # 只提取本标签对应的值
    for sep in ("：", ":"):
        if sep in text:
            parts = text.split("\n")
            for part in parts:
                part = part.strip()
                if label in part:
                    val = part.split(sep, 1)[-1].strip()
                    if val and val != label:
                        return val


def _parse_label_value_pairs(cell_text: str) -> dict[str, str]:
    """从单元格文本中解析所有 label:value 对。

    处理格式如：
    "课程名称:高等数学B"  → {"课程名称": "高等数学B"}
    "开课时间:2024-2025-2\\n考试类别:期末考试"  → {"开课时间": "2024-2025-2", "考试类别": "期末考试"}
    "参评人数：55"  → {"参评人数": "55"}
    """
    result: dict[str, str] = {}
    lines = cell_text.replace("\r\n", "\n").split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            continue

        for sep in ("：", ":"):
            if sep in line:
                label_part, _, value_part = line.partition(sep)
                label_part = label_part.strip()
                value_part = value_part.strip()
                if label_part and value_part and label_part != value_part:
                    result[label_part] = value_part
                break  # 只用第一个分隔符

    return result

    # 标签不在开头——整段文本可能就是值
    return cell_text.strip()


# ── 核心提取逻辑 ──

def _extract_field_values(
    doc,  # python-docx Document object
    analysis_result: dict,
) -> dict[str, str]:
    """从已填写的 docx 中反向提取字段值。

    Args:
        doc: python-docx Document 对象（旧报告）
        analysis_result: 模板分析结果（字段位置和标签信息）

    Returns:
        {raw_label: value} 映射，仅包含有值的字段
    """
    extracted: dict[str, str] = {}

    # 按 table 分组字段，以便一次性检测合并单元格
    table_fields: dict[int, list] = {}
    for field in analysis_result.get("label_fields", []):
        pos = _parse_field_id(field.get("field_id", ""))
        if pos:
            t_idx = pos[0]
            table_fields.setdefault(t_idx, []).append(field)

    # ── 1. 处理标签字段（append / set / replace / check） ──
    for t_idx, fields in table_fields.items():
        if t_idx >= len(doc.tables):
            continue
        table = doc.tables[t_idx]

        # 按 row 分组
        row_fields: dict[int, list] = {}
        for field in fields:
            pos = _parse_field_id(field.get("field_id", ""))
            if pos:
                row_fields.setdefault(pos[1], []).append(field)

        for r_idx, row_field_list in row_fields.items():
            if r_idx >= len(table.rows):
                continue

            row = table.rows[r_idx]

            # Build: merged_col → merge_source_col mapping
            merged_to_source: dict[int, int] = {}
            seen_tc: dict[int, int] = {}  # tc_id → first col index
            for ci, cell in enumerate(row.cells):
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    merged_to_source[ci] = seen_tc[tc_id]
                else:
                    seen_tc[tc_id] = ci

            # 预提取：每个合并组的源单元格文本，解析出 label→value 对
            source_cell_pairs: dict[int, dict[str, str]] = {}
            row_label_value_pairs: dict[int, dict[str, str]] = {}  # 所有单元格（含被合并的）
            for ci in set(merged_to_source.values()) | set(seen_tc.values()):
                cell_text = _read_cell_text(doc, t_idx, r_idx, ci)
                if cell_text:
                    pairs = _parse_label_value_pairs(cell_text)
                    source_cell_pairs[ci] = pairs
                    row_label_value_pairs[ci] = pairs
            # 合并单元格也继承源单元格的 pairs
            for merged_ci, src_ci in merged_to_source.items():
                if src_ci in source_cell_pairs:
                    row_label_value_pairs[merged_ci] = source_cell_pairs[src_ci]

            for field in row_field_list:
                field_id = field.get("field_id", "")
                raw_label = field.get("raw_label", field.get("label", ""))
                fill_mode = field.get("fill_mode", "set")

                pos = _parse_field_id(field_id)
                if not pos:
                    continue

                _, _, c_idx, _ = pos

                source_col = merged_to_source.get(c_idx, c_idx)
                cell_text = _read_cell_text(doc, t_idx, r_idx, source_col)
                if not cell_text:
                    continue

                # 策略1: 直接从预解析的 label→value 对中查找（处理合并单元格标签错位）
                pairs = source_cell_pairs.get(source_col, {})
                if raw_label in pairs and pairs[raw_label]:
                    extracted[raw_label] = pairs[raw_label]
                    continue

                # 策略1.5: 在整行所有单元格的 pairs 中查找 raw_label（跨单元格匹配）
                found = False
                for col_pairs in row_label_value_pairs.values():
                    if raw_label in col_pairs and col_pairs[raw_label]:
                        extracted[raw_label] = col_pairs[raw_label]
                        found = True
                        break
                if found:
                    continue

                # 策略2: 按 fill_mode 提取值（兜底）
                value = _extract_value_by_mode(cell_text, raw_label, fill_mode)
                if value and _is_meaningful_value(value):
                    extracted[raw_label] = value

    # ── 2. 处理段落下划线字段 ──
    for field in analysis_result.get("paragraph_fields", []):
        raw_label = field.get("raw_label", field.get("label", ""))
        para_idx = field.get("paragraph_idx", -1)
        run_idx = field.get("run_idx", -1)

        if para_idx < 0 or para_idx >= len(doc.paragraphs):
            continue

        para = doc.paragraphs[para_idx]
        if run_idx >= 0 and run_idx < len(para.runs):
            text = para.runs[run_idx].text.strip()
        else:
            text = para.text.strip()

        if text and text != raw_label:
            extracted[raw_label] = text

    return extracted


def _read_cell_text(doc, t_idx: int, r_idx: int, c_idx: int) -> str:
    """安全读取 docx 表格单元格文本"""
    if t_idx >= len(doc.tables):
        return ""
    table = doc.tables[t_idx]
    if r_idx >= len(table.rows):
        return ""
    row = table.rows[r_idx]
    if c_idx >= len(row.cells):
        return ""
    return row.cells[c_idx].text.strip()


def _extract_value_by_mode(cell_text: str, raw_label: str, fill_mode: str) -> str:
    """根据 fill_mode 从单元格文本提取值"""
    if fill_mode == "append":
        return _strip_label(cell_text, raw_label)

    if fill_mode in ("set", "replace"):
        # set/replace 模式：先剥离标签，再排除空值/占位符
        value = _strip_label(cell_text, raw_label)
        if not value:
            return ""
        if value == raw_label:
            return ""
        if value in _PLACEHOLDER_TEXTS:
            return ""
        return value

    if fill_mode == "check":
        if cell_text in _CHECK_VALUES:
            return cell_text

    return ""


_PLACEHOLDER_TEXTS = frozenset({"%", "…", "……", "xxx", "yyy", "yyyy", "请填写"})
_CHECK_VALUES = frozenset({"√", "✓", "☑", "是", "合格", "有"})


def _is_meaningful_value(value: str) -> bool:
    """判断提取的值是否为有意义的用户数据（非标签/占位符/空值）。"""
    if not value or not value.strip():
        return False
    v = value.strip()
    # 纯占位符
    if v in _PLACEHOLDER_TEXTS:
        return False
    # 看起来像标签（以冒号结尾，没有实际数据）
    if v.endswith(("：", ":")) and len(v) <= 20:
        return False
    # 纯标签组合（如 "实现途径:\n评价方法:"）
    if all(part.strip().endswith(("：", ":")) or not part.strip() for part in v.split("\n")):
        return False
    return True


# ── 数据准备清单 ──

# 字段分类关键词
_MUST_FILL_KEYWORDS = [
    "课程名称", "教学班级", "教师", "开课", "学期", "学院",
    "课程性质", "考核方式", "命题形式", "考试形式",
]
_COMPUTABLE_KEYWORDS = [
    "平均分", "最高分", "最低分", "标准差", "达成度", "达成评价",
    "百分比", "占比", "比例", "及格率",
]
_INHERITABLE_HINT = "可从上学期旧报告自动继承"
_COMPUTABLE_HINT = "可从成绩数据自动计算"


def _categorize_field(raw_label: str, fill_mode: str) -> str:
    """将字段分类为 must_fill / computable / inheritable / optional"""
    label_lower = raw_label.lower() if raw_label else ""

    for kw in _MUST_FILL_KEYWORDS:
        if kw in label_lower:
            return "must_fill"

    for kw in _COMPUTABLE_KEYWORDS:
        if kw in label_lower:
            return "computable"

    if fill_mode == "check":
        return "inheritable"

    if any(kw in label_lower for kw in ("签字", "参与人", "备注", "说明", "建议", "改进")):
        return "optional"

    return "inheritable"


def build_fill_checklist(analysis_result: dict) -> dict:
    """根据模板分析结果生成数据准备清单。

    Returns:
        {
            "template_name": str,
            "total_fields": int,
            "categories": {
                "must_fill": {"description": ..., "fields": [...], "count": N},
                "computable": {"description": ..., "fields": [...], "count": N, "hint": ...},
                "inheritable": {"description": ..., "fields": [...], "count": N, "hint": ...},
                "optional":   {"description": ..., "fields": [...], "count": N},
            }
        }
    """
    categories = {
        "must_fill": {"description": "每次必填（无法自动获取）", "fields": [], "hint": "请提前准备好"},
        "computable": {
            "description": "可自动计算（需提供成绩数据）",
            "fields": [], "hint": _COMPUTABLE_HINT,
        },
        "inheritable": {
            "description": "可从旧报告继承（跨学期基本不变）",
            "fields": [], "hint": _INHERITABLE_HINT,
        },
        "optional": {"description": "可选填写", "fields": []},
    }

    for field in analysis_result.get("label_fields", []):
        raw_label = field.get("raw_label", field.get("label", ""))
        fill_mode = field.get("fill_mode", "set")

        cat = _categorize_field(raw_label, fill_mode)

        categories[cat]["fields"].append({
            "field_id": field.get("field_id", ""),
            "label": raw_label,
            "fill_mode": fill_mode,
        })

    # 统计
    for cat_data in categories.values():
        cat_data["count"] = len(cat_data["fields"])

    return {
        "template_name": analysis_result.get("template_name", ""),
        "total_fields": len(analysis_result.get("label_fields", [])),
        "categories": categories,
    }


# ── @tool 函数 ──


@tool
def extract_from_old_report(
    file_path: str,
    template_name_or_path: str = "",
) -> str:
    """从已填写的旧报告 docx 中反向提取字段值。

    传入旧报告文件路径和对应模板名称（如"评价报告"），
    返回提取到的 {字段名: 值} 映射。可用于 update_form_fields 预填。

    Args:
        file_path: 旧报告 docx 文件路径
        template_name_or_path: 模板名称（如"评价报告"）或模板文件路径
    """
    ctx = request_context.get() or new_context(method="extract_from_old_report")

    try:
        if not os.path.exists(file_path):
            return json.dumps(
                {"success": False, "message": f"文件不存在: {file_path}"},
                ensure_ascii=False,
            )

        # 解析模板路径
        template_path = _resolve_template_path(template_name_or_path)

        # 分析模板结构（获取字段位置信息）
        analysis = analyze_template(template_path)

        # 打开旧报告
        doc = Document(file_path)

        # 反向提取
        extracted = _extract_field_values(doc, analysis)

        return json.dumps({
            "success": True,
            "template_name": template_name_or_path,
            "extracted_count": len(extracted),
            "extracted_fields": extracted,
            "message": f"从旧报告提取到 {len(extracted)} 个字段值",
        }, ensure_ascii=False)

    except Exception as e:
        return json.dumps(
            {"success": False, "message": f"提取失败: {e}"},
            ensure_ascii=False,
        )


@tool
def prefill_from_old_report(
    session_id: str,
    old_report_path: str,
) -> str:
    """从旧报告提取字段值并自动预填入当前填写会话。

    在 init_form_filling 之后调用，自动将旧报告中的数据填入状态机，
    返回预填进度和仍需手动填写的字段清单。

    Args:
        session_id: 填写会话 ID（init_form_filling 时创建）
        old_report_path: 旧报告 docx 文件路径
    """
    ctx = request_context.get() or new_context(method="prefill_from_old_report")

    try:
        # 获取会话状态
        state = _active_form_states.get(session_id)
        if not state:
            return json.dumps(
                {"success": False, "message": f"会话不存在: {session_id}，请先调用 init_form_filling"},
                ensure_ascii=False,
            )

        if not os.path.exists(old_report_path):
            return json.dumps(
                {"success": False, "message": f"文件不存在: {old_report_path}"},
                ensure_ascii=False,
            )

        # 解析模板路径
        template_path = _resolve_template_path(state.template_name or state.template_path)

        # 分析模板结构
        analysis = analyze_template(template_path)

        # 打开旧报告并提取
        doc = Document(old_report_path)
        extracted = _extract_field_values(doc, analysis)

        if not extracted:
            return json.dumps({
                "success": True,
                "prefilled_count": 0,
                "message": "旧报告中未提取到可填入的字段值",
                "progress": state.get_progress(),
            }, ensure_ascii=False)

        # 批量写入状态机
        matched, unmatched = state.bulk_fill(
            extracted, confidence=0.9, source="old_report"
        )

        return json.dumps({
            "success": True,
            "prefilled_count": len(matched),
            "unmatched_count": len(unmatched),
            "prefilled_fields": {k: v for k, v in extracted.items() if k in matched},
            "still_missing": state.get_missing_important(),
            "progress": state.get_progress(),
            "message": f"从旧报告预填了 {len(matched)} 个字段",
        }, ensure_ascii=False)

    except Exception as e:
        return json.dumps(
            {"success": False, "message": f"预填失败: {e}"},
            ensure_ascii=False,
        )


@tool
def get_fill_checklist(
    template_name_or_path: str = "",
    session_id: str = "",
) -> str:
    """获取数据准备清单，按必填/可继承/可计算/可选分类展示。

    帮助用户了解需要准备哪些数据，以及哪些数据可以自动获取。
    可在 init_form_filling 之前调用（提供模板名），也可在之后调用（提供 session_id）。

    Args:
        template_name_or_path: 模板名称或路径（session_id 为空时必填）
        session_id: 填写会话 ID（提供时自动附带已填进度）
    """
    ctx = request_context.get() or new_context(method="get_fill_checklist")

    try:
        analysis = None
        state = None

        # 优先从会话获取分析结果
        if session_id and session_id in _active_form_states:
            state = _active_form_states[session_id]
            if state._analysis_result:
                analysis = state._analysis_result
            elif state.template_path:
                template_path = _resolve_template_path(
                    state.template_name or state.template_path
                )
                analysis = analyze_template(template_path)

        # 回退到模板名
        if not analysis and template_name_or_path:
            template_path = _resolve_template_path(template_name_or_path)
            analysis = analyze_template(template_path)

        if not analysis:
            return json.dumps({
                "success": False,
                "message": "请提供 template_name_or_path 或 session_id",
            }, ensure_ascii=False)

        # 生成清单
        checklist = build_fill_checklist(analysis)

        # 如果有会话状态，附加已填进度
        if state:
            checklist["current_progress"] = state.get_progress()
            # 标记已填字段
            filled_labels = set()
            for fid, finfo in state._fields.items():
                if finfo["status"] != FormFillingState.EMPTY and finfo["value"]:
                    filled_labels.add(finfo["raw_label"] or finfo["label"])

            for cat_name, cat_data in checklist["categories"].items():
                for f in cat_data["fields"]:
                    f["already_filled"] = f["label"] in filled_labels

        return json.dumps({
            "success": True,
            "checklist": checklist,
        }, ensure_ascii=False)

    except Exception as e:
        return json.dumps(
            {"success": False, "message": f"生成清单失败: {e}"},
            ensure_ascii=False,
        )
