"""模板解析器 - 自动识别Word模板中所有可填充字段，输出精简的用户收集清单

支持的字段模式：
1. "标签：值"模式 - 冒号后为空则待填（如 "课程名称："）
2. "标签格+空白格"模式 - 标签文字单独在一个格，紧邻空白格为待填（如 |课程名称|[空白]|）
3. "标签格+占位符格"模式 - 标签文字后有占位符（如 |题号|[%]|）需替换占位符
4. "多列数据行"模式 - 表头行下有空白/占位符行，按表头名识别（如 题型百分比、分数段人数等）
"""

import os
import re
import json
import logging
from dataclasses import dataclass, field, asdict
from typing import Optional
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# 复杂度阈值：列数超过此值的行组不自动填充
_MAX_SIMPLE_GROUP_COLS = 15

# 已知的标签文字黑名单（这些不是待填字段，而是表头或装饰文字）
_LABEL_BLACKLIST = {'序号', '编号', '合计', '总计', '备注', '说明', '项', '次', '类', '号'}

# 选项类标签不算待填字段
_OPTION_WORDS = {'选修', '必修', '开卷', '闭卷', '半开卷', '试题库', '试卷库', '教师组题', '是', '否',
                '优秀', '良好', '中等', '及格', '不及格', 'A', 'B', 'C', 'D', '√', '✓', '○', '●',
                '本人阅卷', '同行阅卷', '集体阅卷', '机器阅卷', '其他'}

# 占位符文字（需要替换为实际值的格）
_PLACEHOLDER_TEXTS = {'%', '…', '……', '...', '—', '--', '___', '□', '○'}

# 日期占位符模式（年月日及其变体，含全角/半角空格）
_DATE_PLACEHOLDER_RE = re.compile(r'^[年]\s*[月]\s*[日]?\s*$|^[年][\s　]*[月][\s　]*[日]?\s*$')

# 宽松日期占位符模式：包含"年月日"及其上下文（如"上午""下午"）
_DATE_PLACEHOLDER_LOOSE_RE = re.compile(
    r'年\s*月\s*日'  # 核心模式
    r'(?:\s*(?:上午|下午|上|下)\s*)*'  # 可选的上午/下午
)

# 嵌入式日期占位符模式：标签文字中夹带 年/月/日/上午/下午 等占位符
# 匹配如 "考试时间  年  月  日  上午  下午" 或 "负责人签名  年  月  日"
_EMBEDDED_DATE_RE = re.compile(
    r'^(.+?)\s*年\s*月\s*日'   # 标签 + 年月日
    r'(?:\s*(上午|下午|上|下))?'  # 可选的上午/下午
    r'\s*$'
)

# 章节标题行标签（这些行不包含待填字段，只是标题）
_SECTION_TITLE_LABELS = {'试卷分析', '试　卷　分　析', '课程目标与试卷题目的对应关系'}


def _get_unique_cells(row) -> list:
    """获取行中的独立单元格 (跳过被合并重复的)"""
    seen = set()
    cells = []
    for cell in row.cells:
        cid = id(cell._element)
        if cid not in seen:
            seen.add(cid)
            cells.append(cell)
    return cells


def _is_section_title_row(unique_cells) -> bool:
    """判断是否是章节标题行"""
    if len(unique_cells) != 1:
        return False
    text = unique_cells[0].text.strip()
    # 标准编号标题
    if re.match(r'^[一二三四五六七八九十]+[、.．]', text):
        return True
    # 全宽空格标题（如"试　卷　分　析"）
    if text in _SECTION_TITLE_LABELS:
        return True
    # gridSpan跨整行的标题格
    return False


def _is_empty_data_row(unique_cells) -> bool:
    """判断是否是空白数据行：大部分格为空，且不含标签模式。"""
    if not unique_cells:
        return True
    empty = sum(1 for c in unique_cells if not c.text.strip())
    # 含冒号的行 = 标签行
    label_count = sum(1 for c in unique_cells if re.search(r'[：:]', c.text.strip()))
    # 含短标签文字（2-8字）且不含冒号的格 = 可能是标签格
    short_label_count = sum(1 for c in unique_cells
                           if c.text.strip() and 2 <= len(c.text.strip()) <= 8
                           and not re.search(r'[：:]', c.text.strip())
                           and c.text.strip() not in _LABEL_BLACKLIST
                           and c.text.strip() not in _OPTION_WORDS)
    return empty >= len(unique_cells) * 0.5 and label_count == 0 and short_label_count == 0


def _is_header_row(unique_cells) -> bool:
    """判断是否是表头行"""
    if not unique_cells:
        return False
    non_empty = sum(1 for c in unique_cells if c.text.strip())
    has_colon = any('：' in c.text or ':' in c.text for c in unique_cells)
    return non_empty >= len(unique_cells) * 0.6 and not has_colon


def _find_header_for_row(table, row_idx) -> Optional[int]:
    """向上查找最近的表头行"""
    for r in range(row_idx - 1, -1, -1):
        unique = _get_unique_cells(table.rows[r])
        if _is_header_row(unique):
            return r
    return None


def _extract_labels_from_cell(cell) -> list[tuple[str, str]]:
    """从单元格中提取所有"标签: 值"对（冒号模式），返回 [(label, existing_value), ...]

    也识别"标签+嵌入式日期占位符"模式（如"考试时间  年  月  日  上午  下午"）
    以及多段组合单元格（如"审批意见  负责人签名：  年月日  公章"）
    """
    results = []
    text = cell.text.strip().replace('\r', '')
    if not text:
        return results

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        # 模式1：冒号模式 "标签：值"
        # 对于"审批意见  负责人签名：  年月日  公章"这种多段文本，
        # 用split分段提取每个"短标签：值"对，而非只取第一个冒号
        if '：' in line or ':' in line:
            # 按冒号分割，每次取"冒号前的短标签（最后2-15字）:冒号后的值"
            seg_parts = re.split(r'([：:])', line)
            for seg_i in range(0, len(seg_parts) - 2, 2):
                before = seg_parts[seg_i].strip()
                after = seg_parts[seg_i + 2].strip() if seg_i + 2 < len(seg_parts) else ''
                # 尝试从before的后部提取短标签（回退到前一个空白处，取最后2-15字）
                # 如"数据申请部门审批意见     负责人签名" → 标签="负责人签名"
                short_label = before
                # 如果before太长，取最后一个空白分隔的短词
                if len(before) > 15:
                    space_parts = re.split(r'\s{2,}', before)
                    if len(space_parts) >= 2 and 1 <= len(space_parts[-1]) <= 15:
                        short_label = space_parts[-1]
                    # 也尝试用单个空格/全角空格分隔
                    alt_parts = re.split(r'[\s　]+', before)
                    if len(alt_parts) >= 2 and 1 <= len(alt_parts[-1]) <= 15:
                        short_label = alt_parts[-1]
                # 标签不能太长或太短
                if 1 <= len(short_label) <= 20:
                    # 清理value：截断掉可能混入的后段标签（如"年月日"后还有"公章"）
                    value = after
                    # 如果value中有"申请部门公章"等非占位符词，尝试精简
                    if value and len(value) > 20:
                        # 保留到日期占位符结束为止
                        date_end = 0
                        for m in re.finditer(r'年\s*月\s*日', value):
                            date_end = m.end()
                        if date_end > 0 and date_end < len(value):
                            value = value[:date_end]
                    results.append((short_label, value))
                    if seg_i + 2 >= len(seg_parts) - 2:
                        break  # 只处理最后一个冒号（最内层的标签）
            continue

        # 模式2：嵌入式日期占位符 "标签  年  月  日" 或 "标签  年  月  日  上午  下午"
        date_match = _EMBEDDED_DATE_RE.match(line)
        if date_match:
            label = date_match.group(1).strip()
            remaining = line[len(label):].strip()
            results.append((label, remaining))
            continue

        # 模式3：标签+填写说明 "主要教学经历（授课名称、起止时间...）"
        hint_match = re.match(r'^(.+?)（[^）]{2,}）\s*$', line)
        if hint_match:
            label = hint_match.group(1).strip()
            results.append((label, line))
            continue

    return results


def _is_label_cell(cell) -> bool:
    """判断单元格是否是"标签格"（有短文字，不含冒号，不是纯数字）"""
    text = cell.text.strip()
    if not text:
        return False
    # 太长的不算标签
    if len(text) > 15:
        return False
    # 含冒号的由冒号模式处理
    if '：' in text or ':' in text:
        return False
    # 纯数字不是标签
    if re.match(r'^[\d.]+$', text):
        return False
    # 纯标点符号不算标签
    if re.match(r'^[^\w\u4e00-\u9fff]+$', text):
        return False
    # 黑名单
    if text in _LABEL_BLACKLIST:
        return False
    # 选项类（也检查去空格版本，如"闭  卷"→"闭卷"）
    if text in _OPTION_WORDS or text.replace(" ", "").replace("　", "") in _OPTION_WORDS:
        return False
    # 常见表头词不算标签
    header_words = {'试题', '题号', '分数', '得分', '评卷人', '项目', '内容', '签名', '日期'}
    if text in header_words:
        return False
    # 日期占位符不算标签
    if _DATE_PLACEHOLDER_RE.match(text) or _DATE_PLACEHOLDER_LOOSE_RE.match(text):
        return False
    return True


def _is_placeholder_cell(cell) -> bool:
    """判断单元格是否是占位符格（如%、……、年月日、提示文本等）"""
    text = cell.text.strip()
    if text in _PLACEHOLDER_TEXTS:
        return True
    if _DATE_PLACEHOLDER_RE.match(text):
        return True
    # 宽松匹配：以"年月日"开头且不包含其他实质性文字
    if _DATE_PLACEHOLDER_LOOSE_RE.match(text) and len(text) <= 50:
        return True
    # 长提示文本：如"请在此详细说明…"、"请在方框内勾选…"
    if text.startswith(('请在', '请填写', '请勾选', '请在此', '请选择')):
        return True
    # 带括号的描述性提示：如"（授课名称、起止时间…）"
    if text.startswith('（') and len(text) > 10:
        return True
    # 勾选框选项文本：如"□通修基础□专业基础…"
    if text.startswith('□') and len(text) > 5:
        return True
    return False


def _is_vmerge_continue(cell) -> bool:
    """判断单元格是否是垂直合并的延续格（非起始格）"""
    tc = cell._element
    vm = tc.find(qn('w:tcPr'))
    if vm is not None:
        vm_elem = vm.find(qn('w:vMerge'))
        if vm_elem is not None:
            val = vm_elem.get(qn('w:val'))
            return val is None or val == "continue"
    return False


def _get_grid_span(cell) -> int:
    """获取单元格的gridSpan值"""
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        gs = tcPr.find(qn('w:gridSpan'))
        if gs is not None:
            val = gs.get(qn('w:val'))
            if val:
                return int(val)
    return 1


def _scan_paragraph_underline_fields(doc) -> list:
    """扫描文档正文段落中的下划线横线字段（如"教材名称______"）。
    
    这些字段不在表格中，而是文档正文段落里的"标签+下划线空白"格式。
    
    Returns:
        list[dict]: 字段列表，每个字段包含:
            - field_id: "P{paragraph_idx}"
            - label: 标签文字
            - pattern: "paragraph_underline"
            - table_idx: -1 (不在表格中)
            - row_idx: 段落索引
            - col_idx: -1
            - existing_value: "" (空白待填)
            - fill_mode: "paragraph_underline"
            - description: 描述文字
            - row_indices: [段落索引]
            - repeat_count: 1
    """
    fields = []
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # 检查段落中是否包含下划线run
        label_text = ""
        has_underline_blank = False
        for run in para.runs:
            is_underline = False
            # 检查python-docx的underline属性
            if run.underline and run.underline not in (False, 0):
                is_underline = True
            # 也检查XML中的w:u元素
            if not is_underline:
                rPr = run._element.find(qn('w:rPr'))
                if rPr is not None:
                    u_elem = rPr.find(qn('w:u'))
                    if u_elem is not None:
                        val = u_elem.get(qn('w:val'))
                        if val and val not in ('none',):
                            is_underline = True
            
            if is_underline:
                # 下划线run，内容应该是空白/空格（待填区域）
                if not run.text.strip():
                    has_underline_blank = True
                else:
                    # 下划线run有内容，可能已有值，暂不处理
                    pass
            else:
                # 非下划线run，收集标签文字
                label_text += run.text
        
        label_text = label_text.strip()
        
        if not has_underline_blank or not label_text:
            continue
        
        # 过滤：标签太短或太长的不合理
        if len(label_text) < 2 or len(label_text) > 20:
            continue
        
        # 过滤黑名单
        if label_text in _LABEL_BLACKLIST:
            continue
        
        # 过滤选项类
        normalized = label_text.replace(" ", "").replace("　", "")
        if label_text in _OPTION_WORDS or normalized in _OPTION_WORDS:
            continue
        
        fields.append({
            "field_id": f"P{p_idx}",
            "label": label_text,
            "pattern": "paragraph_underline",
            "table_idx": -1,
            "row_idx": p_idx,
            "col_idx": -1,
            "existing_value": "",
            "fill_mode": "paragraph_underline",
            "description": f"请填写{label_text}",
            "row_indices": [p_idx],
            "repeat_count": 1,
        })
    
    return fields


def analyze_template(template_path: str) -> dict:
    """
    解析模板文件，返回精简的字段清单和行组信息。

    返回结构:
    {
        "label_fields": [...],
        "row_groups": [...],
        "summary": {...}
    }
    """
    doc = Document(template_path)
    label_fields = []
    row_groups = []

    # 记录已经被行组覆盖的行（避免标签字段和行组重复）
    row_group_rows = {}  # {(table_idx, row_idx): group_id}

    for t_idx, table in enumerate(doc.tables):
        # ── 0.5 预扫描：提取每行的section标题（整行合并单元格） ──
        # section标题是单元格文本的第一行或冒号前的简短文字
        _row_section_titles = {}  # {(t_idx, r_idx): section_title}
        for r_idx, row in enumerate(table.rows):
            unique = _get_unique_cells(row)
            if len(unique) == 1:
                full_text = unique[0].text.strip().replace('\n', '|')
                # 尝试取第一行（以|分隔）
                first_line = full_text.split('|')[0].strip()
                # 尝试取冒号/换行前的文字
                for sep in ['：', ':']:
                    if sep in first_line:
                        first_line = first_line.split(sep)[0].strip()
                        break
                # section标题特征：文字简短（2~25字），不是占位符
                if 2 <= len(first_line) <= 25 and not _is_placeholder_cell(unique[0]):
                    _row_section_titles[(t_idx, r_idx)] = first_line

        # ── 1. 扫描冒号模式字段 ("标签：值") ──
        for r_idx, row in enumerate(table.rows):
            unique = _get_unique_cells(row)
            if _is_section_title_row(unique):
                continue

            # 查找最近的上方section标题作为上下文
            section_ctx = None
            for up_r in range(r_idx - 1, -1, -1):
                if (t_idx, up_r) in _row_section_titles:
                    section_ctx = _row_section_titles[(t_idx, up_r)]
                    break

            for c_idx, cell in enumerate(unique):
                labels = _extract_labels_from_cell(cell)
                for line_idx, (label, existing_value) in enumerate(labels):
                    field_id = f"T{t_idx}_R{r_idx}_C{c_idx}_L{line_idx}"
                    # 如果有section上下文，用"section标题-label"格式区分
                    display_label = label
                    desc = f"请填写{label}" if not existing_value else f"{label}(已有:{existing_value})"
                    if section_ctx and label in ("负责人签名", "申请人签名", "签名", "日期", "审批意见"):
                        # 截取section标题的关键区分部分
                        short_ctx = section_ctx
                        if len(section_ctx) > 8:
                            # 取关键词：去掉"审批意见"/"审批"等后缀
                            for suffix in ["审批意见", "审批", "意见"]:
                                if section_ctx.endswith(suffix) and len(section_ctx) > len(suffix) + 2:
                                    short_ctx = section_ctx[:-len(suffix)]
                                    break
                            # 如果仍超8字，进一步缩短
                            if len(short_ctx) > 8:
                                for suffix2 in ["教学运行数据", "管理部门"]:
                                    if short_ctx.endswith(suffix2):
                                        short_ctx = short_ctx[:-len(suffix2)]
                                        break
                        display_label = f"{short_ctx}-{label}"
                        desc = f"请填写{short_ctx}的{label}" if not existing_value else f"{short_ctx}-{label}(已有:{existing_value})"
                    # 日期占位符需要replace而非append
                    fill_mode = "append"
                    if existing_value and _DATE_PLACEHOLDER_LOOSE_RE.match(existing_value.strip()):
                        fill_mode = "replace"
                    # 标签+填写说明模式：整个单元格需要被替换（如"主要教学经历（授课名称...）"）
                    if existing_value and re.match(r'^.+?（[^）]{2,}）\s*$', existing_value):
                        fill_mode = "replace"
                    label_fields.append({
                        "field_id": field_id,
                        "table_idx": t_idx,
                        "row_idx": r_idx,
                        "col_idx": c_idx,
                        "line_idx": line_idx,
                        "label": display_label,
                        "raw_label": label,
                        "existing_value": existing_value,
                        "description": desc,
                        "fill_mode": fill_mode,
                        "pattern": "colon",
                        "section_context": section_ctx,
                    })

        # ── 2. 扫描"标签格+空白格/占位符格"模式 ──
        # 扩展：不仅检查相邻格，还检查中间跳过选项标签格的情况
        # 注意：当一行有"行标签+多列待填"模式时，跳过label_blank扫描（交给multi_col处理）
        for r_idx, row in enumerate(table.rows):
            unique = _get_unique_cells(row)
            if _is_section_title_row(unique):
                continue

            # 检测该行是否是"行标签+多列待填"模式
            # 多列数据行特征：一个标签后连续>=2个空白格/占位符格，且中间没有标签格或选项格
            # 标签-空白交替行特征：标签-空白-标签-空白 交替出现
            # 选项-空白交替行特征：选项-空白-选项-空白 交替出现（勾选框行）
            is_multi_col_row = False
            first_label_idx = -1
            consecutive_fillable = 0
            max_consecutive_fillable = 0
            label_blank_alternating = False
            option_blank_alternating = False
            
            prev_was_label = False
            prev_was_fillable = False
            prev_was_option = False
            for ci, cell in enumerate(unique):
                if _is_vmerge_continue(cell):
                    continue
                text = cell.text.strip()
                is_label = text and _is_label_cell(cell)
                is_option = text and (text in _OPTION_WORDS or text.replace(" ", "").replace("　", "") in _OPTION_WORDS)
                is_fillable = not text or _is_placeholder_cell(cell)
                
                if is_label and first_label_idx < 0:
                    first_label_idx = ci
                
                if (is_label or is_option) and prev_was_fillable:
                    if is_label:
                        label_blank_alternating = True
                    if is_option:
                        option_blank_alternating = True
                
                if is_fillable and first_label_idx >= 0:
                    consecutive_fillable += 1
                    max_consecutive_fillable = max(max_consecutive_fillable, consecutive_fillable)
                elif is_label or is_option:
                    consecutive_fillable = 0
                
                prev_was_label = is_label
                prev_was_fillable = is_fillable
                prev_was_option = is_option
            
            # 多列数据行：有行标签，后面连续>=2个待填格，且不是交替模式
            is_multi_col_row = (first_label_idx >= 0 and max_consecutive_fillable >= 2 
                               and not label_blank_alternating and not option_blank_alternating)
            
            i = 0
            while i < len(unique):
                cell = unique[i]
                
                if _is_label_cell(cell):
                    label = cell.text.strip()
                    
                    # 检查这个标签是否已经被冒号模式识别过
                    already_found = any(
                        f["table_idx"] == t_idx and f["row_idx"] == r_idx
                        and f["label"] == label and f["pattern"] == "colon"
                        for f in label_fields
                    )
                    
                    if not already_found and not is_multi_col_row:
                        # 向后寻找第一个空白格或占位符格
                        for j in range(i + 1, min(i + 4, len(unique))):  # 最多往后看3格
                            next_cell = unique[j]
                            next_text = next_cell.text.strip()
                            
                            if _is_vmerge_continue(next_cell):
                                continue
                            
                            if not next_text or _is_placeholder_cell(next_cell):
                                # 找到待填格
                                fill_mode = "set"
                                if _is_placeholder_cell(next_cell):
                                    fill_mode = "replace"  # 替换占位符
                                
                                field_id = f"T{t_idx}_R{r_idx}_C{j}"
                                label_fields.append({
                                    "field_id": field_id,
                                    "table_idx": t_idx,
                                    "row_idx": r_idx,
                                    "col_idx": j,  # 指向待填格
                                    "line_idx": 0,
                                    "label": label,
                                    "existing_value": next_text if next_text else "",
                                    "description": f"请填写{label}",
                                    "fill_mode": fill_mode,
                                    "pattern": "label_blank",
                                })
                                break
                            elif next_text in _OPTION_WORDS:
                                # 跳过选项格继续寻找
                                continue
                            else:
                                # 遇到非选项非空白格，停止寻找
                                break
                i += 1

            # ── 2.4 扫描独立占位符行（如"年　月　日"） ──
            # 当一行的所有非合并单元格都是占位符，没有标签格
            # 给它们友好的标签名
            _scan_standalone_placeholder_row(unique, t_idx, r_idx, label_fields)
            
            # ── 2.45 扫描勾选框行（如"闭卷  开卷  半开卷  其他"） ──
            _scan_checkbox_fields(unique, t_idx, r_idx, table, label_fields)

        # ── 2.5 扫描"表头+数据行"模式（多列数据行） ──
        # 处理试卷分析中的"题型百分比"、"分数段人数/比例"等区域
        # 这些区域特征：表头行有多个标签，下一行有空白格或占位符格待填
        _scan_table_data_sections(table, t_idx, label_fields)

        # ── 3. 检测重复行组 ──
        r = 0
        while r < len(table.rows):
            unique = _get_unique_cells(table.rows[r])
            if _is_empty_data_row(unique):
                header_row_idx = _find_header_for_row(table, r)
                if header_row_idx is not None:
                    header_unique = _get_unique_cells(table.rows[header_row_idx])
                    col_labels = []
                    for c in header_unique:
                        text = c.text.strip().replace('\n', '|')
                        col_labels.append(text)

                    # 统计连续空行
                    start = r
                    count = 0
                    while r < len(table.rows):
                        row_unique = _get_unique_cells(table.rows[r])
                        if _is_empty_data_row(row_unique):
                            start_unique = _get_unique_cells(table.rows[start])
                            if len(row_unique) == len(start_unique):
                                count += 1
                                row_group_rows[(t_idx, r)] = f"T{t_idx}_G{len(row_groups)}"
                                r += 1
                                continue
                        break

                    if count > 0:
                        num_cols = len(col_labels)
                        group_id = f"T{t_idx}_G{len(row_groups)}"
                        row_groups.append({
                            "group_id": group_id,
                            "table_idx": t_idx,
                            "start_row": start,
                            "template_row_count": count,
                            "column_labels": col_labels,
                            "header_row_idx": header_row_idx,
                            "header_text": " | ".join(col_labels),
                            "num_cols": num_cols,
                            "is_complex": num_cols > _MAX_SIMPLE_GROUP_COLS,
                        })
                        for marked_r in range(start, start + count):
                            row_group_rows[(t_idx, marked_r)] = group_id
                        continue
            r += 1

    # ── 3.5 扫描文档正文段落中的下划线横线字段（如"教材名称______"） ──
    paragraph_fields = _scan_paragraph_underline_fields(doc)
    label_fields.extend(paragraph_fields)

    # ── 4. 去重 label_fields: 同一表格同一列的重复标签合并 ──
    # 注意：有 section_context 的字段用 (t, col, section, label) 去重，
    #       不同审批环节的同名签名不应合并
    label_occurrences = {}
    for f in label_fields:
        section = f.get("section_context")
        if section:
            key = (f["table_idx"], f["col_idx"], section, f["label"])
        else:
            key = (f["table_idx"], f["col_idx"], f["label"])
        if key not in label_occurrences:
            label_occurrences[key] = []
        label_occurrences[key].append(f)

    deduped_fields = []
    for key, occurrences in label_occurrences.items():
        first = occurrences[0].copy()
        first["row_indices"] = [o["row_idx"] for o in occurrences]
        first["repeat_count"] = len(occurrences)
        if len(occurrences) > 1:
            first["description"] = f"请填写{first['label']} (共{len(occurrences)}处)"
        deduped_fields.append(first)

    # ── 5. 过滤掉已有值的字段（不需要用户填写） ──
    # 注意：占位符字段（fill_mode="replace"）的existing_value是占位符文本，需要保留
    deduped_fields = [f for f in deduped_fields
                      if not f["existing_value"] or f.get("fill_mode") == "replace"]
    
    # ── 5.5 过滤掉完全属于行组区域的multi_col字段 ──
    # 行组区域由行组填充逻辑处理，不需要单独的字段
    # 但保留部分在行组外的multi_col字段
    def _fully_in_row_group(field):
        """如果字段的所有行都在行组区域内，返回True"""
        for ri in field["row_indices"]:
            if (field["table_idx"], ri) not in row_group_rows:
                return False
        return True
    
    deduped_fields = [f for f in deduped_fields if not _fully_in_row_group(f)]

    # ── 5.6 消除占位符行与冒号模式/多列模式的冲突 ──
    # 场景：R17有"任课教师签字："（colon模式，空值），R18有"年月日"（standalone_placeholder+multi_col）
    # 策略1：同一行有standalone_placeholder和multi_col时，保留multi_col（有更具体的标签）
    # 策略2：colon字段空值且下一行有multi_col覆盖同位置时，移除colon字段（multi_col才是真正的填写目标）
    
    # 策略1：按行分组，如果某行同时有standalone_placeholder和multi_col，去掉standalone_placeholder
    from collections import defaultdict
    row_fields = defaultdict(list)
    for f in deduped_fields:
        for ri in f["row_indices"]:
            row_fields[(f["table_idx"], ri)].append(f)
    
    standalone_to_remove = set()
    for (t, r), fields_in_row in row_fields.items():
        has_multi_col = any(f["pattern"] == "multi_col" for f in fields_in_row)
        has_standalone = any(f["pattern"] == "standalone_placeholder" for f in fields_in_row)
        if has_multi_col and has_standalone:
            for f in fields_in_row:
                if f["pattern"] == "standalone_placeholder":
                    standalone_to_remove.add(id(f))
    
    deduped_fields = [f for f in deduped_fields if id(f) not in standalone_to_remove]
    
    # 策略2：colon字段空值 + 下一行有multi_col覆盖 → 移除colon字段
    # 重新按行分组
    row_fields2 = defaultdict(list)
    for f in deduped_fields:
        for ri in f["row_indices"]:
            row_fields2[(f["table_idx"], ri)].append(f)
    
    colon_to_remove = set()
    for f in deduped_fields:
        if f["pattern"] == "colon" and not f.get("existing_value"):
            t = f["table_idx"]
            r = f["row_idx"]
            c = f["col_idx"]
            # 检查下一行是否有multi_col在同列
            next_row_fields = row_fields2.get((t, r + 1), [])
            for nf in next_row_fields:
                if nf["pattern"] == "multi_col" and nf["col_idx"] == c:
                    colon_to_remove.add(id(f))
                    break
    
    deduped_fields = [f for f in deduped_fields if id(f) not in colon_to_remove]

    # ── 6. 按表格和行号排序 ──
    deduped_fields.sort(key=lambda f: (f["table_idx"], min(f["row_indices"]), f["col_idx"]))

    # ── 7. 标记行组覆盖的字段 ──
    for f in deduped_fields:
        for ri in f["row_indices"]:
            if (f["table_idx"], ri) in row_group_rows:
                f["in_row_group"] = row_group_rows[(f["table_idx"], ri)]
                break

    # ── 8. 生成收集计划 ──
    plan_lines = []
    plan_lines.append("【标签字段 - 待填写】")
    for f in deduped_fields:
        repeat = f" (x{f['repeat_count']})" if f["repeat_count"] > 1 else ""
        in_rg = f" [属于行组{f['in_row_group']}]" if f.get("in_row_group") else ""
        plan_lines.append(f"  - {f['label']}{repeat}: {f['description']}{in_rg}")

    plan_lines.append("\n【数据行组 - 表格区域待填】")
    simple_groups = [g for g in row_groups if not g.get("is_complex", False)]
    complex_groups = [g for g in row_groups if g.get("is_complex", False)]

    for g in simple_groups:
        plan_lines.append(f"  - {g['group_id']}: {g['template_row_count']}行x{g['num_cols']}列, 表头: {g['header_text']}")
    if complex_groups:
        plan_lines.append(f"\n【复杂行组 - 将跳过不填】")
        for g in complex_groups:
            plan_lines.append(f"  - {g['group_id']}: {g['num_cols']}列(>{_MAX_SIMPLE_GROUP_COLS}列), 表头: {g['header_text']}")

    result = {
        "label_fields": deduped_fields,
        "row_groups": row_groups,
        "summary": {
            "total_unique_labels": len(deduped_fields),
            "total_row_groups": len(row_groups),
            "simple_row_groups": len(simple_groups),
            "complex_row_groups": len(complex_groups),
            "collection_plan": "\n".join(plan_lines),
        }
    }

    return result


def _get_placeholder_label(text: str) -> str | None:
    """根据占位符文本返回友好标签名，如果不是占位符返回None"""
    stripped = text.strip()
    # 日期占位符（宽松匹配，包含"上午""下午"等）
    if _DATE_PLACEHOLDER_RE.match(stripped) or _DATE_PLACEHOLDER_LOOSE_RE.match(stripped):
        if '上午' in stripped or '下午' in stripped:
            return "考试时间"
        return "日期"
    return None


def _scan_standalone_placeholder_row(unique_cells, t_idx: int, r_idx: int, label_fields: list):
    """扫描独立占位符行——整行都是占位符（如'年　月　日'），没有标签格"""
    seen_elements = set()
    for c_idx, cell in enumerate(unique_cells):
        if _is_vmerge_continue(cell):
            continue
        # 去重：合并单元格指向同一XML元素
        elem_id = id(cell._element)
        if elem_id in seen_elements:
            continue
        seen_elements.add(elem_id)
        text = cell.text.strip()
        if not text:
            continue
        # 只处理占位符
        friendly_label = _get_placeholder_label(text)
        if not friendly_label:
            continue
        # 检查该格是否已在 label_fields 中（可能被前面的步骤识别了）
        fid = f"T{t_idx}_R{r_idx}_C{c_idx}"
        if any(f['field_id'] == fid for f in label_fields):
            continue
        # 添加为独立字段
        label_fields.append({
            "field_id": fid,
            "table_idx": t_idx,
            "row_idx": r_idx,
            "col_idx": c_idx,
            "label": friendly_label,
            "description": f"请填写{friendly_label}",
            "fill_mode": "set",
            "pattern": "standalone_placeholder",
            "existing_value": "",
        })
        # 只添加一次，不重复
        break


def _scan_checkbox_fields(unique_cells, t_idx: int, r_idx: int, table, label_fields: list):
    """扫描勾选框行，识别互斥选项组并将其作为字段暴露给Agent。
    
    模式: [行标签] [选项1] [空白1] [选项2] [空白2] ...
    或: [选项1] [空白1] [选项2] [空白2] ... (无行标签，上方有章节标题)
    """
    # 复用检测逻辑
    groups = _detect_checkbox_groups(unique_cells)
    if not groups:
        return
    
    for group in groups:
        row_label = group["label"]
        options = list(group["option_blanks"].keys())
        if not options:
            continue
        
        # 如果没有行标签，尝试从上方获取章节标题
        if not row_label or row_label == "_checkbox_group":
            section_title = _find_section_title_for_checkbox(table, r_idx)
            if section_title:
                row_label = section_title.replace("\n", "").strip()
                # 去掉编号前缀（如"二、考试方式"→"考试方式"）
                import re as _re
                row_label = _re.sub(r'^[一二三四五六七八九十]+[、.．]\s*', '', row_label)
            else:
                row_label = "选项"
        
        field_id = f"T{t_idx}_R{r_idx}_CHK"
        # 检查是否已存在
        if any(f['field_id'] == field_id for f in label_fields):
            continue
        
        label_fields.append({
            "field_id": field_id,
            "table_idx": t_idx,
            "row_idx": r_idx,
            "col_idx": 0,  # 整行级别，无特定列
            "line_idx": 0,
            "label": row_label,
            "existing_value": "",
            "description": f"请选择{row_label}（可选值: {'/'.join(options)}）",
            "fill_mode": "check",
            "pattern": "checkbox",
        })


def _detect_checkbox_groups(unique_cells):
    """检测一行中的勾选框组，返回标签组列表。"""
    groups = []
    current_label = None
    current_options = {}
    
    for ci in range(len(unique_cells) - 1):
        cell_text = unique_cells[ci].text.strip()
        next_text = unique_cells[ci + 1].text.strip()
        
        normalized = cell_text.replace(" ", "").replace("　", "")
        
        if (normalized in _OPTION_WORDS or cell_text in _OPTION_WORDS) and not next_text:
            opt_key = normalized if normalized in _OPTION_WORDS else cell_text
            current_options[opt_key] = ci + 1
        elif cell_text and normalized not in _OPTION_WORDS and cell_text not in _OPTION_WORDS and not _is_vmerge_continue(unique_cells[ci]):
            if current_options and current_label:
                groups.append({"label": current_label, "option_blanks": current_options})
            elif current_options:
                groups.append({"label": "_checkbox_group", "option_blanks": current_options})
            current_label = cell_text
            current_options = {}
    
    if current_options:
        if current_label:
            groups.append({"label": current_label, "option_blanks": current_options})
        else:
            groups.append({"label": "_checkbox_group", "option_blanks": current_options})
    
    return groups


def _find_section_title_for_checkbox(table, row_idx: int) -> str:
    """从上方行查找章节标题（用于无行标签的勾选框组）。"""
    for r in range(row_idx - 1, max(row_idx - 3, -1), -1):
        unique = _get_unique_cells(table.rows[r])
        if len(unique) == 1:
            text = unique[0].text.strip()
            if text and len(text) >= 2:
                return text
    return ""


def _scan_table_data_sections(table, t_idx: int, label_fields: list):
    """扫描表格中的"表头+数据行"区域，识别多列数据待填字段。
    
    处理的场景：
    1. "标签+空白/占位符"在同行（如 |人数|[空白]|[空白]|...）
    2. "表头行+数据行"配对（如"分数段"行下有"人数"行和"比例"行）
    3. 纯占位符行（如全是%的行，需要替换）
    """
    for r_idx, row in enumerate(table.rows):
        unique = _get_unique_cells(row)
        if len(unique) < 2:
            continue
        if _is_section_title_row(unique):
            continue
        
        # 检测"多列标签+待填"模式
        # 模式A：行内有多个空白格/占位符格，且前面有标签格提供上下文
        # 模式B：行内有"行标签"+"占位符格/空白格"序列（如 |人数|空白|空白|空白|...）
        
        _detect_multi_column_fields(table, t_idx, r_idx, unique, label_fields)


def _detect_multi_column_fields(table, t_idx: int, r_idx: int, unique_cells: list, label_fields: list):
    """检测多列待填字段。
    
    典型场景：
    - 试卷分析行11: |分数分布(续)|人数|空白|空白|空白|空白|空白|
    - 试卷分析行12: |分数分布(续)|比例|%|%|%|%|%|
    - 试卷分析行9: |及百分比|%|%|%|%|%|%|%|%|%|
    """
    # 如果该行已有checkbox字段，跳过multi_col检测
    if any(f.get("pattern") == "checkbox" and f["table_idx"] == t_idx and f["row_idx"] == r_idx for f in label_fields):
        return
    
    # 构建已被label_blank模式覆盖的格集合 (table_idx, row_idx, col_idx)
    covered_cells = set()
    for f in label_fields:
        if f["pattern"] in ("label_blank", "colon"):
            covered_cells.add((f["table_idx"], f["row_idx"], f["col_idx"]))
    
    # 检查此行是否有"行标签+多列待填"模式
    row_label = None
    fillable_cells = []  # (col_idx_in_unique, cell, existing_text)
    
    for ci, cell in enumerate(unique_cells):
        if _is_vmerge_continue(cell):
            continue
        
        text = cell.text.strip()
        
        # 第一个非空非选项格作为行标签
        if row_label is None and text and _is_label_cell(cell):
            row_label = text
            continue
        
        # 跳过已被label_blank模式识别的格
        if (t_idx, r_idx, ci) in covered_cells:
            continue
        
        # 跳过选项文字格（勾选框行的选项不算待填格）
        normalized = text.replace(" ", "").replace("　", "")
        if text in _OPTION_WORDS or normalized in _OPTION_WORDS:
            continue
        
        if not text or _is_placeholder_cell(cell):
            fillable_cells.append((ci, cell, text))
    
    if not fillable_cells or len(fillable_cells) < 2:
        logger.debug(f"multi_col T{t_idx}_R{r_idx}: no fillable cells (count={len(fillable_cells)})")
        return
    
    logger.debug(f"multi_col T{t_idx}_R{r_idx}: row_label={row_label}, fillable_count={len(fillable_cells)}")
    
    # 需要为这些待填格生成有意义的标签
    # 策略1：如果有行标签，用"行标签+列标签"
    # 策略2：看上方行对应的格，获取列标签
    
    if row_label:
        # 清理行标签中的换行符
        row_label = row_label.replace('\n', '').replace('\r', '')
        # 这是一个多列数据行（如人数行、比例行）
        col_labels = _get_column_labels_for_row(table, t_idx, r_idx, unique_cells, fillable_cells)
        
        for idx, (ci, cell, existing) in enumerate(fillable_cells):
            col_label = col_labels[idx] if idx < len(col_labels) else f"第{idx+1}列"
            full_label = f"{row_label}_{col_label}"
            
            # 如果col_label仍是无语义的"第N列"，尝试补充上方表头文本到description
            if col_label.startswith('第') and '列' in col_label:
                # 尝试获取更具体的列标签
                specific_label = _find_specific_col_label(table, r_idx, unique_cells, ci)
                desc = f"请填写{row_label}的{col_label}"
                if specific_label:
                    desc = f"请填写{row_label}的{col_label}(表头:{specific_label})"
            else:
                desc = f"请填写{row_label}的{col_label}"
            
            # 检查是否已存在
            already = any(f["label"] == full_label and f["table_idx"] == t_idx for f in label_fields)
            if already:
                continue
            
            fill_mode = "replace" if existing else "set"
            field_id = f"T{t_idx}_R{r_idx}_C{ci}"
            label_fields.append({
                "field_id": field_id,
                "table_idx": t_idx,
                "row_idx": r_idx,
                "col_idx": ci,
                "line_idx": 0,
                "label": full_label,
                "existing_value": existing if existing else "",
                "description": desc,
                "fill_mode": fill_mode,
                "pattern": "multi_col",
            })
    
    else:
        # 没有行标签，但有多个待填格
        col_labels = _get_column_labels_for_row(table, t_idx, r_idx, unique_cells, fillable_cells)
        above_label = _get_row_label_from_above(table, t_idx, r_idx)
        
        for idx, (ci, cell, existing) in enumerate(fillable_cells):
            col_label = col_labels[idx] if idx < len(col_labels) else f"第{idx+1}列"
            if above_label:
                full_label = f"{above_label}_{col_label}"
            else:
                full_label = col_label
            
            # 如果col_label仍是无语义的"第N列"，尝试补充上方表头文本到description
            if col_label.startswith('第') and '列' in col_label:
                specific_label = _find_specific_col_label(table, r_idx, unique_cells, ci)
                desc = f"请填写{full_label}"
                if specific_label:
                    desc = f"请填写{full_label}(表头:{specific_label})"
            else:
                desc = f"请填写{full_label}"
            
            already = any(f["label"] == full_label and f["table_idx"] == t_idx for f in label_fields)
            if already:
                continue
            
            fill_mode = "replace" if existing else "set"
            field_id = f"T{t_idx}_R{r_idx}_C{ci}"
            label_fields.append({
                "field_id": field_id,
                "table_idx": t_idx,
                "row_idx": r_idx,
                "col_idx": ci,
                "line_idx": 0,
                "label": full_label,
                "existing_value": existing if existing else "",
                "description": desc,
                "fill_mode": fill_mode,
                "pattern": "multi_col",
            })


def _find_specific_col_label(table, r_idx: int, unique_cells: list, ci: int) -> str:
    """当_get_column_labels_for_row返回"第N列"时，尝试更宽泛地查找列标签。
    
    策略：放宽条件，即使文本较长或包含换行，也提取关键词作为上下文。
    """
    # 计算当前行的grid col位置
    grid_cols = _compute_grid_col_positions(unique_cells)
    target_grid_col = grid_cols.get(ci, -1)
    
    for above_r in range(r_idx - 1, max(r_idx - 8, -1), -1):
        above_unique = _get_unique_cells(table.rows[above_r])
        above_header_map = _build_col_header_map(table, above_r, above_unique)
        
        if target_grid_col >= 0 and target_grid_col in above_header_map:
            text = above_header_map[target_grid_col]
            # 放宽长度限制到50，但截断
            if text and text not in _PLACEHOLDER_TEXTS and len(text) > 0:
                # 清理文本：去换行，截断
                text = text.replace('\n', ' ').replace('\r', '').strip()
                if len(text) > 20:
                    text = text[:20] + '...'
                return text
    
    return ""


def _compute_grid_col_positions(unique_cells: list) -> dict:
    """计算每个unique cell在grid中的起始列位置。
    
    Returns:
        {unique_cell_index: grid_col_start}
    """
    from docx.oxml.ns import qn as _qn
    
    positions = {}
    grid_col = 0
    
    for ci, cell in enumerate(unique_cells):
        positions[ci] = grid_col
        tc = cell._tc
        tcPr = tc.find(_qn("w:tcPr"))
        grid_span = 1
        if tcPr is not None:
            gs = tcPr.find(_qn("w:gridSpan"))
            if gs is not None:
                try:
                    grid_span = int(gs.get(_qn("w:val"), "1"))
                except (ValueError, TypeError):
                    grid_span = 1
        grid_col += grid_span
    
    return positions


def _build_col_header_map(table, row_idx: int, unique_cells: list) -> dict:
    """构建列索引到表头文本的映射，支持合并单元格。
    
    当一个unique cell通过gridSpan横跨多列时，该格的文本会映射到
    所有被横跨的列索引上。
    
    Returns:
        {grid_col_idx: header_text}  grid_col_idx是基于gridCol计算的绝对列索引
    """
    from docx.oxml.ns import qn as _qn
    
    col_map = {}
    grid_col = 0  # 当前在grid中的列位置
    
    for ci, cell in enumerate(unique_cells):
        # 获取cell的gridSpan
        tc = cell._tc
        tcPr = tc.find(_qn("w:tcPr"))
        grid_span = 1
        if tcPr is not None:
            gs = tcPr.find(_qn("w:gridSpan"))
            if gs is not None:
                try:
                    grid_span = int(gs.get(_qn("w:val"), "1"))
                except (ValueError, TypeError):
                    grid_span = 1
        
        text = cell.text.strip().replace('\n', '')
        if text:
            # 将文本映射到该cell覆盖的所有grid列
            for g in range(grid_span):
                col_map[grid_col + g] = text
        
        grid_col += grid_span
    
    return col_map


def _get_column_labels_for_row(table, t_idx: int, r_idx: int, unique_cells: list, fillable_cells: list) -> list:
    """获取待填格对应的列标签（从上方表头行获取）
    
    改进：支持合并单元格表头——当上方行的某个格横跨多列时，
    该格的文本作为所有被横跨列的标签。
    """
    from docx.oxml.ns import qn as _qn
    
    col_labels = []
    
    # 计算当前行每个unique cell的gridCol起始位置
    current_grid_cols = _compute_grid_col_positions(unique_cells)
    
    # 向上查找表头行
    for above_r in range(r_idx - 1, max(r_idx - 5, -1), -1):
        above_unique = _get_unique_cells(table.rows[above_r])
        if len(above_unique) < len(fillable_cells):
            continue
        
        # 构建上方行的列映射：grid_col → header_text
        col_header_map = _build_col_header_map(table, above_r, above_unique)
        
        for idx, (ci, cell, _) in enumerate(fillable_cells):
            # 获取该fillable cell的grid列位置
            grid_col = current_grid_cols.get(ci, ci)
            
            # 先尝试从col_header_map精确匹配
            header_text = col_header_map.get(grid_col)
            if header_text and header_text not in _PLACEHOLDER_TEXTS and len(header_text) <= 30:
                col_labels.append(header_text)
            elif ci < len(above_unique):
                # 回退：直接读上方行同位置格
                above_text = above_unique[ci].text.strip().replace('\n', '')
                if above_text and above_text not in _PLACEHOLDER_TEXTS and len(above_text) <= 30:
                    col_labels.append(above_text)
                else:
                    col_labels.append(f"第{idx+1}列")
            else:
                col_labels.append(f"第{idx+1}列")
        
        # 检查是否获取到有意义的标签
        meaningful = sum(1 for l in col_labels if not l.startswith('第'))
        if meaningful > 0:
            return col_labels
    
    # 没找到，用序号
    return [f"第{idx+1}列" for idx in range(len(fillable_cells))]


def _get_row_label_from_above(table, t_idx: int, r_idx: int) -> Optional[str]:
    """从上方行获取行标签"""
    for above_r in range(r_idx - 1, max(r_idx - 3, -1), -1):
        above_unique = _get_unique_cells(table.rows[above_r])
        if above_unique:
            first_text = above_unique[0].text.strip()
            if first_text and _is_label_cell(above_unique[0]):
                return first_text
    return None


if __name__ == "__main__":
    template_path = os.path.join(
        os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects"),
        "assets", "template_exam_lingnan.docx"
    )
    result = analyze_template(template_path)
    print(json.dumps(result, ensure_ascii=False, indent=2))
