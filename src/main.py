import argparse
import asyncio
import json
import os
import sys
import threading
import traceback
import logging

# 确保 src/ 在 sys.path 中，让 coze_coding_utils 内部的 from utils.helper import graph_helper 能找到
_src_dir = os.path.dirname(os.path.abspath(__file__))
if _src_dir not in sys.path:
    sys.path.insert(0, _src_dir)
from typing import Any, Dict, Iterable, AsyncIterable, AsyncGenerator, Optional
import cozeloop
import uvicorn
import time
from fastapi import FastAPI, HTTPException, Request, UploadFile, File
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from langchain_core.runnables import RunnableConfig
from langgraph.graph import StateGraph, END
from langgraph.graph.state import CompiledStateGraph
from coze_coding_utils.runtime_ctx.context import new_context, Context
from coze_coding_utils.helper import graph_helper
from coze_coding_utils.log.node_log import LOG_FILE
from coze_coding_utils.log.write_log import setup_logging, request_context
from coze_coding_utils.log.config import LOG_LEVEL
from coze_coding_utils.error.classifier import ErrorClassifier, classify_error
from coze_coding_utils.helper.stream_runner import AgentStreamRunner, WorkflowStreamRunner,agent_stream_handler,workflow_stream_handler, RunOpt

setup_logging(
    log_file=LOG_FILE,
    max_bytes=100 * 1024 * 1024, # 100MB
    backup_count=5,
    log_level=LOG_LEVEL,
    use_json_format=True,
    console_output=True
)

logger = logging.getLogger(__name__)
from coze_coding_utils.helper.agent_helper import to_stream_input
from coze_coding_utils.openai.handler import OpenAIChatHandler
from coze_coding_utils.log.parser import LangGraphParser
from coze_coding_utils.log.err_trace import extract_core_stack
from coze_coding_utils.log.loop_trace import init_run_config, init_agent_config


# 超时配置常量
TIMEOUT_SECONDS = 900  # 15分钟

class GraphService:
    def __init__(self):
        # 用于跟踪正在运行的任务（使用asyncio.Task）
        self.running_tasks: Dict[str, asyncio.Task] = {}
        # 错误分类器
        self.error_classifier = ErrorClassifier()
        # stream runner
        self._agent_stream_runner = AgentStreamRunner()
        self._workflow_stream_runner = WorkflowStreamRunner()
        self._graph = None
        self._graph_lock = threading.Lock()

    def _get_graph(self, ctx=Context):
        if graph_helper.is_agent_proj():
            return graph_helper.get_agent_instance("agents.agent", ctx)

        if self._graph is not None:
            return self._graph
        with self._graph_lock:
            if self._graph is not None:
                return self._graph
            self._graph = graph_helper.get_graph_instance("graphs.graph")
            return self._graph

    @staticmethod
    def _sse_event(data: Any, event_id: Any = None) -> str:
        id_line = f"id: {event_id}\n" if event_id else ""
        return f"{id_line}event: message\ndata: {json.dumps(data, ensure_ascii=False, default=str)}\n\n"

    def _get_stream_runner(self):
        if graph_helper.is_agent_proj():
            return self._agent_stream_runner
        else:
            return self._workflow_stream_runner

    # 流式运行（原始迭代器）：本地调用使用
    def stream(self, payload: Dict[str, Any], run_config: RunnableConfig, ctx=Context) -> Iterable[Any]:
        graph = self._get_graph(ctx)
        stream_runner = self._get_stream_runner()
        for chunk in stream_runner.stream(payload, graph, run_config, ctx):
            yield chunk

    # 同步运行：本地/HTTP 通用
    async def run(self, payload: Dict[str, Any], ctx=None) -> Dict[str, Any]:
        if ctx is None:
            ctx = new_context("run")

        run_id = ctx.run_id
        logger.info(f"Starting run with run_id: {run_id}")

        try:
            graph = self._get_graph(ctx)
            # custom tracer
            run_config = init_run_config(graph, ctx)
            run_config["configurable"] = {"thread_id": ctx.run_id}

            # 直接调用，LangGraph会在当前任务上下文中执行
            # 如果当前任务被取消，LangGraph的执行也会被取消
            return await graph.ainvoke(payload, config=run_config, context=ctx)

        except asyncio.CancelledError:
            logger.info(f"Run {run_id} was cancelled")
            return {"status": "cancelled", "run_id": run_id, "message": "Execution was cancelled"}
        except Exception as e:
            # 使用错误分类器分类错误
            err = self.error_classifier.classify(e, {"node_name": "run", "run_id": run_id})
            # 记录详细的错误信息和堆栈跟踪
            logger.error(
                f"Error in GraphService.run: [{err.code}] {err.message}\n"
                f"Category: {err.category.name}\n"
                f"Traceback:\n{extract_core_stack()}"
            )
            # 保留原始异常堆栈，便于上层返回真正的报错位置
            raise
        finally:
            # 清理任务记录
            self.running_tasks.pop(run_id, None)

    # 流式运行（SSE 格式化）：HTTP 路由使用
    async def stream_sse(self, payload: Dict[str, Any], ctx=None, run_opt: Optional[RunOpt] = None) -> AsyncGenerator[str, None]:
        if ctx is None:
            ctx = new_context(method="stream_sse")
        if run_opt is None:
            run_opt = RunOpt()

        run_id = ctx.run_id
        logger.info(f"Starting stream with run_id: {run_id}")
        graph = self._get_graph(ctx)
        if graph_helper.is_agent_proj():
            run_config = init_agent_config(graph, ctx)
        else:
            run_config = init_run_config(graph, ctx)  # vibeflow

        is_workflow = not graph_helper.is_agent_proj()

        try:
            async for chunk in self.astream(payload, graph, run_config=run_config, ctx=ctx, run_opt=run_opt):
                if is_workflow and isinstance(chunk, tuple):
                    event_id, data = chunk
                    yield self._sse_event(data, event_id)
                else:
                    yield self._sse_event(chunk)
        finally:
            # 清理任务记录
            self.running_tasks.pop(run_id, None)
            cozeloop.flush()

    # 取消执行 - 使用asyncio的标准方式
    def cancel_run(self, run_id: str, ctx: Optional[Context] = None) -> Dict[str, Any]:
        """
        取消指定run_id的执行

        使用asyncio.Task.cancel()来取消任务,这是标准的Python异步取消机制。
        LangGraph会在节点之间检查CancelledError,实现优雅的取消。
        """
        logger.info(f"Attempting to cancel run_id: {run_id}")

        # 查找对应的任务
        if run_id in self.running_tasks:
            task = self.running_tasks[run_id]
            if not task.done():
                # 使用asyncio的标准取消机制
                # 这会在下一个await点抛出CancelledError
                task.cancel()
                logger.info(f"Cancellation requested for run_id: {run_id}")
                return {
                    "status": "success",
                    "run_id": run_id,
                    "message": "Cancellation signal sent, task will be cancelled at next await point"
                }
            else:
                logger.info(f"Task already completed for run_id: {run_id}")
                return {
                    "status": "already_completed",
                    "run_id": run_id,
                    "message": "Task has already completed"
                }
        else:
            logger.warning(f"No active task found for run_id: {run_id}")
            return {
                "status": "not_found",
                "run_id": run_id,
                "message": "No active task found with this run_id. Task may have already completed or run_id is invalid."
            }

    # 运行指定节点：本地/HTTP 通用
    async def run_node(self, node_id: str, payload: Dict[str, Any], ctx=None) -> Any:
        if ctx is None or Context.run_id == "":
            ctx = new_context(method="node_run")

        _graph = self._get_graph()
        node_func, input_cls, output_cls = graph_helper.get_graph_node_func_with_inout(_graph.get_graph(), node_id)
        if node_func is None or input_cls is None:
            raise KeyError(f"node_id '{node_id}' not found")

        parser = LangGraphParser(_graph)
        metadata = parser.get_node_metadata(node_id) or {}

        _g = StateGraph(input_cls, input_schema=input_cls, output_schema=output_cls)
        _g.add_node("sn", node_func, metadata=metadata)
        _g.set_entry_point("sn")
        _g.add_edge("sn", END)
        _graph = _g.compile()

        run_config = init_run_config(_graph, ctx)
        return await _graph.ainvoke(payload, config=run_config)

    def graph_inout_schema(self) -> Any:
        if graph_helper.is_agent_proj():
            return {"input_schema": {}, "output_schema": {}}
        builder = getattr(self._get_graph(), 'builder', None)
        if builder is not None:
            input_cls = getattr(builder, 'input_schema', None) or self.graph.get_input_schema()
            output_cls = getattr(builder, 'output_schema', None) or self.graph.get_output_schema()
        else:
            logger.warning(f"No builder input schema found for graph_inout_schema, using graph input schema instead")
            input_cls = self.graph.get_input_schema()
            output_cls = self.graph.get_output_schema()

        return {
            "input_schema": input_cls.model_json_schema(), 
            "output_schema": output_cls.model_json_schema(),
            "code":0,
            "msg":""
        }

    async def astream(self, payload: Dict[str, Any], graph: CompiledStateGraph, run_config: RunnableConfig, ctx=Context, run_opt: Optional[RunOpt] = None) -> AsyncIterable[Any]:
        stream_runner = self._get_stream_runner()
        async for chunk in stream_runner.astream(payload, graph, run_config, ctx, run_opt):
            yield chunk


service = GraphService()
app = FastAPI()

# CORS 中间件 - 允许前端页面访问 API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 挂载前端静态文件
_web_dir = os.path.join(os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects"), "web")
if os.path.isdir(_web_dir):
    app.mount("/web", StaticFiles(directory=_web_dir, html=True), name="web")
    logger.info(f"Frontend mounted at /web from {_web_dir}")

# 根路径重定向到 /web/
from starlette.responses import RedirectResponse
@app.get("/", include_in_schema=False)
async def root_redirect():
    return RedirectResponse(url="/web/")

# OpenAI 兼容接口处理器
openai_handler = OpenAIChatHandler(service)


HEADER_X_RUN_ID = "x-run-id"
@app.post("/run")
async def http_run(request: Request) -> Dict[str, Any]:
    global result
    raw_body = await request.body()
    try:
        body_text = raw_body.decode("utf-8")
    except Exception as e:
        body_text = str(raw_body)
        raise HTTPException(status_code=400,
                            detail=f"Invalid JSON format: {body_text}, traceback: {traceback.format_exc()}, error: {e}")

    ctx = new_context(method="run", headers=request.headers)
    # 优先使用上游指定的 run_id，保证 cancel 能精确匹配
    upstream_run_id = request.headers.get(HEADER_X_RUN_ID)
    if upstream_run_id:
        ctx.run_id = upstream_run_id
    run_id = ctx.run_id
    request_context.set(ctx)

    logger.info(
        f"Received request for /run: "
        f"run_id={run_id}, "
        f"query={dict(request.query_params)}, "
        f"body={body_text}"
    )

    try:
        payload = await request.json()

        # 创建任务并记录 - 这是关键，让我们可以通过run_id取消任务
        task = asyncio.create_task(service.run(payload, ctx))
        service.running_tasks[run_id] = task

        try:
            result = await asyncio.wait_for(task, timeout=float(TIMEOUT_SECONDS))
        except asyncio.TimeoutError:
            logger.error(f"Run execution timeout after {TIMEOUT_SECONDS}s for run_id: {run_id}")
            task.cancel()
            try:
                result = await task
            except asyncio.CancelledError:
                return {
                    "status": "timeout",
                    "run_id": run_id,
                    "message": f"Execution timeout: exceeded {TIMEOUT_SECONDS} seconds"
                }

        if not result:
            result = {}
        if isinstance(result, dict):
            result["run_id"] = run_id
        return result

    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error in http_run: {e}, traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=400, detail=f"Invalid JSON format, {extract_core_stack()}")

    except asyncio.CancelledError:
        logger.info(f"Request cancelled for run_id: {run_id}")
        result = {"status": "cancelled", "run_id": run_id, "message": "Execution was cancelled"}
        return result

    except Exception as e:
        # 使用错误分类器获取错误信息
        error_response = service.error_classifier.get_error_response(e, {"node_name": "http_run", "run_id": run_id})
        logger.error(
            f"Unexpected error in http_run: [{error_response['error_code']}] {error_response['error_message']}, "
            f"traceback: {traceback.format_exc()}", exc_info=True
        )
        raise HTTPException(
            status_code=500,
            detail={
                "error_code": error_response["error_code"],
                "error_message": error_response["error_message"],
                "stack_trace": extract_core_stack(),
            }
        )
    finally:
        cozeloop.flush()


HEADER_X_WORKFLOW_STREAM_MODE = "x-workflow-stream-mode"


def _register_task(run_id: str, task: asyncio.Task):
    service.running_tasks[run_id] = task


@app.post("/stream_run")
async def http_stream_run(request: Request):
    ctx = new_context(method="stream_run", headers=request.headers)
    # 优先使用上游指定的 run_id，保证 cancel 能精确匹配
    upstream_run_id = request.headers.get(HEADER_X_RUN_ID)
    if upstream_run_id:
        ctx.run_id = upstream_run_id
    workflow_stream_mode = request.headers.get(HEADER_X_WORKFLOW_STREAM_MODE, "").lower()
    workflow_debug = workflow_stream_mode == "debug"
    request_context.set(ctx)
    raw_body = await request.body()
    try:
        body_text = raw_body.decode("utf-8")
    except Exception as e:
        body_text = str(raw_body)
        raise HTTPException(status_code=400,
                            detail=f"Invalid JSON format: {body_text}, traceback: {extract_core_stack()}, error: {e}")
    run_id = ctx.run_id
    is_agent = graph_helper.is_agent_proj()
    logger.info(
        f"Received request for /stream_run: "
        f"run_id={run_id}, "
        f"is_agent_project={is_agent}, "
        f"query={dict(request.query_params)}, "
        f"body={body_text}"
    )
    try:
        payload = await request.json()
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error in http_stream_run: {e}, traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=400, detail=f"Invalid JSON format:{extract_core_stack()}")

    if is_agent:
        stream_generator = agent_stream_handler(
            payload=payload,
            ctx=ctx,
            run_id=run_id,
            stream_sse_func=service.stream_sse,
            sse_event_func=service._sse_event,
            error_classifier=service.error_classifier,
            register_task_func=_register_task,
        )
    else:
        stream_generator = workflow_stream_handler(
            payload=payload,
            ctx=ctx,
            run_id=run_id,
            stream_sse_func=service.stream_sse,
            sse_event_func=service._sse_event,
            error_classifier=service.error_classifier,
            register_task_func=_register_task,
            run_opt=RunOpt(workflow_debug=workflow_debug),
        )

    response = StreamingResponse(stream_generator, media_type="text/event-stream")
    return response

@app.post("/cancel/{run_id}")
async def http_cancel(run_id: str, request: Request):
    """
    取消指定run_id的执行

    使用asyncio.Task.cancel()实现取消,这是Python标准的异步任务取消机制。
    LangGraph会在节点之间的await点检查CancelledError,实现优雅取消。
    """
    ctx = new_context(method="cancel", headers=request.headers)
    request_context.set(ctx)
    logger.info(f"Received cancel request for run_id: {run_id}")
    result = service.cancel_run(run_id, ctx)
    return result


@app.post(path="/node_run/{node_id}")
async def http_node_run(node_id: str, request: Request):
    raw_body = await request.body()
    try:
        body_text = raw_body.decode("utf-8")
    except UnicodeDecodeError:
        body_text = str(raw_body)
        raise HTTPException(status_code=400, detail=f"Invalid JSON format: {body_text}")
    ctx = new_context(method="node_run", headers=request.headers)
    request_context.set(ctx)
    logger.info(
        f"Received request for /node_run/{node_id}: "
        f"query={dict(request.query_params)}, "
        f"body={body_text}",
    )

    try:
        payload = await request.json()
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error in http_node_run: {e}, traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=400, detail=f"Invalid JSON format:{extract_core_stack()}")
    try:
        return await service.run_node(node_id, payload, ctx)
    except KeyError:
        raise HTTPException(status_code=404,
                            detail=f"node_id '{node_id}' not found or input miss required fields, traceback: {extract_core_stack()}")
    except Exception as e:
        # 使用错误分类器获取错误信息
        error_response = service.error_classifier.get_error_response(e, {"node_name": node_id})
        logger.error(
            f"Unexpected error in http_node_run: [{error_response['error_code']}] {error_response['error_message']}, "
            f"traceback: {traceback.format_exc()}", exc_info=True
        )
        raise HTTPException(
            status_code=500,
            detail={
                "error_code": error_response["error_code"],
                "error_message": error_response["error_message"],
                "stack_trace": extract_core_stack(),
            }
        )
    finally:
        cozeloop.flush()


@app.post("/v1/chat/completions")
async def openai_chat_completions(request: Request):
    """OpenAI Chat Completions API 兼容接口"""
    ctx = new_context(method="openai_chat", headers=request.headers)
    request_context.set(ctx)

    logger.info(f"Received request for /v1/chat/completions: run_id={ctx.run_id}")

    try:
        payload = await request.json()
        return await openai_handler.handle(payload, ctx)
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error in openai_chat_completions: {e}")
        raise HTTPException(status_code=400, detail="Invalid JSON format")
    finally:
        cozeloop.flush()


@app.get("/health")
async def health_check():
    try:
        # 这里可以添加更多的健康检查逻辑
        return {
            "status": "ok",
            "message": "Service is running",
        }
    except Exception as e:
        raise HTTPException(status_code=503, detail=str(e))


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """接收前端上传的文件，解析文本内容后返回"""
    import tempfile
    from tools.knowledge_tool import extract_text_from_upload

    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    # 文件大小限制 10MB
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 10MB)")

    # 保存到临时文件
    suffix = os.path.splitext(file.filename)[1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        result = extract_text_from_upload(tmp_path)
        result["filename"] = file.filename
        return JSONResponse(content=result)
    finally:
        os.unlink(tmp_path)


@app.post("/upload-template")
async def upload_template_file(file: UploadFile = File(...)):
    """接收前端上传的docx模板文件，保存并提取文本供Agent分析"""
    import tempfile
    from docx import Document

    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    if not file.filename.lower().endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Only .docx/.doc files are supported")

    # 文件大小限制 10MB
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 10MB)")

    # 保存到临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    # 提取文档文本内容（包括表格）
    extracted_text = ""
    try:
        doc = Document(tmp_path)
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text += para.text.strip() + "\n"
        # 提取表格
        for t_idx, table in enumerate(doc.tables):
            extracted_text += f"\n--- 表格{t_idx+1} ---\n"
            for row in table.rows:
                cells_text = []
                for cell in row.cells:
                    cells_text.append(cell.text.strip())
                extracted_text += " | ".join(cells_text) + "\n"
    except Exception as e:
        logger.warning(f"Failed to extract text from template: {e}")
        extracted_text = ""

    return JSONResponse(content={
        "success": True,
        "filename": file.filename,
        "template_path": tmp_path,
        "extracted_text": extracted_text,
        "message": f"模板文件已上传：{file.filename}，请告诉Agent分析此模板"
    })


@app.get("/template-preview")
async def template_preview(path: str = ""):
    """将 docx 模板转为 HTML 预览，返回 HTML + 字段映射"""
    if not path:
        return JSONResponse(content={"success": False, "message": "缺少 path 参数"})

    # 支持 path 是模板名（评价报告/试卷分析/关联矩阵）或文件路径
    workspace = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")
    full_path = path

    # 内置模板名
    template_map = {
        "评价报告": "2023-2024-2《xxx》 岭南师范学院专业课程目标达成度评价报告模板.docx",
        "试卷分析": "2023-2024-2《xxx》 试卷分析模板.docx",
        "关联矩阵": "2023-2024-2《xxx》岭南师范学院考题与课程目标及毕业要求关联矩阵表模板.docx",
    }
    if path in template_map:
        full_path = os.path.join(workspace, "assets", template_map[path])
    elif not os.path.isabs(path):
        full_path = os.path.join(workspace, path)

    if not os.path.exists(full_path):
        return JSONResponse(content={"success": False, "message": f"文件不存在: {path}"})

    try:
        from tools.docx_preview import docx_to_html
        result = docx_to_html(full_path)
        return JSONResponse(content={
            "success": True,
            "html": result["html"],
            "field_map": result["field_map"],
            "label_to_fields": result["label_to_fields"],
        })
    except Exception as e:
        logger.exception("template-preview failed")
        return JSONResponse(content={"success": False, "message": f"预览生成失败: {e}"})


@app.get("/generated-preview")
async def generated_preview(local_path: str = ""):
    """将生成的 docx 文件转为 HTML 预览（含填写后的内容）"""
    if not local_path:
        return JSONResponse(content={"success": False, "message": "缺少 local_path 参数"})

    if not os.path.exists(local_path):
        return JSONResponse(content={"success": False, "message": f"文件不存在: {local_path}"})

    try:
        from tools.docx_preview import docx_to_html
        result = docx_to_html(local_path)
        return JSONResponse(content={
            "success": True,
            "html": result["html"],
            "field_map": result["field_map"],
            "label_to_fields": result["label_to_fields"],
        })
    except Exception as e:
        logger.exception("generated-preview failed")
        return JSONResponse(content={"success": False, "message": f"预览生成失败: {e}"})


@app.get(path="/download-docx")
async def http_download_docx(file_path: str = ""):
    """代理下载 docx 文件（支持本地路径和远程URL）"""
    if not file_path:
        return JSONResponse(content={"success": False, "message": "缺少 file_path"})

    # 如果是远程URL（对象存储），先下载到 /tmp 再返回
    if file_path.startswith("http://") or file_path.startswith("https://"):
        import httpx
        try:
            async with httpx.AsyncClient(timeout=60, follow_redirects=True) as client:
                resp = await client.get(file_path)
                resp.raise_for_status()
                # 从URL提取文件名
                from urllib.parse import urlparse, unquote
                url_path = unquote(urlparse(file_path).path)
                fname = os.path.basename(url_path).split("?")[0]
                if not fname.endswith(".docx"):
                    fname = "document.docx"
                tmp_path = os.path.join("/tmp", fname)
                with open(tmp_path, "wb") as f:
                    f.write(resp.content)
                return FileResponse(tmp_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=fname)
        except Exception as e:
            return JSONResponse(content={"success": False, "message": f"下载失败: {e}"})

    # 本地文件路径
    full_path = file_path
    if not os.path.isabs(file_path):
        from tools.docx_preview import _resolve_template_path
        full_path = _resolve_template_path(file_path)
    if not os.path.isfile(full_path):
        return JSONResponse(content={"success": False, "message": f"文件不存在: {file_path}"})
    fname = os.path.basename(full_path)
    return FileResponse(full_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=fname)


@app.get(path="/convert-pdf")
async def http_convert_pdf(file_path: str = ""):
    """将 docx 转为 PDF（用 LibreOffice），保留原格式"""
    import subprocess
    if not file_path:
        return JSONResponse(content={"success": False, "message": "缺少 file_path"})
    # 支持模板名解析（如"试卷分析"）
    from tools.docx_preview import _resolve_template_path
    full_path = _resolve_template_path(file_path)
    if not os.path.isfile(full_path):
        return JSONResponse(content={"success": False, "message": f"文件不存在: {file_path}"})
    try:
        out_dir = "/tmp/pdf_output"
        os.makedirs(out_dir, exist_ok=True)
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, full_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            return JSONResponse(content={"success": False, "message": f"转换失败: {result.stderr}"})
        pdf_name = os.path.splitext(os.path.basename(full_path))[0] + ".pdf"
        pdf_path = os.path.join(out_dir, pdf_name)
        if not os.path.isfile(pdf_path):
            return JSONResponse(content={"success": False, "message": "PDF 文件未生成"})
        return FileResponse(pdf_path, media_type="application/pdf", filename=pdf_name)
    except Exception as e:
        return JSONResponse(content={"success": False, "message": f"转换出错: {e}"})


@app.get(path="/graph_parameter")
async def http_graph_inout_parameter(request: Request):
    return service.graph_inout_schema()

def parse_args():
    parser = argparse.ArgumentParser(description="Start FastAPI server")
    parser.add_argument("-m", type=str, default="http", help="Run mode, support http,flow,node")
    parser.add_argument("-n", type=str, default="", help="Node ID for single node run")
    parser.add_argument("-p", type=int, default=5000, help="HTTP server port")
    parser.add_argument("-i", type=str, default="", help="Input JSON string for flow/node mode")
    return parser.parse_args()


def parse_input(input_str: str) -> Dict[str, Any]:
    """Parse input string, support both JSON string and plain text"""
    if not input_str:
        return {"text": "你好"}

    # Try to parse as JSON first
    try:
        return json.loads(input_str)
    except json.JSONDecodeError:
        # If not valid JSON, treat as plain text
        return {"text": input_str}

def start_http_server(port):
    workers = 1
    reload = False
    if graph_helper.is_dev_env():
        reload = True

    logger.info(f"Start HTTP Server, Port: {port}, Workers: {workers}")
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=reload, workers=workers)

if __name__ == "__main__":
    args = parse_args()
    if args.m == "http":
        start_http_server(args.p)
    elif args.m == "flow":
        payload = parse_input(args.i)
        result = asyncio.run(service.run(payload))
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif args.m == "node" and args.n:
        payload = parse_input(args.i)
        result = asyncio.run(service.run_node(args.n, payload))
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif args.m == "agent":
        agent_ctx = new_context(method="agent")
        for chunk in service.stream(
                {
                    "type": "query",
                    "session_id": "1",
                    "message": "你好",
                    "content": {
                        "query": {
                            "prompt": [
                                {
                                    "type": "text",
                                    "content": {"text": "现在几点了？请调用工具获取当前时间"},
                                }
                            ]
                        }
                    },
                },
                run_config={"configurable": {"session_id": "1"}},
                ctx=agent_ctx,
        ):
            print(chunk)
