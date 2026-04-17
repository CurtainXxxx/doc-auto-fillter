"""深入分析模板中所有可填充字段的模式，为自动识别方案提供依据"""
from docx import Document
from lxml import etree
import re

doc = Document("assets/template.docx")
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def get_unique_cells(row):
    """获取行中的独立单元格（跳过被合并重复的）"""
    seen = set()
    cells = []
    for cell in row.cells:
        cid = id(cell._element)
        if cid not in seen:
            seen.add(cid)
            cells.append(cell)
    return cells

# ========== 分析所有单元格的标签模式 ==========
print("=" * 100)
print("一、所有含文本的单元格，标注: 是否含冒号、是否含待填空白、是否有值")
print("=" * 100)

for t_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格{t_idx} ({len(table.rows)}行 x {len(table.columns)}列) ---")
    for r_idx, row in enumerate(table.rows):
        unique = get_unique_cells(row)
        for c_idx, cell in enumerate(unique):
            text = cell.text.strip().replace('\n', '|')
            if not text:
                # 空单元格 - 可能是待填字段
                # 检查是否是数据行（同行有标签的）
                row_has_label = any('：' in c.text or ':' in c.text for c in unique)
                print(f"  T{t_idx}.R{r_idx}.C{c_idx}: [空] row_has_label={row_has_label}")
                continue
            
            # 分析标签模式
            has_colon = '：' in text or ':' in text
            parts = [p.strip() for p in re.split(r'[：:|]', text) if p.strip()]
            
            # 判断是否有待填部分
            # 含冒号的文本：冒号后如果为空，则待填
            fillable = False
            label = ""
            value_hint = ""
            if has_colon:
                for line in text.split('|'):
                    if '：' in line:
                        before, after = line.split('：', 1)
                        if not after.strip():
                            fillable = True
                            label = before.strip()
                        else:
                            label = before.strip()
                            value_hint = after.strip()
                    elif ':' in line:
                        before, after = line.split(':', 1)
                        if not after.strip():
                            fillable = True
                            label = before.strip()
                        else:
                            label = before.strip()
                            value_hint = after.strip()
            
            mark = "⬜待填" if fillable else ("📝有值" if value_hint else "📌标签")
            print(f"  T{t_idx}.R{r_idx}.C{c_idx}: {mark} | label='{label}' | value='{value_hint}' | full='{text[:60]}'")

# ========== 分析重复行模式 ==========
print("\n" + "=" * 100)
print("二、重复行模式检测（找出结构相同但数据不同的连续行）")
print("=" * 100)

for t_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格{t_idx} ---")
    # 对比相邻行的结构相似度
    for r_idx in range(1, len(table.rows)):
        prev_unique = get_unique_cells(table.rows[r_idx - 1])
        curr_unique = get_unique_cells(table.rows[r_idx])
        
        # 比较独立单元格数量
        if len(prev_unique) == len(curr_unique) and len(prev_unique) > 1:
            # 比较合并模式
            prev_gs = []
            curr_gs = []
            for cell in prev_unique:
                tcPr = cell._element.find('w:tcPr', nsmap)
                if tcPr is not None:
                    gs = tcPr.find('w:gridSpan', nsmap)
                    prev_gs.append(gs.get(nsmap['w']+'val') if gs is not None else '1')
            for cell in curr_unique:
                tcPr = cell._element.find('w:tcPr', nsmap)
                if tcPr is not None:
                    gs = tcPr.find('w:gridSpan', nsmap)
                    curr_gs.append(gs.get(nsmap['w']+'val') if gs is not None else '1')
            
            if prev_gs == curr_gs and prev_gs:
                prev_texts = [c.text.strip()[:20] for c in prev_unique]
                curr_texts = [c.text.strip()[:20] for c in curr_unique]
                # 检查是否都是空的或者含数字递增
                both_mostly_empty = sum(1 for t in prev_texts if not t) >= len(prev_texts)//2 and \
                                    sum(1 for t in curr_texts if not t) >= len(curr_texts)//2
                if both_mostly_empty:
                    print(f"  R{r_idx-1}-R{r_idx}: 相似空行 | gs={prev_gs} | prev={prev_texts} | curr={curr_texts}")

# ========== 分析含"课程目标N"的模式 ==========
print("\n" + "=" * 100)
print("三、含'课程目标'的模式")
print("=" * 100)
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if '课程目标' in cell.text:
                text = cell.text.strip().replace('\n', '|')[:80]
                print(f"  T{t_idx}.R{r_idx}: '{text}'")
