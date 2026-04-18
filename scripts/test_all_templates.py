"""自动化测试：三个模板的填空能力验证"""
import sys, json
sys.path.insert(0, '/workspace/projects/src')

from tools.template_analyzer import analyze_template
from tools.edu_report_tool import _build_report_docx, _get_template_path, _MAX_SIMPLE_GROUP_COLS
from docx import Document

def test_template(name, data):
    print(f"\n{'='*70}")
    print(f"测试模板: {name}")
    print(f"{'='*70}")
    
    # 1. 解析模板
    path = _get_template_path(name)
    analysis = analyze_template(path)
    
    labels = [f['label'] for f in analysis['label_fields']]
    repeats = {f['label']: f['repeat_count'] for f in analysis['label_fields'] if f['repeat_count'] > 1}
    simple_groups = [g for g in analysis['row_groups'] if len(g['column_labels']) <= _MAX_SIMPLE_GROUP_COLS]
    complex_groups = [g for g in analysis['row_groups'] if len(g['column_labels']) > _MAX_SIMPLE_GROUP_COLS]
    
    print(f"标签字段({len(labels)}个): {labels}")
    if repeats:
        print(f"重复标签: {repeats}")
    print(f"简单行组({len(simple_groups)}个): {[g['group_id'] for g in simple_groups]}")
    print(f"复杂行组({len(complex_groups)}个): {[g['group_id'] for g in complex_groups]} — 跳过")
    
    # 2. 生成文档
    try:
        doc_bytes = _build_report_docx(path, data)
        print(f"✅ 文档生成成功，大小: {len(doc_bytes)} bytes")
    except Exception as e:
        print(f"❌ 文档生成失败: {e}")
        import traceback
        traceback.print_exc()
        return
    
    # 3. 验证填充结果
    import io
    doc = Document(io.BytesIO(doc_bytes))
    
    # 检查标签字段
    print(f"\n--- 标签字段验证 ---")
    for label in labels:
        value = data.get(label, "")
        if isinstance(value, list):
            value = value[0] if value else ""
        if value:
            # 在文档中搜索该值
            found = False
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if str(value) in cell.text:
                            found = True
                            break
                    if found:
                        break
                if found:
                    break
            status = "✅" if found else "❌"
            print(f"  {status} {label}: '{value}' {'已填入' if found else '未找到'}")
    
    # 检查行组
    print(f"\n--- 行组验证 ---")
    for g in simple_groups:
        gid = g['group_id']
        gdata = data.get(gid, [])
        if not gdata:
            print(f"  ⏭️ {gid}: 无数据，跳过")
            continue
        
        t_idx = g['table_idx']
        table = doc.tables[t_idx]
        start_row = g['start_row']
        
        # 检查第一行数据是否填入
        if start_row < len(table.rows):
            first_data = gdata[0]
            row = table.rows[start_row]
            row_text = " | ".join(c.text.strip()[:20] for c in row.cells if c.text.strip())
            
            found_any = any(str(v) in row_text for v in first_data if v)
            status = "✅" if found_any else "❌"
            print(f"  {status} {gid}: 数据行首行='{row_text[:80]}'")
    
    # 保存到临时文件供人工检查
    out_path = f"/tmp/test_{name}.docx"
    with open(out_path, "wb") as f:
        f.write(doc_bytes)
    print(f"\n📄 已保存到: {out_path}")


# ========== 测试数据 ==========

# 模板1：评价报告
evaluation_data = {
    "课程名称": "《软件工程导论》",
    "开课时间": "2023-2024-2",
    "考试类别": "闭卷考试",
    "平时": "30%",
    "期末": "70%",
    "参评人数": "45",
    "教学班级": "软件工程24级1班",
    "评价责任人": "张明",
    "参与人": "李华、王强、赵丽",
    "课程总结": "本学期课程教学总体达到预期目标",
    "改进措施": "1.增加软件测试实验课时 2.引入更多项目案例教学",
    # 简单行组 T0_G0
    "T0_G0": [
        ["工程知识", "1.3", "课程目标1：掌握软件工程基本概念和方法"],
        ["问题分析", "2.1", "课程目标2：能够运用面向对象分析方法"],
        ["设计开发解决方案", "3.2", "课程目标3：具备软件测试基本能力"],
        ["使用现代工具", "4.1", "课程目标4：理解软件项目管理"],
    ],
    # 简单行组 T1_G1
    "T1_G1": [
        ["期末考试", "0.7", "0.6", "0.5", "0.4"],
        ["平时成绩", "0.2", "0.2", "0.3", "0.3"],
        ["实验成绩", "0.1", "0.2", "0.2", "0.3"],
    ],
    # T1_G2 是复杂行组(26列)，跳过
    # T1_G3 是复杂行组(35列)，跳过
    # T1_G4 是简单行组
    "T1_G4": [
        ["课程目标1", "0.89"],
        ["课程目标2", "0.91"],
    ],
}

# 模板2：试卷分析
exam_data = {
    "任课教师签字": "张明",
    "系主任签字": "王强",
    "注": ["注1", "注2", "注3"],
    # 行组 T0_G0: 学期信息
    "T0_G0": [
        ["2023-2024学年第2学期"],
    ],
    # 行组 T0_G1: 班级考勤
    "T0_G1": [
        ["软件工程24级1班", "", "45", "43", "2", "0", "0", "0"],
    ],
    # 行组 T0_G2: 分数分布
    "T0_G2": [
        ["", "分数段", "3", "8", "15", "12", "5"],
    ],
}

# 模板3：关联矩阵
matrix_data = {
    # 行组 T0_G0: 毕业要求与课程目标关联
    "T0_G0": [
        ["工程知识", "1.3", "课程目标1(强支撑)"],
        ["问题分析", "2.1", "课程目标2(强支撑)"],
        ["设计开发解决方案", "3.2", "课程目标3(较强支撑)"],
        ["使用现代工具", "4.1", "课程目标4(一般支撑)"],
        ["个人和团队", "5.2", "课程目标2(一般支撑)"],
    ],
    # 行组 T1_G1: 课程目标与题号
    "T1_G1": [
        ["课程目标1", "1,2,3,4"],
        ["课程目标2", "5,6,7"],
        ["课程目标3", "8,9,10"],
        ["课程目标4", "11,12"],
        ["课程目标5", "13"],
    ],
    # T2_G2 是复杂行组(11列)，跳过
}

# ========== 运行测试 ==========
test_template("评价报告", evaluation_data)
test_template("试卷分析", exam_data)
test_template("关联矩阵", matrix_data)

print(f"\n{'='*70}")
print("全部测试完成")
