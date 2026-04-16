"""进一步验证页面设置和表格格式是否与模板一致"""
from docx import Document

# 对比模板和输出文档
template = Document("assets/template.docx")
output = Document("/tmp/verify_output.docx")

print("=== 页面设置对比 ===")
for name, doc in [("模板", template), ("输出", output)]:
    sec = doc.sections[0]
    print(f"  {name}: 宽={sec.page_width} 高={sec.page_height}")
    print(f"  {name}: 上={sec.top_margin} 下={sec.bottom_margin} 左={sec.left_margin} 右={sec.right_margin}")

print("\n=== 表格0结构对比 ===")
for name, doc in [("模板", template), ("输出", output)]:
    t = doc.tables[0]
    print(f"  {name}: {len(t.rows)}行 x {len(t.columns)}列")

print("\n=== 表格1结构对比 ===")
for name, doc in [("模板", template), ("输出", output)]:
    t = doc.tables[1]
    print(f"  {name}: {len(t.rows)}行 x {len(t.columns)}列")

print("\n=== 表格0 行0 合并信息对比 ===")
from lxml import etree
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
for name, doc in [("模板", template), ("输出", output)]:
    t = doc.tables[0]
    row = t.rows[0]
    for c_idx, cell in enumerate(row.cells):
        tc = cell._element
        tcPr = tc.find('w:tcPr', nsmap)
        gs = tcPr.find('w:gridSpan', nsmap) if tcPr is not None else None
        if gs is not None:
            print(f"  {name} 行0列{c_idx}: gs={gs.get(nsmap['w']+'val')}")

print("\n=== 表格0 行4 (新增行) 合并信息 ===")
t = output.tables[0]
row = t.rows[4]
for c_idx, cell in enumerate(row.cells):
    tc = cell._element
    tcPr = tc.find('w:tcPr', nsmap)
    gs = tcPr.find('w:gridSpan', nsmap) if tcPr is not None else None
    if gs is not None:
        print(f"  输出 行4列{c_idx}: gs={gs.get(nsmap['w']+'val')}")

print("\n=== 表格0 行7 (新增的第4数据行) 合并信息 ===")
row = t.rows[7]
for c_idx, cell in enumerate(row.cells):
    tc = cell._element
    tcPr = tc.find('w:tcPr', nsmap)
    gs = tcPr.find('w:gridSpan', nsmap) if tcPr is not None else None
    if gs is not None:
        print(f"  输出 行7列{c_idx}: gs={gs.get(nsmap['w']+'val')}")

print("\n=== 字体一致性检查（表格0行1教学班级） ===")
for name, doc in [("模板", template), ("输出", output)]:
    t = doc.tables[0]
    cell = t.rows[1].cells[0]
    for run in cell.paragraphs[0].runs:
        rPr = run._element.rPr
        rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts') if rPr is not None else None
        east_asia = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia') if rFonts is not None else 'N/A'
        print(f"  {name}: run='{run.text}' font={run.font.name} eastAsia={east_asia} size={run.font.size}")
