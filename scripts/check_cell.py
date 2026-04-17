"""检查表格1行18-20列1的原始内容"""
from docx import Document
doc = Document("assets/template.docx")
t1 = doc.tables[1]
for r_idx in [18, 19, 20]:
    cell = t1.rows[r_idx].cells[1]
    print(f"行{r_idx}列1:")
    for p_idx, para in enumerate(cell.paragraphs):
        text = "".join(r.text for r in para.runs if r.text)
        print(f"  P{p_idx}: '{text}'")
        for run in para.runs:
            print(f"    Run: '{run.text}' bold={run.font.bold}")
