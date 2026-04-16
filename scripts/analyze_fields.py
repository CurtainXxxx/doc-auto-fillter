"""详细分析模板中所有需要用户填写的字段"""
from docx import Document

doc = Document("assets/template.docx")

print("=" * 80)
print("表格0: 基本信息表 + 课程目标与毕业要求对应关系")
print("=" * 80)
t0 = doc.tables[0]

for r_idx, row in enumerate(t0.rows):
    seen = set()
    for c_idx, cell in enumerate(row.cells):
        cell_id = id(cell._element)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        text = cell.text.replace('\n', '|').strip()
        # 标记空字段和待填字段
        if text == '':
            print(f"  行{r_idx} 列{c_idx}: [空 - 待填]")
        else:
            print(f"  行{r_idx} 列{c_idx}: '{text}'")

print("\n" + "=" * 80)
print("表格1: 评价依据 + 考核分布 + 评价结果 + 总结改进")
print("=" * 80)
t1 = doc.tables[1]

for r_idx, row in enumerate(t1.rows):
    seen = set()
    cells_data = []
    for c_idx, cell in enumerate(row.cells):
        cell_id = id(cell._element)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        text = cell.text.replace('\n', '|').strip()
        if text:
            cells_data.append(f"'{text[:40]}'")
        else:
            cells_data.append("[空]")
    print(f"  行{r_idx}: {' | '.join(cells_data)}")
