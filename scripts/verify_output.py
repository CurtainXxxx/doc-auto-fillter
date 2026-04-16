"""验证生成的文档是否正确填充了用户数据"""
from docx import Document
import io

# 重新生成一份到本地以便验证
import sys
sys.path.insert(0, '/workspace/projects/src')

from tools.edu_report_tool import _build_report_docx

doc_bytes = _build_report_docx(
    teaching_class="软件工程24级1班",
    evaluator="张明",
    participants="李华、王强、赵丽",
    course_objectives=[
        "掌握软件工程基本概念和方法",
        "能够运用面向对象分析方法进行系统建模",
        "具备软件测试与质量保证的基本能力",
        "理解软件项目管理的基本流程与方法",
    ],
)

# 保存到本地
with open("/tmp/verify_output.docx", "wb") as f:
    f.write(doc_bytes)

# 读取并验证
doc = Document("/tmp/verify_output.docx")

print("=== 标题 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"段落{i}: '{para.text}'")

print("\n=== 表格0 内容验证 ===")
t0 = doc.tables[0]
# Row 1: 教学班级 / 评价责任人 / 参与人
print(f"行1-教学班级: '{t0.rows[1].cells[0].text}'")
print(f"行1-评价责任人: '{t0.rows[1].cells[3].text}'")
print(f"行1-参与人: '{t0.rows[1].cells[6].text}'")

# Row 2: 章节标题
print(f"行2: '{t0.rows[2].cells[0].text}'")

# 数据行: 课程目标
for r_idx in range(4, min(8, len(t0.rows))):
    row = t0.rows[r_idx]
    obj_text = row.cells[4].text
    print(f"行{r_idx}-课程目标: '{obj_text}'")

print(f"\n表格0总行数: {len(t0.rows)}")

print("\n=== 表格1 关键内容验证 ===")
t1 = doc.tables[1]
# Row 0: 二、课程目标评价依据
print(f"行0: '{t1.rows[0].cells[0].text[:30]}'")
# Row 1: 考核环节/课程目标1-4
print(f"行1列0: '{t1.rows[1].cells[0].text}'")
# 评价结果行
for r_idx in [18, 21, 24, 27]:
    if r_idx < len(t1.rows):
        print(f"行{r_idx}-课程目标: '{t1.rows[r_idx].cells[0].text}'")

print(f"\n表格1总行数: {len(t1.rows)}")

# 验证格式保持
print("\n=== 格式验证 ===")
print(f"标题字体: {doc.paragraphs[0].runs[0].font.name}")
print(f"标题字号: {doc.paragraphs[0].runs[0].font.size}")
print(f"标题加粗: {doc.paragraphs[0].runs[0].font.bold}")
# 表格0行1的字体
r1_cell = t0.rows[1].cells[0]
for para in r1_cell.paragraphs:
    for run in para.runs:
        print(f"教学班级单元格-run: '{run.text}' font={run.font.name} bold={run.font.bold}")
