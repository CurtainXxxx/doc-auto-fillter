"""直接检查模板和输出行4的XML"""
from docx import Document
from lxml import etree

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

template = Document("assets/template.docx")
output = Document("/tmp/verify_output.docx")

print("=== 模板 表格0 行4 第2个独立单元格的tcPr XML ===")
t = template.tables[0]
row = t.rows[4]
# 获取第2个单元格的完整XML
seen = set()
unique_cells = []
for cell in row.cells:
    cell_id = id(cell._element)
    if cell_id not in seen:
        seen.add(cell_id)
        unique_cells.append(cell._element)

for idx, tc in enumerate(unique_cells):
    tcPr = tc.find('w:tcPr', nsmap)
    if tcPr is not None:
        print(f"  单元格{idx}:")
        # 只打印tcPr内容，去掉长命名空间
        xml_str = etree.tostring(tcPr, pretty_print=True).decode()
        for line in xml_str.split('\n')[:10]:
            # 简化命名空间
            line = line.replace('{' + nsmap['w'] + '}', 'w:')
            print(f"    {line.strip()}")

print("\n=== 输出 表格0 行4 第2个独立单元格的tcPr XML ===")
t = output.tables[0]
row = t.rows[4]
seen = set()
unique_cells = []
for cell in row.cells:
    cell_id = id(cell._element)
    if cell_id not in seen:
        seen.add(cell_id)
        unique_cells.append(cell._element)

for idx, tc in enumerate(unique_cells):
    tcPr = tc.find('w:tcPr', nsmap)
    if tcPr is not None:
        print(f"  单元格{idx}:")
        xml_str = etree.tostring(tcPr, pretty_print=True).decode()
        for line in xml_str.split('\n')[:10]:
            line = line.replace('{' + nsmap['w'] + '}', 'w:')
            print(f"    {line.strip()}")

print("\n=== 输出 表格0 行7 (新增行) 的tcPr XML ===")
row = t.rows[7]
seen = set()
unique_cells = []
for cell in row.cells:
    cell_id = id(cell._element)
    if cell_id not in seen:
        seen.add(cell_id)
        unique_cells.append(cell._element)

for idx, tc in enumerate(unique_cells):
    tcPr = tc.find('w:tcPr', nsmap)
    if tcPr is not None:
        print(f"  单元格{idx}:")
        xml_str = etree.tostring(tcPr, pretty_print=True).decode()
        for line in xml_str.split('\n')[:10]:
            line = line.replace('{' + nsmap['w'] + '}', 'w:')
            print(f"    {line.strip()}")
