"""解析模板文档的完整格式信息"""
import json
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("assets/template.docx")

# 1. 页面设置
for section in doc.sections:
    print("=" * 60)
    print("【页面设置】")
    print(f"  页面宽度: {section.page_width} EMU = {section.page_width / 914400:.2f} inches")
    print(f"  页面高度: {section.page_height} EMU = {section.page_height / 914400:.2f} inches")
    print(f"  上边距: {section.top_margin} EMU = {section.top_margin / 914400:.2f} inches = {section.top_margin / 360000:.2f} cm")
    print(f"  下边距: {section.bottom_margin} EMU = {section.bottom_margin / 914400:.2f} inches = {section.bottom_margin / 360000:.2f} cm")
    print(f"  左边距: {section.left_margin} EMU = {section.left_margin / 914400:.2f} inches = {section.left_margin / 360000:.2f} cm")
    print(f"  右边距: {section.right_margin} EMU = {section.right_margin / 914400:.2f} inches = {section.right_margin / 360000:.2f} cm")
    print(f"  页眉距离: {section.header_distance} EMU")
    print(f"  页脚距离: {section.footer_distance} EMU")

# 2. 所有段落
print("\n" + "=" * 60)
print("【段落详细分析】")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "" and len(para.runs) == 0:
        print(f"\n--- 段落 {i}: [空段落] ---")
        pf = para.paragraph_format
        print(f"  对齐: {pf.alignment}")
        print(f"  段前间距: {pf.space_before}")
        print(f"  段后间距: {pf.space_after}")
        print(f"  行距: {pf.line_spacing} (规则: {pf.line_spacing_rule})")
        continue
    
    print(f"\n--- 段落 {i}: \"{para.text[:80]}\" ---")
    pf = para.paragraph_format
    print(f"  对齐: {pf.alignment}")
    print(f"  段前间距: {pf.space_before}")
    print(f"  段后间距: {pf.space_after}")
    print(f"  行距: {pf.line_spacing} (规则: {pf.line_spacing_rule})")
    print(f"  首行缩进: {pf.first_line_indent}")
    print(f"  左缩进: {pf.left_indent}")
    
    style_name = para.style.name if para.style else "None"
    print(f"  样式名: {style_name}")
    
    for j, run in enumerate(para.runs):
        text_preview = run.text[:50].replace('\n', '\\n')
        print(f"  Run {j}: \"{text_preview}\"")
        font = run.font
        print(f"    字体名: {font.name}")
        print(f"    字号: {font.size} ({Pt(font.size.pt).pt if font.size else 'None'}pt)" if font.size else f"    字号: None")
        print(f"    加粗: {font.bold}")
        print(f"    斜体: {font.italic}")
        print(f"    下划线: {font.underline}")
        print(f"    颜色RGB: {font.color.rgb if font.color and font.color.rgb else 'None'}")
        # 检查东亚字体
        rPr = run._element.rPr
        if rPr is not None:
            rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
            if rFonts is not None:
                print(f"    东亚字体(eastAsia): {rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')}")
                print(f"    ASCII字体: {rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')}")
                print(f"    hAnsi字体: {rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi')}")

# 3. 所有表格
print("\n" + "=" * 60)
print("【表格详细分析】")
for t_idx, table in enumerate(doc.tables):
    print(f"\n=== 表格 {t_idx} ({len(table.rows)}行 x {len(table.columns)}列) ===")
    
    # 表格样式
    print(f"  表格样式: {table.style.name if table.style else 'None'}")
    
    # 表格宽度
    tbl_pr = table._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
    if tbl_pr is not None:
        tbl_w = tbl_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
        if tbl_w is not None:
            print(f"  表格宽度: {tbl_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')} {tbl_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')}")
        jc = tbl_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc')
        if jc is not None:
            print(f"  表格对齐: {jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')}")
    
    # 每行每列
    for r_idx, row in enumerate(table.rows):
        row_height = row.height
        print(f"\n  行 {r_idx} (高度: {row_height}):")
        for c_idx, cell in enumerate(row.cells):
            # 检查合并
            tc_pr = cell._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
            merge_info = ""
            if tc_pr is not None:
                v_merge = tc_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                if v_merge is not None:
                    val = v_merge.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    merge_info = f" [vMerge: {'continue' if val is None else val}]"
                grid_span = tc_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
                if grid_span is not None:
                    merge_info += f" [gridSpan: {grid_span.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')}]"
                tc_w = tc_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
                if tc_w is not None:
                    merge_info += f" [宽度: {tc_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')} {tc_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')}]"
            
            cell_text = cell.text.replace('\n', '\\n')[:60]
            print(f"    列{c_idx}: \"{cell_text}\"{merge_info}")
            
            # 单元格内段落和run详情
            for p_idx, para in enumerate(cell.paragraphs):
                if para.text.strip():
                    pf = para.paragraph_format
                    print(f"      P{p_idx}: align={pf.alignment}, line_spacing={pf.line_spacing}")
                    for run in para.runs:
                        if run.text.strip():
                            font = run.font
                            east_asia = ""
                            rPr = run._element.rPr
                            if rPr is not None:
                                rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                                if rFonts is not None:
                                    east_asia = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '')
                            print(f"        Run: \"{run.text[:40]}\" font={font.name} size={font.size} bold={font.bold} eastAsia={east_asia}")

# 4. 页眉页脚
print("\n" + "=" * 60)
print("【页眉页脚】")
for section in doc.sections:
    header = section.header
    if header and header.paragraphs:
        for p in header.paragraphs:
            if p.text.strip():
                print(f"  页眉: \"{p.text[:80]}\"")
                for run in p.runs:
                    print(f"    Run: \"{run.text}\" font={run.font.name} size={run.font.size} bold={run.font.bold}")
    footer = section.footer
    if footer and footer.paragraphs:
        for p in footer.paragraphs:
            if p.text.strip():
                print(f"  页脚: \"{p.text[:80]}\"")
