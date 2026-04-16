"""精炼解析模板结构 - 提取表格骨架和合并信息"""
from docx import Document
from lxml import etree

doc = Document("assets/template.docx")

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# 提取每个表格的简化结构
for t_idx, table in enumerate(doc.tables):
    rows = table.rows
    cols_count = len(table.columns)
    print(f"\n{'='*80}")
    print(f"表格 {t_idx}: {len(rows)}行 x {cols_count}列")
    
    for r_idx, row in enumerate(rows):
        cells_info = []
        seen = set()
        for c_idx, cell in enumerate(row.cells):
            # 跳过被合并的重复cell
            cell_id = id(cell._element)
            if cell_id in seen:
                continue
            seen.add(cell_id)
            
            tc = cell._element
            tcPr = tc.find('w:tcPr', nsmap)
            
            grid_span = 1
            v_merge_val = None
            width_val = None
            
            if tcPr is not None:
                gs = tcPr.find('w:gridSpan', nsmap)
                if gs is not None:
                    grid_span = int(gs.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '1'))
                vm = tcPr.find('w:vMerge', nsmap)
                if vm is not None:
                    v_merge_val = vm.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'CONTINUE')
                tw = tcPr.find('w:tcW', nsmap)
                if tw is not None:
                    width_val = tw.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
            
            # 获取文本（简化）
            text = cell.text.replace('\n', '|')[:80]
            merge_str = f"gs={grid_span}" if grid_span > 1 else ""
            if v_merge_val is not None:
                merge_str += f" vm={v_merge_val}"
            if width_val:
                merge_str += f" w={width_val}"
            
            info = f'"{text}"'
            if merge_str:
                info += f" [{merge_str}]"
            cells_info.append(info)
        
        print(f"  行{r_idx}: {'; '.join(cells_info)}")

# 解析表格的gridCols
print(f"\n{'='*80}")
print("表格Grid列宽定义:")
for t_idx, table in enumerate(doc.tables):
    tbl = table._element
    grid = tbl.find('w:tblGrid', nsmap)
    if grid is not None:
        cols = grid.findall('w:gridCol', nsmap)
        widths = [c.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '?') for c in cols]
        print(f"  表格{t_idx}: {len(cols)}列, 宽度={widths}")
