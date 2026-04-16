"""深入解析模板XML以获取精确的格式信息"""
from docx import Document
from lxml import etree

doc = Document("assets/template.docx")

# 1. 获取默认段落样式中的字号
style = doc.styles["Normal"]
print(f"Normal样式字号: {style.font.size}")
print(f"Normal样式字体: {style.font.name}")

# 2. 获取表格0行5的详细XML（考试类别那行，包含换行）
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

print("\n\n=== 表格0 行0列5 详细XML (考试类别) ===")
table0 = doc.tables[0]
cell = table0.rows[0].cells[5]
xml_str = etree.tostring(cell._element, pretty_print=True).decode()
# 只打印段落部分
for p_elem in cell._element.findall('.//w:p', nsmap):
    print(etree.tostring(p_elem, pretty_print=True).decode()[:500])

# 3. 表格0行1列3 详细XML（评价责任人）
print("\n=== 表格0 行1列3 详细XML (评价责任人) ===")
cell = table0.rows[1].cells[3]
for p_elem in cell._element.findall('.//w:p', nsmap):
    print(etree.tostring(p_elem, pretty_print=True).decode()[:500])

# 4. 表格0 行3 详细XML（表头行：毕业要求/毕业要求指标点/课程目标）
print("\n=== 表格0 行3 详细XML ===")
for c_idx in [0, 1, 3]:
    cell = table0.rows[3].cells[c_idx]
    for p_elem in cell._element.findall('.//w:p', nsmap):
        print(f"列{c_idx}:")
        print(etree.tostring(p_elem, pretty_print=True).decode()[:600])

# 5. 表格1行1的详细（考核环节/课程目标1-4）
print("\n=== 表格1 行1 详细XML (考核环节头) ===")
table1 = doc.tables[1]
for c_idx in [0, 1, 2, 3, 4]:
    cell = table1.rows[1].cells[c_idx]
    tc = cell._element
    tcPr = tc.find('w:tcPr', nsmap)
    # 只看gridSpan
    gs = tcPr.find('w:gridSpan', nsmap) if tcPr is not None else None
    print(f"  列{c_idx}: text='{cell.text[:30]}' gs={gs.get(nsmap['w']+'val') if gs is not None else '1'}")

# 6. 表格1行17（表头行：课程目标/实现途径/目标分值/实际平均分/目标达成评价值）
print("\n=== 表格1 行17 详细XML ===")
for c_idx in [0, 1, 2, 3, 4]:
    cell = table1.rows[17].cells[c_idx]
    for p_elem in cell._element.findall('.//w:p', nsmap):
        print(f"  列{c_idx}:")
        print(etree.tostring(p_elem, pretty_print=True).decode()[:500])

# 7. 表格1行18（课程目标1的期末考试行）
print("\n=== 表格1 行18 详细XML ===")
for c_idx in [0, 2]:
    cell = table1.rows[18].cells[c_idx]
    for p_elem in cell._element.findall('.//w:p', nsmap):
        print(f"  列{c_idx}:")
        print(etree.tostring(p_elem, pretty_print=True).decode()[:500])

# 8. 表格1行33（课程总结与改进）
print("\n=== 表格1 行33 详细XML ===")
cell = table1.rows[33].cells[0]
for p_idx, p_elem in enumerate(cell._element.findall('.//w:p', nsmap)):
    text = ""
    for r in p_elem.findall('.//w:r', nsmap):
        t = r.find('w:t', nsmap)
        if t is not None and t.text:
            text += t.text
    if text.strip():
        print(f"  P{p_idx}: '{text}'")

# 9. 获取表格边框设置
print("\n=== 表格0 边框设置 ===")
tbl_pr = table0._element.find('w:tblPr', nsmap)
borders = tbl_pr.find('w:tblBorders', nsmap) if tbl_pr is not None else None
if borders is not None:
    print(etree.tostring(borders, pretty_print=True).decode()[:1000])

print("\n=== 表格1 边框设置 ===")
tbl_pr = table1._element.find('w:tblPr', nsmap)
borders = tbl_pr.find('w:tblBorders', nsmap) if tbl_pr is not None else None
if borders is not None:
    print(etree.tostring(borders, pretty_print=True).decode()[:1000])
