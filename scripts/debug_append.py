"""调试_append_after_label在行19格1的行为"""
import sys
sys.path.insert(0, '/workspace/projects/src')
from docx import Document
from tools.edu_report_tool import _append_after_label, _get_unique_cells
import copy
from docx.oxml.ns import qn

# 先对行18做追加
doc = Document("assets/template.docx")
t = doc.tables[1]
unique18 = _get_unique_cells(t.rows[18])
cell18 = unique18[1]  # 实现途径/评价方法 格
print(f"行18格1追加前:")
for r in cell18.paragraphs[0].runs:
    print(f"  run: '{r.text}'")

_append_after_label(cell18, "实现途径", "闭卷笔试")
_append_after_label(cell18, "评价方法", "百分制")
print(f"\n行18格1追加后:")
for r in cell18.paragraphs[0].runs:
    print(f"  run: '{r.text}'")

# 再对行19做追加
unique19 = _get_unique_cells(t.rows[19])
cell19 = unique19[1]
print(f"\n行19格1追加前:")
for r in cell19.paragraphs[0].runs:
    print(f"  run: '{r.text}'")

_append_after_label(cell19, "实现途径", "课堂作业")
_append_after_label(cell19, "评价方法", "等级制")
print(f"\n行19格1追加后:")
for r in cell19.paragraphs[0].runs:
    print(f"  run: '{r.text}'")
