"""验证合并单元格在新行中是否正确保留"""
from docx import Document
from lxml import etree

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

template = Document("assets/template.docx")
output = Document("/tmp/verify_output.docx")

print("=== 模板 表格0 行4-6 每个独立单元格的gridSpan ===")
t = template.tables[0]
for r_idx in [4, 5, 6]:
    row = t.rows[r_idx]
    seen = set()
    col_pos = 0
    for c_idx, cell in enumerate(row.cells):
        cell_id = id(cell._element)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        tc = cell._element
        tcPr = tc.find('w:tcPr', nsmap)
        gs_val = "1"
        if tcPr is not None:
            gs = tcPr.find('w:gridSpan', nsmap)
            if gs is not None:
                gs_val = gs.get(nsmap['w']+'val')
        print(f"  行{r_idx} 独立单元格{col_pos}: text='{cell.text[:30]}' gridSpan={gs_val}")
        col_pos += 1

print("\n=== 输出 表格0 行4-7 每个独立单元格的gridSpan ===")
t = output.tables[0]
for r_idx in [4, 5, 6, 7]:
    row = t.rows[r_idx]
    seen = set()
    col_pos = 0
    for c_idx, cell in enumerate(row.cells):
        cell_id = id(cell._element)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        tc = cell._element
        tcPr = tc.find('w:tcPr', nsmap)
        gs_val = "1"
        if tcPr is not None:
            gs = tcPr.find('w:gridSpan', nsmap)
            if gs is not None:
                gs_val = gs.get(nsmap['w']+'val')
        print(f"  行{r_idx} 独立单元格{col_pos}: text='{cell.text[:30]}' gridSpan={gs_val}")
        col_pos += 1

print("\n=== 输出 表格1 行18-29 课程目标列验证 ===")
t1 = output.tables[1]
for r_idx in [18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29]:
    cell = t1.rows[r_idx].cells[0]
    print(f"  行{r_idx}: '{cell.text[:30]}'")
