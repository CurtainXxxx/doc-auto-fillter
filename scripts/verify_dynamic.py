"""快速验证动态填充后的文档"""
import sys
sys.path.insert(0, '/workspace/projects/src')
from tools.template_analyzer import analyze_template
from tools.edu_report_tool import _build_report_docx
import json

data = {
    "课程名称": "《软件工程导论》",
    "开课时间": "2023-2024-2",
    "考试类别": "闭卷考试",
    "平时": "30%",
    "期末": "70%",
    "参评人数": "45",
    "教学班级": "软件工程24级1班",
    "评价责任人": "张明",
    "参与人": "李华、王强、赵丽",
    "实现途径": ["闭卷笔试","课堂作业","实验操作"]*4,
    "评价方法": ["百分制","等级制","实操考核"]*4,
    "课程总结": "本学期课程教学总体达到预期目标",
    "改进措施": "1.增加软件测试实验课时",
    "T0_G0": [
        ["工程知识","1.3","课程目标1：掌握软件工程基本概念和方法"],
        ["问题分析","2.1","课程目标2：能够运用面向对象分析方法"],
        ["设计开发解决方案","3.2","课程目标3：具备软件测试基本能力"],
        ["使用现代工具","4.1","课程目标4：理解软件项目管理"],
    ],
    "T1_G1": [
        ["期末考试","0.7","0.6","0.5","0.4"],
        ["平时成绩","0.2","0.2","0.3","0.3"],
        ["实验成绩","0.1","0.2","0.2","0.3"],
    ],
}

doc_bytes = _build_report_docx(data)
with open("/tmp/verify_dynamic.docx", "wb") as f:
    f.write(doc_bytes)

from docx import Document
doc = Document("/tmp/verify_dynamic.docx")

# 验证表格0
t0 = doc.tables[0]
print("=== 表格0 验证 ===")
print(f"行数: {len(t0.rows)} (模板7行, 新增1行→8行)")
# 基本信息
r0_texts = [c.text.strip().replace('\n','|')[:40] for c in t0.rows[0].cells if c.text.strip()]
print(f"行0: {r0_texts}")
r1_texts = [c.text.strip().replace('\n','|')[:40] for c in t0.rows[1].cells if c.text.strip()]
print(f"行1: {r1_texts}")

# 课程目标行
for r_idx in range(4, len(t0.rows)):
    unique = []
    seen = set()
    for c in t0.rows[r_idx].cells:
        cid = id(c._element)
        if cid not in seen:
            seen.add(cid)
            unique.append(c.text.strip()[:50])
    print(f"行{r_idx}: {unique}")

# 验证表格1
t1 = doc.tables[1]
print(f"\n=== 表格1 验证 ===")
# 评价依据行
for r_idx in [2,3,4]:
    unique = []
    seen = set()
    for c in t1.rows[r_idx].cells:
        cid = id(c._element)
        if cid not in seen:
            seen.add(cid)
            unique.append(c.text.strip()[:30])
    print(f"行{r_idx}: {unique}")

# 评价结果 - 实现途径/评价方法
for r_idx in [18, 19, 20]:
    cell1 = t1.rows[r_idx].cells[1]
    print(f"行{r_idx}列1: '{cell1.text[:60]}'")

# 课程总结/改进
cell33 = t1.rows[33].cells[0]
print(f"行33: '{cell33.text[:100]}'")

print(f"\n文件大小: {len(doc_bytes)} bytes")
print("验证完成!")
